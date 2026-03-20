#!/usr/bin/env python3
"""
common.py — Module partagé pour tous les scripts build v4
Contient : chargement CSV, graphiques, utilitaires HTML/DOCX, CSS global.
"""
import csv
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D
import matplotlib.ticker
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─── Chemins ────────────────────────────────────────────────
# Ce module est dans Formants/Scripts/v4-html-docx-enriched/
# BASE = Formants/
BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUT_DIR = os.path.join(BASE, "Versions-html-and-docx")
OUT_IMG = os.path.join(OUT_DIR, "media")
os.makedirs(OUT_IMG, exist_ok=True)

# ─── Chargement CSV ─────────────────────────────────────────
NAME_MAPPING = {
    'Bass_Clarinet_Bb': 'Clarinette basse en Sib',
    'Bass_Flute': 'Flûte basse',
    'Bass_Trombone': 'Trombone basse',
    'Bass_Tuba': 'Tuba basse',
    'Bass_Tuba+sordina': 'Tuba basse+sourdine',
    'Bassoon': 'Basson',
    'Bassoon+sordina': 'Basson+sourdine',
    'Clarinet_Bb': 'Clarinette en Sib',
    'Clarinet_Eb': 'Clarinette en Mib',
    'Contrabass': 'Contrebasse',
    'Contrabass+sordina': 'Contrebasse+sourdine',
    'Contrabass_Clarinet_Bb': 'Clarinette contrebasse en Sib',
    'Contrabass_Ensemble': 'Ensemble de contrebasses',
    'Contrabass_Flute': 'Flûte contrebasse',
    'Contrabass_Tuba': 'Tuba contrebasse',
    'Contrabassoon': 'Contrebasson',
    'English_Horn': 'Cor anglais',
    'Flute': 'Flûte',
    'Horn': 'Cor',
    'Horn+sordina': 'Cor+sourdine',
    'Oboe': 'Hautbois',
    'Oboe+sordina': 'Hautbois+sourdine',
    'Piccolo': 'Petite flûte',
    'Sax_Alto': 'Saxophone alto',
    'Trombone': 'Trombone',
    'Trombone+sordina_cup': 'Trombone+sourdine bol',
    'Trombone+sordina_harmon': 'Trombone+sourdine harmon',
    'Trombone+sordina_straight': 'Trombone+sourdine sèche',
    'Trombone+sordina_wah': 'Trombone+sourdine wah',
    'Trumpet_C': 'Trompette en Ut',
    'Trumpet_C+sordina_cup': 'Trompette en Ut+sourdine bol',
    'Trumpet_C+sordina_harmon': 'Trompette en Ut+sourdine harmon',
    'Trumpet_C+sordina_straight': 'Trompette en Ut+sourdine sèche',
    'Trumpet_C+sordina_wah': 'Trompette en Ut+sourdine wah',
    'Viola': 'Alto',
    'Viola+sordina': 'Alto+sourdine',
    'Viola+sordina_piombo': 'Alto+sourdine piombo',
    'Viola_Ensemble': "Ensemble d'altos",
    'Viola_Ensemble+sordina': "Ensemble d'altos+sourdine",
    'Violin': 'Violon',
    'Violin+sordina': 'Violon+sourdine',
    'Violin+sordina_piombo': 'Violon+sourdine piombo',
    'Violin_Ensemble': 'Ensemble de violons',
    'Violin_Ensemble+sordina': 'Ensemble de violons+sourdine',
    'Violoncello': 'Violoncelle',
    'Violoncello+sordina': 'Violoncelle+sourdine',
    'Violoncello+sordina_piombo': 'Violoncelle+sourdine piombo',
    'Violoncello_Ensemble': 'Ensemble de violoncelles',
    'Violoncello_Ensemble+sordina': 'Ensemble de violoncelles+sourdine',
}
V2_TO_LEGACY = {v: k for k, v in NAME_MAPPING.items()}

DATA = {}

def sf(v):
    try:
        return float(v)
    except Exception:
        return 0.0

def load_data(path):
    if not os.path.exists(path):
        print(f"  ⚠ CSV introuvable : {path}")
        return
    with open(path, 'r', encoding='utf-8') as f:
        for row in csv.DictReader(f):
            inst = row['instrument']
            tech = row['technique']
            DATA[(inst, tech)] = row
            if inst in NAME_MAPPING:
                DATA[(NAME_MAPPING[inst], tech)] = row
            if inst in V2_TO_LEGACY:
                DATA[(V2_TO_LEGACY[inst], tech)] = row

def load_all_csvs():
    csv_all = os.path.join(BASE, 'Resultats/formants_all_techniques_v2.csv')
    csv_yan = os.path.join(BASE, 'Resultats/formants_yan_adds_v2.csv')
    if not os.path.exists(csv_all):
        csv_all = os.path.join(BASE, 'Resultats/formants_all_techniques.csv')
    if not os.path.exists(csv_yan):
        csv_yan = os.path.join(BASE, 'Resultats/formants_yan_adds.csv')
    load_data(csv_all)
    load_data(csv_yan)

def get_f(inst, tech):
    r = DATA.get((inst, tech))
    if not r:
        return None
    return {
        'n': int(r['n_samples']),
        'F': [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)],
    }

def fmt_hz(v):
    return '—' if v == 0 else f"{int(v):,}".replace(',', '\u202f')

# ─── Zones vocaliques ───────────────────────────────────────
VOWEL_ZONES = [
    (100,  400,  '#DCEEFB', 'u (oo)\nProfondeur'),
    (400,  600,  '#D5ECD5', 'o (oh)\nPlénitude'),
    (600,  800,  '#FDE8CE', 'å (aw)\nTransition'),
    (800,  1250, '#F8D5D5', 'a (ah)\nPuissance'),
    (1250, 2600, '#E8D5F0', 'e (eh)\nClarté'),
    (2600, 6000, '#FFF8D0', 'i (ee)\nBrillance'),
]
FC = ['#D32F2F', '#E64A19', '#F57C00', '#FFA000', '#FBC02D', '#CDDC39']
FA = [1.0, 0.85, 0.7, 0.55, 0.4, 0.3]

# ─── Génération graphiques ───────────────────────────────────
def make_graph(display, filename, n, formants, fp=None, family_color='#2E7D32', family_label=''):
    valid = [(i, f) for i, f in enumerate(formants) if f > 0]
    if not valid:
        return None

    mf = max(f for _, f in valid) + 500
    mf = min(max(mf, 3000), 6500)

    fig, ax = plt.subplots(figsize=(9.6, 4.8), dpi=150)

    for lo, hi, c, l in VOWEL_ZONES:
        if lo < mf:
            ax.axvspan(lo, min(hi, mf), alpha=0.35, color=c, zorder=0)
            mid = (lo + min(hi, mf)) / 2
            if mid < mf * 0.95:
                ax.text(mid, 0.97, l, ha='center', va='top', fontsize=7,
                        color='#666', fontweight='bold',
                        transform=ax.get_xaxis_transform())

    bw = mf * 0.012
    for i, freq in valid:
        bh = 1.0 * (1.0 - i * 0.12)
        ax.bar(freq, bh, width=bw * (1.2 - i * 0.05),
               color=FC[i], alpha=FA[i], edgecolor='#333', linewidth=0.8, zorder=3)
        ax.text(freq, bh + 0.03, f"F{i+1}\n{freq} Hz",
                ha='center', va='bottom', fontsize=8,
                fontweight='bold', color='#333', zorder=5)

    f1 = formants[0]
    if fp and abs(fp - f1) > 30:
        ax.plot(fp, 0.5, marker='D', markersize=14, color='#1B5E20',
                markeredgecolor='black', markeredgewidth=1.5, zorder=6)
        ax.annotate(f"Fp = {fp} Hz\n(centroïde)", xy=(fp, 0.5), xytext=(fp, 0.65),
                    ha='center', fontsize=8, fontweight='bold', color='#1B5E20',
                    arrowprops=dict(arrowstyle='->', color='#1B5E20', lw=1.5), zorder=7)

    # Cluster de convergence
    ax.axvspan(420, 550, alpha=0.12, color='red', zorder=1)
    ax.text(485, 0.02, 'cluster /o/', ha='center', va='bottom', fontsize=7,
            color='#C62828', fontstyle='italic',
            transform=ax.get_xaxis_transform())

    ax.set_xlim(100, mf)
    ax.set_ylim(0, 1.25)
    ax.set_xlabel("Fréquence (Hz)", fontsize=10, fontweight='bold')
    ax.set_ylabel("Importance relative du formant", fontsize=10, fontweight='bold')
    ax.set_title(f"{display} — Formants spectraux F1–F6 (ordinario, N={n})",
                 fontsize=12, fontweight='bold', color=family_color, pad=12)
    ax.set_xscale('log')
    ticks = [t for t in [100, 150, 200, 300, 400, 500, 600, 800,
                          1000, 1500, 2000, 3000, 4000, 5000, 6000] if t <= mf]
    ax.set_xticks(ticks)
    ax.set_xticklabels([str(t) for t in ticks], fontsize=8)
    ax.get_xaxis().set_major_formatter(matplotlib.ticker.ScalarFormatter())
    ax.set_yticks([])
    for s in ['top', 'right', 'left']:
        ax.spines[s].set_visible(False)

    le = [mpatches.Patch(facecolor=FC[i], alpha=FA[i], edgecolor='#333',
                         label=f'F{i+1} = {formants[i]} Hz') for i, _ in valid]
    if fp and abs(fp - f1) > 30:
        le.append(Line2D([0], [0], marker='D', color='w',
                         markerfacecolor='#1B5E20', markeredgecolor='black',
                         markersize=10, label=f'Fp centroïde = {fp} Hz'))
    ax.legend(handles=le, loc='upper right', fontsize=7, framealpha=0.9, edgecolor='#CCC')
    ax.text(0.01, -0.08,
            f"Famille : {family_label or 'Orchestre'}\nSource : CSV v22 (SOL2020 + Yan_Adds)",
            transform=ax.transAxes, fontsize=7, color='#888')
    plt.tight_layout()

    out = os.path.join(OUT_IMG, f"{filename}.png")
    fig.savefig(out, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return out


# ─── Calcul Fp depuis specenv brut ───────────────────────────
SAMPLE_RATE = 44100
FFT_SIZE    = 4096
FREQ_RES    = SAMPLE_RATE / FFT_SIZE   # ~10.77 Hz/bin

# Bandes de centroïde par instrument (même que build2v2_envelopes)
FP_BANDS_SPECENV = {
    'Horn':                  (600,  1400),
    'Horn+sordina':          (600,  1400),
    'Trumpet_C':             (600,  1400),
    'Trumpet_C+sordina_cup':     (600, 1400),
    'Trumpet_C+sordina_straight':(600, 1400),
    'Trumpet_C+sordina_harmon':  (800, 2000),
    'Trumpet_C+sordina_wah':     (600, 1400),
    'Trombone':              (1000, 2000),
    'Trombone+sordina_cup':      (1000, 2000),
    'Trombone+sordina_straight': (1000, 2000),
    'Trombone+sordina_harmon':   (800,  2000),
    'Trombone+sordina_wah':      (1000, 2000),
    'Bass_Tuba':             (1000, 2000),
    'Bass_Tuba+sordina':     (1000, 2000),
    'Bass_Trombone':         (1000, 2000),
    'Contrabass_Tuba':       (1000, 2000),
}

SOL_DIR = os.path.join(BASE, 'Data', 'FullSOL2020_specenv par instrument')
YAN_DIR = os.path.join(BASE, 'Data', 'Yan_Adds-Divers_specenv par instrument')

def _band_centroid(env_vals, lo_hz, hi_hz):
    """Centroïde spectral pondéré en énergie dans une bande Hz."""
    lo_bin = max(0, int(lo_hz / FREQ_RES))
    hi_bin = min(len(env_vals) - 1, int(hi_hz / FREQ_RES))
    region = env_vals[lo_bin:hi_bin + 1]
    if not region:
        return None
    linear = [10 ** (v / 10) for v in region]
    total  = sum(linear)
    if total <= 0:
        return None
    return sum(linear[i] * (lo_bin + i) * FREQ_RES for i in range(len(linear))) / total

def compute_fp_from_specenv(instrument_key, techs=('ordinario',),
                             specenv_dir=None, specenv_filename=None):
    """
    Calcule le Fp (centroïde spectral) directement depuis le fichier specenv brut.
    
    instrument_key  : ex. 'Trombone+sordina_cup', 'Horn+sordina'
    techs           : tuple de techniques à inclure, ex. ('ordinario',) ou ('ordinario_open',)
    specenv_dir     : répertoire du fichier (défaut : SOL_DIR)
    specenv_filename: nom du fichier (défaut : auto-construit depuis instrument_key)
    
    Retourne le Fp médian arrondi à l'entier, ou None si pas de données.
    """
    if specenv_dir is None:
        specenv_dir = SOL_DIR
    if specenv_filename is None:
        specenv_filename = f"FullSOL2020_specenv.db_{instrument_key}.txt"

    filepath = os.path.join(specenv_dir, specenv_filename)
    if not os.path.exists(filepath):
        return None

    # Bande Fp pour cet instrument
    band = FP_BANDS_SPECENV.get(instrument_key, (800, 1800))
    lo_hz, hi_hz = band

    fps = []
    try:
        with open(filepath, encoding='utf-8') as f:
            f.readline()  # en-tête
            for line in f:
                parts = line.strip().split(';')
                if len(parts) < 10:
                    continue
                path = parts[0]
                if not path.startswith('/'):
                    continue
                # Filtrer par technique (champ [3] du chemin /Famille/Instrument/technique/...)
                path_parts = path.split('/')
                if len(path_parts) < 5:
                    continue
                tech = path_parts[3]
                if tech not in techs:
                    continue
                try:
                    vals = [float(v) for v in parts[1:]]
                except ValueError:
                    continue
                if len(vals) < 100:
                    continue
                # Les valeurs specenv sont en dB (négatives) — on vérifie juste
                # que le tableau n'est pas vide et non constant
                max_amp = max(vals)
                min_amp = min(vals)
                if max_amp == min_amp:
                    continue   # spectre plat = données invalides
                fp = _band_centroid(vals, lo_hz, hi_hz)
                if fp:
                    fps.append(fp)
    except Exception:
        return None

    if not fps:
        return None

    fps.sort()
    median = fps[len(fps) // 2]
    return round(median)

# ─── Tableaux techniques HTML ────────────────────────────────
SUSTAINED_TECHS = {
    'ordinario', 'non-vibrato', 'non_vibrato', 'vibrato',
    'ordinario_open', 'ordinario_closed',
    'flautando', 'arco',
}
EXCLUDED_TECHS = {
    'pizzicato_secco', 'pizzicato_bartok', 'pizzicato_l_vib', 'pizzicato-l-vib',
    'col_legno_battuto', 'col_legno_tratto', 'tremolo', 'artificial_harmonic',
    'artificial_harmonic_tremolo', 'behind_the_bridge', 'behind_the_fingerboard',
    'on_the_bridge', 'on_the_tailpiece', 'on_the_tuning_pegs', 'on_the_frog',
    'hit_on_body', 'col_legno', 'natural_harmonics_glissandi',
}

def tech_table_html(inst_csv, filter_sustained=False):
    techs = sorted([(t, r) for (i, t), r in DATA.items() if i == inst_csv],
                   key=lambda x: x[0])
    if not techs:
        return ""
    rows = []
    for tech, r in techs:
        if filter_sustained and tech in EXCLUDED_TECHS:
            continue
        n = r['n_samples']
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        is_ord = any(s in tech.lower() for s in ('ordinario', 'non-vibrato', 'non_vibrato'))
        bg = ' style="background-color:#dff0d8;"' if is_ord else ''
        fvals = ''.join(f'<td{bg}>{fmt_hz(f)}</td>' for f in fs)
        rows.append(f'<tr><td{bg}><b>{tech}</b></td><td{bg}>{n}</td>{fvals}</tr>')
    if not rows:
        return ""
    return (f'<table class="tech-table">'
            f'<tr class="header"><th>Technique</th><th>N</th>'
            f'<th>F1</th><th>F2</th><th>F3</th><th>F4</th><th>F5</th><th>F6</th></tr>'
            + ''.join(rows) + '</table>')

# ─── Tableaux références sources HTML ───────────────────────
def ref_table_html(rows_data):
    """
    rows_data : liste de dicts avec clés :
      source, f1, f2, f3, f4, voyelle, n, accord, note
    """
    html = (
        '<table class="ref-table">'
        '<tr class="header">'
        '<th>Source</th><th>F1 (Hz)</th><th>F2 (Hz)</th>'
        '<th>F3 (Hz)</th><th>F4 (Hz)</th>'
        '<th>Voyelle</th><th>N</th><th>Accord</th>'
        '</tr>'
    )
    for r in rows_data:
        src = r.get('source', '—')
        is_sol = 'SOL' in src or 'Yan' in src
        bg = ' style="background-color:#dff0d8;"' if is_sol else ''
        accord = r.get('accord', '')
        if accord == 'oui':
            accord = '✓'
        elif accord == 'non':
            accord = '✗'
        elif accord == 'approx':
            accord = '~'
        html += (
            f'<tr>'
            f'<td{bg}><b>{src}</b></td>'
            f'<td{bg}>{r.get("f1","—")}</td>'
            f'<td{bg}>{r.get("f2","—")}</td>'
            f'<td{bg}>{r.get("f3","—")}</td>'
            f'<td{bg}>{r.get("f4","—")}</td>'
            f'<td{bg}>{r.get("voyelle","—")}</td>'
            f'<td{bg}>{r.get("n","—")}</td>'
            f'<td{bg}>{accord}</td>'
            f'</tr>'
        )
    html += '</table>'
    return html

# ─── Section doublures HTML ──────────────────────────────────
def doublures_html(items):
    """
    items : liste de dicts avec clés :
      instr, f1_a, f1_b, delta, quality, rapport, note
      rapport : 'Unisson' | 'Octave' | '2 octaves' | 'Complémentaire' | etc.
    """
    if not items:
        return ""
    rows = ""
    for it in items:
        delta = it.get('delta', '')
        if isinstance(delta, (int, float)):
            delta_str = f"Δ={delta} Hz"
        else:
            delta_str = str(delta)
        quality = it.get('quality', '')
        rapport = it.get('rapport', '—')
        # Colorier le rapport
        rapport_color = {
            'Unisson':       '#1B5E20',
            'Octave':        '#1565C0',
            '2 octaves':     '#4A148C',
            'Complémentaire':'#E65100',
        }.get(rapport, '#555')
        rapport_html = f'<span style="color:{rapport_color};font-weight:bold;">{rapport}</span>'
        rows += (
            f'<tr>'
            f'<td><b>{it.get("instr","")}</b></td>'
            f'<td>{it.get("f1_a","—")}</td>'
            f'<td>{it.get("f1_b","—")}</td>'
            f'<td>{delta_str}</td>'
            f'<td>{quality}</td>'
            f'<td>{rapport_html}</td>'
            f'<td>{it.get("note","")}</td>'
            f'</tr>'
        )
    return (
        '<div class="doublures-box">'
        '<h4>Doublures recommandées</h4>'
        '<table class="doublures-table">'
        '<tr class="header">'
        '<th>Instrument associé</th><th>F1 (instrument)</th>'
        '<th>F1 (associé)</th><th>Écart</th><th>Qualité</th>'
        '<th>Rapport de tessiture</th><th>Commentaire</th>'
        '</tr>'
        + rows +
        '</table></div>'
    )

# ─── CSS global ──────────────────────────────────────────────
CSS_GLOBAL = """
body {
  font-family: 'Segoe UI', Helvetica, Arial, sans-serif;
  max-width: 1100px; margin: 0 auto; padding: 24px;
  background: #f5f5f5; color: #333; line-height: 1.6;
}
h1 { color: #1a237e; border-bottom: 3px solid #2e7d32; padding-bottom: 10px; font-size: 1.9em; }
h2 { color: #2e7d32; margin-top: 50px; border-left: 5px solid #2e7d32;
     padding-left: 14px; font-size: 1.4em; }
h3 { color: #1b5e20; margin-top: 36px; font-size: 1.2em; }
h4 { color: #555; margin-top: 16px; font-size: 1.05em; }
.instrument-card {
  background: white; border: 1px solid #ddd; border-radius: 8px;
  padding: 22px; margin: 24px 0; box-shadow: 0 2px 6px rgba(0,0,0,0.08);
}
.formant-graph { max-width: 65%; border: 1px solid #eee; border-radius: 4px; display: block; }
.description {
  font-style: italic; color: #555; background: #e8f5e9;
  padding: 12px 16px; border-left: 4px solid #4caf50; margin: 12px 0; border-radius: 0 4px 4px 0;
}
.fp-note { color: #1b5e20; font-weight: bold; margin: 8px 0; }
.fp-explanation {
  background: #e8f5e9; border: 1px solid #a5d6a7; border-radius: 6px;
  padding: 14px 18px; margin: 16px 0; font-size: 0.95em;
}
.note-v4 {
  background: #fff3e0; border-left: 4px solid #ff9800;
  padding: 10px 14px; margin: 12px 0; font-size: 0.9em; color: #555;
  border-radius: 0 4px 4px 0;
}
.cluster-badge {
  display: inline-block; background: #e53935; color: white;
  padding: 2px 9px; border-radius: 10px; font-size: 0.8em;
  margin-left: 8px; font-weight: bold;
}
.yan-badge {
  display: inline-block; background: #ff9800; color: white;
  padding: 2px 8px; border-radius: 10px; font-size: 0.8em; margin-left: 8px;
}
/* Tables techniques */
.tech-table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 0.88em; }
.tech-table th, .tech-table td { border: 1px solid #ccc; padding: 5px 9px; text-align: center; }
.tech-table .header th { background: #1a3a5c; color: white; font-size: 0.9em; }
.tech-table tr:nth-child(even) { background: #fafafa; }
/* Tables références */
.ref-table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 0.88em; }
.ref-table th, .ref-table td { border: 1px solid #bbb; padding: 6px 10px; text-align: center; }
.ref-table .header th { background: #37474f; color: white; }
.ref-table tr:hover { background: #f0f4f8; }
/* Doublures */
.doublures-box {
  background: #fffde7; border: 1px solid #f9a825; border-radius: 6px;
  padding: 14px 18px; margin: 18px 0;
}
.doublures-box h4 { color: #f57f17; margin: 0 0 10px 0; }
.doublures-table { width: 100%; border-collapse: collapse; font-size: 0.88em; }
.doublures-table th, .doublures-table td { border: 1px solid #f9a825; padding: 5px 9px; text-align: center; }
.doublures-table .header th { background: #f57f17; color: white; }
/* Section intro */
.section-intro {
  padding: 16px 20px; border-radius: 6px; margin: 16px 0;
}
.section-intro.bois   { background: #e8f5e9; border-left: 5px solid #2e7d32; }
.section-intro.cuivres{ background: #fff3e0; border-left: 5px solid #e64a19; }
.section-intro.cordes { background: #e3f2fd; border-left: 5px solid #1565c0; }
.section-intro.sax    { background: #fce4ec; border-left: 5px solid #ad1457; }
.section-intro.general{ background: #f3e5f5; border-left: 5px solid #6a1b9a; }
.section-intro.intro  { background: #e8eaf6; border-left: 5px solid #283593; }
/* Source note */
.source-note {
  font-size: 0.83em; color: #888; margin-top: 40px;
  border-top: 1px solid #ddd; padding-top: 12px;
}
/* Document intro */
.doc-header {
  background: linear-gradient(135deg, #1a237e 0%, #283593 60%, #1565c0 100%);
  color: white; padding: 32px 36px; border-radius: 10px; margin-bottom: 32px;
}
.doc-header h1 { color: white; border: none; margin: 0 0 12px 0; font-size: 2em; }
.doc-header .subtitle { color: #b3c5ff; font-size: 1.05em; margin: 4px 0; }
.doc-header .meta { color: #90a4c8; font-size: 0.88em; margin-top: 16px; }
/* Tableau voyelles */
.vowel-table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 0.9em; }
.vowel-table th, .vowel-table td { border: 1px solid #ccc; padding: 7px 12px; text-align: center; }
.vowel-table .header th { background: #4a148c; color: white; }
/* Convergence / matrice */
.matrix-container { overflow-x: auto; margin: 16px 0; }
.matrix-table { border-collapse: collapse; font-size: 0.72em; white-space: nowrap; }
.matrix-table th, .matrix-table td { border: 1px solid #ccc; padding: 4px 7px; text-align: center; }
.matrix-table th { background: #263238; color: white; }
"""

# ─── Utilitaires DOCX ────────────────────────────────────────
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

def tech_table_docx(doc, inst_csv):
    techs = sorted([(t, r) for (i, t), r in DATA.items() if i == inst_csv],
                   key=lambda x: x[0])
    if not techs:
        return
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Technique', 'N', 'F1', 'F2', 'F3', 'F4', 'F5', 'F6']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], '1A3A5C')
    for tech, r in techs:
        row = table.add_row().cells
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        vals = [tech, r['n_samples']] + [fmt_hz(f) for f in fs]
        is_ord = any(s in tech.lower() for s in ('ordinario', 'non-vibrato', 'non_vibrato'))
        fill = 'DFF0D8' if is_ord else None
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
            if fill:
                set_cell_shading(row[idx], fill)
    widths_cm = [4.8, 1.1, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3]
    for row_obj in table.rows:
        for cell, width in zip(row_obj.cells, widths_cm):
            cell.width = Cm(width)

def ref_table_docx(doc, rows_data):
    if not rows_data:
        return
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Source', 'F1 (Hz)', 'F2 (Hz)', 'F3 (Hz)', 'F4 (Hz)', 'Voyelle', 'N', 'Accord']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], '37474F')
    for r in rows_data:
        row = table.add_row().cells
        is_sol = 'SOL' in r.get('source', '') or 'Yan' in r.get('source', '')
        fill = 'DFF0D8' if is_sol else None
        vals = [r.get('source','—'), r.get('f1','—'), r.get('f2','—'),
                r.get('f3','—'), r.get('f4','—'), r.get('voyelle','—'),
                r.get('n','—'), r.get('accord','—')]
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
            if fill:
                set_cell_shading(row[idx], fill)
    widths_cm = [3.0, 1.8, 1.8, 1.8, 1.8, 1.8, 1.0, 1.2]
    for row_obj in table.rows:
        for cell, width in zip(row_obj.cells, widths_cm):
            cell.width = Cm(width)

def clean_text(text):
    """
    Nettoie un texte pour le DOCX :
    - Supprime l'indentation Python des triples guillemets
    - Supprime les balises HTML simples (<strong>, <em>, <br>, etc.)
    - Retourne un texte propre sur une seule ligne cohérente
    """
    import re
    # Supprimer les balises HTML courantes
    text = re.sub(r'<strong>(.*?)</strong>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<em>(.*?)</em>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<b>(.*?)</b>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<br\s*/?>', ' ', text)
    text = re.sub(r'<[^>]+>', '', text)
    # Normaliser : strip chaque ligne puis rejoindre
    lines = [l.strip() for l in text.split('\n')]
    result = ' '.join(l for l in lines if l)
    # Nettoyer les espaces multiples
    result = re.sub(r'  +', ' ', result)
    return result.strip()

def doublures_table_docx(doc, dbl_items, header_color='F57F17'):
    """
    Rend un tableau de doublures en DOCX avec colonne Rapport de tessiture.
    dbl_items : même format que doublures_html (clé 'rapport' optionnelle).
    """
    if not dbl_items:
        return
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ['Instrument associé', 'F1', 'F1 associé', 'Écart', 'Qualité', 'Rapport tessiture']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], header_color)
    for it in dbl_items:
        row = table.add_row().cells
        delta = it.get('delta', '')
        delta_str = f"Δ={delta} Hz" if isinstance(delta, (int, float)) else str(delta)
        rapport = it.get('rapport', '—')
        vals = [it.get('instr', ''), it.get('f1_a', '—'), it.get('f1_b', '—'),
                delta_str, it.get('quality', ''), rapport]
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
    for row_obj in table.rows:
        for cell, w in zip(row_obj.cells, [3.2, 1.6, 1.6, 1.8, 2.5, 2.5]):
            cell.width = Cm(w)

def add_heading(doc, text, level=1, color=None):
    style_name = f'Heading {level}' if level <= 4 else 'Normal'
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _make_bookmark_id():
    """Génère un ID numérique unique pour chaque bookmark."""
    if not hasattr(_make_bookmark_id, '_counter'):
        _make_bookmark_id._counter = 1
    bid = _make_bookmark_id._counter
    _make_bookmark_id._counter += 1
    return bid


def add_bookmark(paragraph, bookmark_name):
    """
    Ajoute un signet (bookmark) Word sur un paragraphe existant.
    Le signet entoure le run du paragraphe.
    bookmark_name doit être unique dans le document et sans espaces.
    """
    bid = str(_make_bookmark_id())
    # bookmarkStart
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), bid)
    bm_start.set(qn('w:name'), bookmark_name)
    paragraph._p.insert(0, bm_start)
    # bookmarkEnd
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), bid)
    paragraph._p.append(bm_end)


def add_heading_bookmarked(doc, text, level=1, color=None, bookmark=None):
    """
    Ajoute un titre avec un signet Word pour permettre les hyperliens TDM.
    bookmark : nom du signet (str, sans espaces). Si None, généré depuis text.
    """
    import unicodedata, re as _re
    p = add_heading(doc, text, level=level, color=color)
    if bookmark is None:
        # Générer un nom de signet depuis le texte
        bm = unicodedata.normalize('NFD', text)
        bm = ''.join(c for c in bm if unicodedata.category(c) != 'Mn')
        bm = _re.sub(r'[^a-zA-Z0-9]+', '_', bm).strip('_')
        bookmark = bm[:40]
    add_bookmark(p, bookmark)
    return p, bookmark


def add_toc_hyperlink(paragraph, text, bookmark, size=10, bold=False,
                      color=None):
    """
    Ajoute un hyperlien interne (vers un bookmark) dans un paragraphe.
    """
    # Relation hyperlien interne
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark)

    run_elem = OxmlElement('w:r')

    # Style du lien (souligné, couleur)
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    if color:
        clr = OxmlElement('w:color')
        if isinstance(color, RGBColor):
            clr.set(qn('w:val'), f'{color.red:02X}{color.green:02X}{color.blue:02X}')
        else:
            clr.set(qn('w:val'), color)
        rPr.append(clr)

    run_elem.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)

def add_paragraph(doc, text, italic=False, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def new_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    return doc

# ─── HTML helpers ────────────────────────────────────────────
def html_head(title, extra_css=""):
    return f"""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
{CSS_GLOBAL}
{extra_css}
</style>
</head>
<body>
"""

def html_foot():
    return """</body>
</html>
"""
