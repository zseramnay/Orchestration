#!/usr/bin/env python3
# Regenerate spectral envelope overview images, grouped by family
# + Build HTML + DOCX summary pages
import os
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SAMPLE_RATE = 44100
FFT_SIZE = 4096
FREQ_RES = SAMPLE_RATE / FFT_SIZE
FORMANT_MIN_BIN = int(80 / FREQ_RES)
FORMANT_MAX_BIN = int(6000 / FREQ_RES)
MIN_PEAK_DIST_HZ = 70
PEAK_THRESHOLD_DB = 30
MAX_FORMANTS = 6
freqs = np.arange(1024) * FREQ_RES

BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
SOL = os.path.join(BASE, 'Data', 'FullSOL2020_specenv par instrument')
YAN = os.path.join(BASE, 'Data', 'Yan_Adds-Divers_specenv par instrument')
OUT_DIR = os.path.join(BASE, "Versions-html-and-docx")
OUT = os.path.join(OUT_DIR, 'media')
os.makedirs(OUT, exist_ok=True)


def find_formant_peaks_v22(envelope):
    env_db = list(envelope)
    lo = max(FORMANT_MIN_BIN, 1)
    hi = min(FORMANT_MAX_BIN, len(env_db) - 1)
    if hi <= lo:
        return []
    region = env_db[lo:hi]
    mx = max(region)
    th = mx - PEAK_THRESHOLD_DB
    peaks = []
    for i in range(1, len(region) - 1):
        v = region[i]
        if v < th:
            continue
        if v >= region[i - 1] and v >= region[i + 1]:
            peaks.append(((lo + i) * FREQ_RES, v))
    peaks.sort(key=lambda x: x[1], reverse=True)
    sel = []
    for f, a in peaks:
        if not any(abs(f - sf) < MIN_PEAK_DIST_HZ for sf, _ in sel):
            sel.append((f, a))
        if len(sel) >= MAX_FORMANTS:
            break
    sel.sort(key=lambda x: x[0])
    return sel


def compute_band_centroid(envelope, lo_hz, hi_hz):
    lo_bin = max(0, min(int(lo_hz / FREQ_RES), len(envelope) - 1))
    hi_bin = max(0, min(int(hi_hz / FREQ_RES), len(envelope) - 1))
    region = envelope[lo_bin:hi_bin + 1]
    linear = np.array([10 ** (v / 10) for v in region])
    total = linear.sum()
    if total <= 0:
        return None
    return sum(linear[i] * (lo_bin + i) * FREQ_RES for i in range(len(linear))) / total


def load_ordinario(filepath, techs=('ordinario',)):
    samples = []
    with open(filepath, encoding='utf-8') as f:
        f.readline()
        for line in f:
            parts = line.strip().split(';')
            if len(parts) < 10:
                continue
            path = parts[0]
            if not path.startswith('/'):
                continue
            pp = path.split('/')
            if len(pp) < 5 or pp[3] not in techs:
                continue
            try:
                vals = [float(v) for v in parts[1:]]
            except Exception:
                continue
            if len(vals) < 100:
                continue
            formants = find_formant_peaks_v22(vals)
            if len(formants) >= 1:
                samples.append({'formants': formants, 'envelope': np.array(vals)})
    return samples


FP_BANDS = {
    'Piccolo': (400, 1000), 'Flute': (1000, 2000), 'Bass_Flute': (1000, 2000), 'Contrabass_Flute': (800, 1600),
    'Oboe': (1000, 2000), 'English_Horn': (800, 1600), 'Clarinet_Eb': (800, 1600), 'Clarinet_Bb': (1000, 2000),
    'Bass_Clarinet': (800, 1600), 'Contrabass_Clarinet': (600, 1400), 'Bassoon': (800, 1600), 'Contrabassoon': (1000, 2000),
    'Sax_Alto': (1000, 2000), 'Horn': (600, 1400), 'Trumpet_C': (600, 1400), 'Trombone': (1000, 2000),
    'Bass_Tuba': (1000, 2000), 'Bass_Trombone': (1000, 2000), 'Contrabass_Tuba': (1000, 2000),
    'Violin': (600, 1400), 'Viola': (800, 1600), 'Violoncello': (600, 1400), 'Contrabass': (1000, 2000),
    'Violin_Ensemble': (600, 1400), 'Viola_Ensemble': (600, 1400), 'Violoncello_Ensemble': (1000, 2000), 'Contrabass_Ensemble': (800, 1600),
}

FAMILIES = {
    'Bois': [
        ('Piccolo', YAN, 'Yan_Adds-Divers_specenv.db_Piccolo.txt', ('ordinario',)),
        ('Flûte', SOL, 'FullSOL2020_specenv.db_Flute.txt', ('ordinario',)),
        ('Flûte basse', YAN, 'Yan_Adds-Divers_specenv.db_Bass-Flute.txt', ('ordinario',)),
        ('Flûte c.basse', YAN, 'Yan_Adds-Divers_specenv.db_Contrabass-Flute.txt', ('ordinario',)),
        ('Hautbois', SOL, 'FullSOL2020_specenv.db_Oboe.txt', ('ordinario',)),
        ('Cor anglais', YAN, 'Yan_Adds-Divers_specenv.db_EnglishHorn.txt', ('ordinario',)),
        ('Cl. Mib', YAN, 'Yan_Adds-Divers_specenv.db_Clarinet-Eb.txt', ('ordinario',)),
        ('Cl. Sib', SOL, 'FullSOL2020_specenv.db_Clarinet_Bb.txt', ('ordinario',)),
        ('Cl. basse', YAN, 'Yan_Adds-Divers_specenv.db_Bass-Clarinet-Bb.txt', ('ordinario',)),
        ('Cl. c.basse', YAN, 'Yan_Adds-Divers_specenv.db_Contrabass-Clarinet-Bb.txt', ('ordinario',)),
        ('Basson', SOL, 'FullSOL2020_specenv.db_Bassoon.txt', ('ordinario',)),
        ('Contrebasson', YAN, 'Yan_Adds-Divers_specenv.db_Contrebasson.txt', ('non-vibrato',)),
        ('Sax alto', SOL, 'FullSOL2020_specenv.db_Sax_Alto.txt', ('ordinario',)),
    ],
    'Cuivres': [
        ('Cor', SOL, 'FullSOL2020_specenv.db_Horn.txt', ('ordinario',)),
        ('Trompette Do', SOL, 'FullSOL2020_specenv.db_Trumpet_C.txt', ('ordinario',)),
        ('Trombone', SOL, 'FullSOL2020_specenv.db_Trombone.txt', ('ordinario',)),
        ('Trb. basse', YAN, 'Yan_Adds-Divers_specenv.db_Bass-Trombone.txt', ('ordinario',)),
        ('Tuba basse', SOL, 'FullSOL2020_specenv.db_Bass_Tuba.txt', ('ordinario',)),
        ('Tuba c.basse', YAN, 'Yan_Adds-Divers_specenv.db_Contrabass-Tuba.txt', ('ordinario',)),
    ],
    'Cordes': [
        ('Violon', SOL, 'FullSOL2020_specenv.db_Violin.txt', ('ordinario',)),
        ('Alto', SOL, 'FullSOL2020_specenv.db_Viola.txt', ('ordinario',)),
        ('Violoncelle', SOL, 'FullSOL2020_specenv.db_Violoncello.txt', ('ordinario',)),
        ('Contrebasse', SOL, 'FullSOL2020_specenv.db_Contrabass.txt', ('ordinario',)),
        ('Vln ens.', YAN, 'Yan_Adds-Divers_specenv.db_Violin-Ensemble.txt', ('ordinario',)),
        ('Alt ens.', YAN, 'Yan_Adds-Divers_specenv.db_Viola-Ensemble.txt', ('ordinario',)),
        ('Vcl ens.', YAN, 'Yan_Adds-Divers_specenv.db_Violoncello-Ensemble.txt', ('ordinario',)),
        ('Cb ens.', YAN, 'Yan_Adds-Divers_specenv.db_Contrabass-Ensemble.txt', ('non-vibrato',)),
    ],
}
FAMILY_COLORS = {'Bois': '#1565C0', 'Cuivres': '#C62828', 'Cordes': '#2E7D32'}
FP_KEY_MAP = {
    'Piccolo': 'Piccolo', 'Flûte': 'Flute', 'Flûte basse': 'Bass_Flute', 'Flûte c.basse': 'Contrabass_Flute',
    'Hautbois': 'Oboe', 'Cor anglais': 'English_Horn', 'Cl. Mib': 'Clarinet_Eb', 'Cl. Sib': 'Clarinet_Bb',
    'Cl. basse': 'Bass_Clarinet', 'Cl. c.basse': 'Contrabass_Clarinet', 'Basson': 'Bassoon', 'Contrebasson': 'Contrabassoon',
    'Sax alto': 'Sax_Alto', 'Cor': 'Horn', 'Trompette Do': 'Trumpet_C', 'Trombone': 'Trombone', 'Trb. basse': 'Bass_Trombone',
    'Tuba basse': 'Bass_Tuba', 'Tuba c.basse': 'Contrabass_Tuba', 'Violon': 'Violin', 'Alto': 'Viola', 'Violoncelle': 'Violoncello',
    'Contrebasse': 'Contrabass', 'Vln ens.': 'Violin_Ensemble', 'Alt ens.': 'Viola_Ensemble', 'Vcl ens.': 'Violoncello_Ensemble', 'Cb ens.': 'Contrabass_Ensemble',
}

summaries = []
for family_name, instruments in FAMILIES.items():
    n_inst = len(instruments)
    # ncols = 4 if n_inst > 6 else 3
    ncols = 2
    nrows = (n_inst + ncols - 1) // ncols
    fig, axes = plt.subplots(nrows, ncols, figsize=(ncols * 6.5, nrows * 4.5))
    axes_flat = axes.flatten() if n_inst > 1 else [axes]
    color = FAMILY_COLORS[family_name]
    loaded_count = 0
    sample_total = 0
    for idx, (display, directory, fname, techs) in enumerate(instruments):
        ax = axes_flat[idx]
        filepath = os.path.join(directory, fname)
        if not os.path.exists(filepath):
            ax.text(0.5, 0.5, f'{display}\nFICHIER MANQUANT', ha='center', va='center', transform=ax.transAxes)
            continue
        samples = load_ordinario(filepath, techs)
        if not samples:
            ax.text(0.5, 0.5, f'{display}\nPAS DE DONNÉES', ha='center', va='center', transform=ax.transAxes)
            continue
        loaded_count += 1
        sample_total += len(samples)
        envs = [s['envelope'] for s in samples]
        mean_env = np.mean(envs, axis=0)
        mask = freqs <= 5500
        ax.plot(freqs[mask], mean_env[mask], color=color, linewidth=2)
        if len(envs) > 3:
            std_env = np.std(envs, axis=0)
            ax.fill_between(freqs[mask], (mean_env - std_env)[mask], (mean_env + std_env)[mask], color=color, alpha=0.1)
        for fi in range(min(4, max(len(s['formants']) for s in samples))):
            fi_freqs = sorted([s['formants'][fi][0] for s in samples if fi < len(s['formants'])])
            if fi_freqs:
                fval = fi_freqs[len(fi_freqs) // 2]
                bin_idx = min(int(fval / FREQ_RES), len(mean_env) - 1)
                ax.plot(fval, mean_env[bin_idx], 'v', color='#D32F2F', markersize=6, zorder=5)
                yo = 3 if fi % 2 == 0 else -5
                ax.annotate(f'F{fi + 1}\n{fval:.0f}', (fval, mean_env[bin_idx] + yo), fontsize=6, ha='center', color='#D32F2F', fontweight='bold')
        fp_key = FP_KEY_MAP.get(display)
        fp_band = FP_BANDS.get(fp_key, (600, 1400))
        fps = [compute_band_centroid(s['envelope'], fp_band[0], fp_band[1]) for s in samples]
        fps = [f for f in fps if f]
        if fps:
            fp_med = np.median(fps)
            fp_std = np.std(fps)
            ax.axvline(fp_med, color='#2E7D32', linewidth=2, alpha=0.7)
            ax.text(fp_med, mean_env[min(int(fp_med / FREQ_RES), len(mean_env) - 1)] + 3, f'Fp={fp_med:.0f}\n(σ={fp_std:.0f})', fontsize=7, ha='center', color='#2E7D32', fontweight='bold', bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.8, edgecolor='#2E7D32'))
        f2_all = [s['formants'][1][0] for s in samples if len(s['formants']) > 1]
        f2_std = np.std(f2_all) if f2_all else 0
        ax.set_title(f'{display}\nn={len(samples)} σ(F2)={f2_std:.0f} Hz', fontsize=10, fontweight='bold', color=color)
        ax.set_xlabel('Hz', fontsize=8)
        ax.set_ylabel('dB', fontsize=8)
        ax.grid(True, alpha=0.2)
        ax.set_xlim(50, 5500)
        ax.tick_params(labelsize=7)
    for idx in range(len(instruments), len(axes_flat)):
        axes_flat[idx].set_visible(False)
    plt.suptitle(f'{family_name} — Enveloppes spectrales moyennes (ordinario)\n▼ F1–F4 (rouge) │ Fp centroïde (vert) │ ±1σ (bande)', fontsize=14, fontweight='bold', y=1.0)
    plt.tight_layout(rect=[0, 0, 1, 0.97])
    outpath = os.path.join(OUT, f'enveloppes_{family_name.lower()}.png')
    fig.savefig(outpath, dpi=150, bbox_inches='tight')
    plt.close(fig)
    summaries.append({'family': family_name, 'path': outpath, 'rel': f'media/enveloppes_{family_name.lower()}.png', 'loaded_count': loaded_count, 'instrument_count': len(instruments), 'sample_total': sample_total})
    print(f' ✓ {outpath} ({len(instruments)} instruments)')


def build_html(output_path):
    parts = []
    intros = {
        'Bois': "Les bois montrent la plus grande diversité de profils d'enveloppe, du piccolo à la flûte contrebasse.",
        'Cuivres': "Les cuivres présentent des enveloppes plus concentrées, avec des zones d'énergie très lisibles dans le medium.",
        'Cordes': "Les cordes se distinguent par des plateaux spectraux et des résonances larges ; le repère Fp y est particulièrement utile.",
    }
    slug_map = {'Bois': 'env-bois', 'Cuivres': 'env-cuivres', 'Cordes': 'env-cordes'}
    for s in summaries:
        fam = s['family']
        sid = slug_map.get(fam, fam.lower())
        parts.append(f'''<h2 id="{sid}">{fam}</h2>
<p>{intros.get(fam, '')}</p>
<p><strong>Instruments chargés :</strong> {s['loaded_count']} / {s['instrument_count']} — <strong>Échantillons :</strong> {s['sample_total']}</p>
<img src="{s['rel']}" alt="Enveloppes {fam}" style="max-width:100%;height:auto;"/>''')
    html = f'''<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<title>Référence Formantique — Enveloppes spectrales</title>
<style>
body {{ font-family: 'Segoe UI', Helvetica, Arial, sans-serif; max-width: 1100px; margin: 0 auto; padding: 20px; background: #fafafa; color: #333; }}
h1 {{ color: #1a237e; border-bottom: 3px solid #1a237e; padding-bottom: 10px; }}
h2 {{ color: #283593; margin-top: 40px; border-left: 4px solid #283593; padding-left: 12px; }}
.section-intro {{ background: #e8eaf6; padding: 15px; border-radius: 6px; margin: 15px 0; }}
img {{ max-width: 100%; height: auto; border: 1px solid #ddd; border-radius: 4px; margin: 15px 0; }}
.source-note {{ font-size: 0.85em; color: #888; margin-top: 30px; border-top: 1px solid #ddd; padding-top: 10px; }}
</style>
</head>
<body>
<h1 id="vii-enveloppes">VII. Enveloppes spectrales par famille</h1>
<div class="section-intro">
<p>Ces planches regroupent les enveloppes spectrales moyennes calculées directement à partir des données specenv brutes. Les marqueurs rouges indiquent F1–F4, la ligne verte indique Fp et la bande colorée représente ±1σ.</p>
</div>
{''.join(parts)}
<p class="source-note">Sources : exports specenv bruts SOL2020 + Yan_Adds-Divers.</p>
</body>
</html>'''
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)


def build_docx(output_path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    p = doc.add_paragraph()
    p.style = doc.styles['Title']
    r = p.add_run('VII. Enveloppes spectrales par famille')
    r.font.color.rgb = RGBColor(26, 35, 126)
    for txt in [
        'Enveloppes spectrales moyennes calculées à partir des données specenv brutes.',
        'Marqueurs rouges : F1–F4 ; ligne verte : Fp ; bande : ±1σ.',
    ]:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(txt)
    intros = {
        'Bois': "Les bois montrent la plus grande diversité de profils d'enveloppe, du piccolo à la flûte contrebasse.",
        'Cuivres': "Les cuivres présentent des enveloppes plus concentrées, avec des zones d'énergie très lisibles dans le medium.",
        'Cordes': "Les cordes se distinguent par des plateaux spectraux et des résonances larges ; le repère Fp y est particulièrement utile.",
    }
    for s in summaries:
        p = doc.add_paragraph()
        p.style = doc.styles['Heading 1']
        p.add_run(s['family'])
        doc.add_paragraph(intros.get(s['family'], ''))
        doc.add_paragraph(f"Instruments chargés : {s['loaded_count']} / {s['instrument_count']} — Échantillons : {s['sample_total']}")
        if os.path.exists(s['path']):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(s['path'], width=Inches(6.8))
    p = doc.add_paragraph()
    r = p.add_run('Sources : exports specenv bruts SOL2020 + Yan_Adds-Divers.')
    r.font.size = Pt(9)
    doc.save(output_path)


html_path = os.path.join(OUT_DIR, 'section_enveloppes_v5.html')
docx_path = os.path.join(OUT_DIR, 'section_enveloppes_v5.docx')
build_html(html_path)
build_docx(docx_path)
print(f"\n{'=' * 60}")
print(f'HTML: {html_path}')
print(f'DOCX: {docx_path}')
print(f'Images: {len(summaries)}')
