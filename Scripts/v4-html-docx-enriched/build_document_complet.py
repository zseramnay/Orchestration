#!/usr/bin/env python3
"""
build_document_complet.py — Assemble toutes les sections en un seul fichier HTML + DOCX.

HTML : sidebar de navigation fixe (même style que v33-REFERENCE_FORMANTIQUE_v3.html)
       avec 3 niveaux : sections principales (jaune), instruments (bleu clair), sous-sections (vert clair)
       Contenu dans une zone scrollable avec marge gauche de 320 px.

DOCX : table des matières Word automatique (champs TC + styles Heading),
       chaque section commence sur une nouvelle page.

Usage : lancer depuis la racine du repo Formants/
    python3 Scripts/v4-html-docx-enriched/build_document_complet.py
"""
import os
import sys
import unicodedata
import re

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from common import *

# ─── python-docx spécifique ──────────────────────────────────
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

load_all_csvs()

OUT_HTML = os.path.join(OUT_DIR, 'REFERENCE_FORMANTIQUE_COMPLETE.html')
OUT_DOCX = os.path.join(OUT_DIR, 'REFERENCE_FORMANTIQUE_COMPLETE.docx')

# ═══════════════════════════════════════════════════════════
# STRUCTURE DE NAVIGATION
# Hiérarchie : (niveau, label, anchor, is_section)
# niveau 0 = section principale (jaune, bold, barre gauche)
# niveau 1 = instrument / sous-section (bleu clair)
# niveau 2 = sous-sous-section (vert clair)
# ═══════════════════════════════════════════════════════════

TOC_STRUCTURE = [
    (0, "I. Introduction et Méthodologie",     "i-introduction"),
    (1, "Sources des données",                  "sources"),
    (1, "Méthodologie d'extraction",            "methodo"),
    (1, "Le Formant Principal (Fp)",            "fp-centroide"),
    (1, "Correspondance voyelles–fréquences",   "voyelles"),

    (0, "II. Les Bois",                         "ii-bois"),
    (1, "Flûtes",                               "bois-flutes"),
    (2, "Petite flûte",                         "bois_piccolo"),
    (2, "Flûte",                                "bois_flute"),
    (2, "Flûte basse",                          "bois_bass_flute"),
    (2, "Flûte contrebasse",                    "bois_contrabass_flute"),
    (1, "Anches doubles",                       "bois-anches"),
    (2, "Hautbois",                             "bois_hautbois"),
    (2, "Cor anglais",                          "bois_cor_anglais"),
    (2, "Basson",                               "bois_basson"),
    (2, "Contrebasson",                         "bois_contrebasson"),
    (1, "Clarinettes",                          "bois-clarinettes"),
    (2, "Clarinette en Mib",                    "bois_clarinet_eb"),
    (2, "Clarinette en Sib",                    "bois_clarinet_bb"),
    (2, "Clarinette basse",                     "bois_clarinet_basse"),
    (2, "Clarinette contrebasse",               "bois_clarinet_cb"),

    (0, "III. Les Cuivres",                     "iii-cuivres"),
    (1, "Instruments principaux",               "cuivres-principaux"),
    (2, "Cor en Fa",                            "cuivres_cor"),
    (2, "Trompette en Ut",                      "cuivres_trompette"),
    (2, "Trombone ténor",                       "cuivres_trombone"),
    (2, "Trombone basse",                       "cuivres_trb_basse"),
    (2, "Tuba basse",                           "cuivres_tuba_basse"),
    (2, "Tuba contrebasse",                     "cuivres_tuba_cb"),
    (1, "Cor avec sourdine",                    "cuivres_cor_sord"),
    (1, "Trombone avec sourdines",              "cuivres-trb-sord"),
    (2, "Trombone + sourd. cup",                "cuivres_trb_cup"),
    (2, "Trombone + sourd. sèche",              "cuivres_trb_straight"),
    (2, "Trombone + sourd. harmon",             "cuivres_trb_harmon"),
    (2, "Trombone + sourd. wah (ouvert)",       "cuivres_trb_wah_open"),
    (2, "Trombone + sourd. wah (fermé)",        "cuivres_trb_wah_closed"),
    (1, "Trompette avec sourdines",             "cuivres-tpt-sord"),
    (2, "Trompette + sourd. cup",               "cuivres_tpt_cup"),
    (2, "Trompette + sourd. sèche",             "cuivres_tpt_straight"),
    (2, "Trompette + sourd. harmon",            "cuivres_tpt_harmon"),
    (2, "Trompette + sourd. wah (ouvert)",      "cuivres_tpt_wah_open"),
    (2, "Trompette + sourd. wah (fermé)",       "cuivres_tpt_wah_closed"),
    (1, "Tuba avec sourdine",                   "cuivres_tuba_sord"),

    (0, "IV. Les Saxophones",                   "iv-saxophones"),
    (1, "Saxophone alto",                       "sax_alto"),
    (1, "Saxophones absents du corpus",         "sax-absents"),

    (0, "V. Les Cordes",                        "v-cordes"),
    (1, "Cordes solistes",                      "cordes-solistes"),
    (2, "Violon",                               "cordes_violon"),
    (2, "Violon + sourdine",                    "cordes_vln_sord"),
    (2, "Violon + sourdine piombo",             "cordes_vln_piombo"),
    (2, "Alto",                                 "cordes_alto"),
    (2, "Alto + sourdine",                      "cordes_alt_sord"),
    (2, "Alto + sourdine piombo",               "cordes_alt_piombo"),
    (2, "Violoncelle",                          "cordes_violoncelle"),
    (2, "Violoncelle + sourdine",               "cordes_vcl_sord"),
    (2, "Violoncelle + sourdine piombo",        "cordes_vcl_piombo"),
    (2, "Contrebasse",                          "cordes_contrebasse"),
    (2, "Contrebasse + sourdine",               "cordes_cb_sord"),
    (1, "Cordes d'ensemble",                    "cordes-ensembles"),
    (2, "Ensemble de violons",                  "cordes_vln_ens"),
    (2, "Ensemble de violons + sourdine",       "cordes_vln_ens_sord"),
    (2, "Ensemble d'altos",                     "cordes_alt_ens"),
    (2, "Ensemble d'altos + sourdine",          "cordes_alt_ens_sord"),
    (2, "Ensemble de violoncelles",             "cordes_vcl_ens"),
    (2, "Ensemble de violoncelles + sourdine",  "cordes_vcl_ens_sord"),
    (2, "Ensemble de contrebasses",             "cordes_cb_ens"),

    (0, "VI. Synthèse",                                        "vi-synthese"),
    (1, "Figure 1 — Positions formantiques des 27 instruments", "fig1-positions"),
    (1, "Figure 2 — Espace vocalique F1/F2",                   "fig2-vocalique"),
    (1, "Figure 3 — Enveloppes du cluster de convergence",     "fig3-cluster"),
    (1, "Cluster de convergence — Zone /o/–/å/",               "cluster"),
    (1, "Positions F1/Fp — 20 instruments de base",            "positions-f1"),
    (1, "Matrices de convergence F1",                          "matrices"),
    (1, "Doublures vérifiées",                                 "doublures-verifiees"),
    (1, "6 Principes d'orchestration acoustique",              "principes"),
    (1, "Concordance multi-sources",                           "concordance"),
]

# ═══════════════════════════════════════════════════════════
# HELPER : slugification pour les ancres id=""
# ═══════════════════════════════════════════════════════════
def slugify(text):
    text = unicodedata.normalize('NFD', text)
    text = ''.join(c for c in text if unicodedata.category(c) != 'Mn')
    text = text.lower()
    text = re.sub(r'[^a-z0-9]+', '-', text)
    return text.strip('-')

# ═══════════════════════════════════════════════════════════
# GÉNÉRATION DU SIDEBAR HTML
# ═══════════════════════════════════════════════════════════
def build_sidebar():
    # Couleurs et styles par niveau, calqués sur le v33
    level_styles = {
        0: dict(color='#f6c90e', bold=True,   border='3px solid #f6c90e', pad_left='14px'),
        1: dict(color='#a8d8ea', bold=False,  border='0',                  pad_left='28px'),
        2: dict(color='#b8e0d2', bold=False,  border='0',                  pad_left='42px'),
    }

    lines = []
    lines.append('''
<div id="toc-sidebar">
  <div class="toc-header">
    <strong>📋 Sommaire</strong>
    <button onclick="toggleToc()">✕</button>
  </div>
  <div class="toc-links">''')

    for level, label, anchor in TOC_STRUCTURE:
        s = level_styles[level]
        style = (
            f'display:block; padding:3px 14px 3px {s["pad_left"]}; '
            f'color:{s["color"]}; text-decoration:none; '
            f'border-left:{s["border"]}; line-height:1.4; '
            f'font-weight:{"bold" if s["bold"] else "normal"};'
        )
        lines.append(
            f'    <a href="#{anchor}" style="{style}" '
            f'onmouseover="this.style.background=\'#2a2a4e\'" '
            f'onmouseout="this.style.background=\'\'">'
            f'{label}</a>'
        )

    lines.append('  </div>\n</div>')
    return '\n'.join(lines)

# ═══════════════════════════════════════════════════════════
# IMPORTATION DES SCRIPTS DE SECTION
# Chaque script expose build_html(output_path) mais on veut
# juste récupérer le <body> pour l'injecter ici.
# On importe les données et fonctions directement.
# ═══════════════════════════════════════════════════════════

def extract_body(script_path):
    """
    Exécute le build_html d'un script dans un fichier temporaire
    et retourne uniquement le contenu entre <body> et </body>.
    """
    import subprocess, tempfile
    tmp = tempfile.mktemp(suffix='.html')
    # Exécuter le script en redirigeant OUT_DIR vers un tmp
    env_code = f"""
import sys, os
sys.path.insert(0, os.path.dirname('{script_path}'))
# Override OUT_DIR to tmp location
import common
common.OUT_DIR = '{os.path.dirname(tmp)}'
common.OUT_IMG = os.path.join(common.OUT_DIR, 'media')
"""
    result = subprocess.run(
        ['python3', script_path],
        capture_output=True, text=True,
        cwd=BASE
    )
    return None  # On n'utilise pas cette approche


# Alternative propre : chaque script a une fonction get_body_html()
# On va plutôt construire le HTML complet en appelant les fonctions
# directement après import. Mais les scripts ont des globals (all_info)
# qui se construisent à l'import. On va donc construire le HTML
# section par section en répliquant la logique ici,
# en s'appuyant sur les modules déjà importés.

# ─── Import des données de chaque section ────────────────────
# On importe les modules de section en les chargeant comme des modules
# Le plus propre : chaque build_xxx exporte une fonction build_html_body()
# Puisqu'ils n'en ont pas, on les exécute dans un sous-process et on lit
# le fichier temporaire généré.

import subprocess, tempfile, shutil

def get_section_body(script_name, tmp_dir):
    """
    Lance un script, lit son HTML de sortie, extrait le contenu <body>.
    Réécrit les chemins d'images relatifs pour pointer vers media/.
    """
    script_path = os.path.join(BASE, 'Scripts', 'v4-html-docx-enriched', script_name)
    out_file_map = {
        'build_intro_html_docx.py':    'section_intro_v4.html',
        'build_bois_html_docx.py':     'section_bois_v4.html',
        'build_cuivres_html_docx.py':  'section_cuivres_v4.html',
        'build_sax_html_docx.py':      'section_sax_v4.html',
        'build_cordes_html_docx.py':   'section_cordes_v4.html',
        'build_synthese_html_docx.py': 'section_synthese_v4.html',
    }
    out_filename = out_file_map[script_name]
    out_path = os.path.join(OUT_DIR, out_filename)

    print(f"  → Génération {script_name} ...", flush=True)
    result = subprocess.run(
        ['python3', script_path],
        cwd=BASE,
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print(f"  ⚠ Erreur dans {script_name}:\n{result.stderr[:500]}")
        return f'<div class="section-error">Erreur : {script_name}</div>'

    if not os.path.exists(out_path):
        print(f"  ⚠ Fichier non trouvé : {out_path}")
        return ''

    with open(out_path, 'r', encoding='utf-8') as f:
        html = f.read()

    # Extraire contenu entre <body> et </body>
    m = re.search(r'<body[^>]*>(.*)</body>', html, re.DOTALL | re.IGNORECASE)
    if not m:
        return html

    body = m.group(1).strip()

    # Supprimer les <style> résiduels qui pourraient dupliquer le CSS global
    body = re.sub(r'<style[^>]*>.*?</style>', '', body, flags=re.DOTALL)

    print(f"  ✓ {script_name} — {len(body):,} caractères")
    return body


# ═══════════════════════════════════════════════════════════
# CSS DOCUMENT COMPLET (adapté pour sidebar fixe)
# ═══════════════════════════════════════════════════════════
CSS_COMPLET = """
/* ─── Reset & layout ─── */
*, *::before, *::after { box-sizing: border-box; }
html { margin: 0; padding: 0; background: #f5f5f5; }
body {
  font-family: 'Segoe UI', Helvetica, Arial, sans-serif;
  font-size: 13px; line-height: 1.5; color: #333;
  background: #ffffff;
  margin-left: 340px;   /* laisser place au sidebar (320px + 20px) */
  margin-right: 0;
  padding: 24px 32px;
  min-height: 100vh;
  box-shadow: -3px 0 12px rgba(0,0,0,0.06);
  transition: margin-left 0.3s;
}

/* ─── Sidebar ─── */
#toc-sidebar {
  position: fixed; top: 0; left: 0;
  width: 320px; height: 100vh;
  overflow-y: auto; overflow-x: hidden;
  background: #1a1a2e; color: #e0e0e0;
  font-family: sans-serif; font-size: 12px;
  z-index: 9999;
  box-shadow: 3px 0 10px rgba(0,0,0,0.4);
  padding: 0;
  transform: translateX(0);
  transition: transform 0.3s;
}
.toc-header {
  padding: 10px 14px;
  border-bottom: 1px solid #333;
  display: flex;
  justify-content: space-between;
  align-items: center;
  position: sticky; top: 0;
  background: #1a1a2e;
  z-index: 2;
}
.toc-header strong { font-size: 13px; color: #a8d8ea; }
.toc-header button {
  background: #333; color: #ccc; border: none;
  cursor: pointer; padding: 2px 8px;
  border-radius: 3px; font-size: 11px;
}
.toc-links { padding: 6px 0; }
.toc-links a:hover { background: #2a2a4e !important; }

/* ─── Bouton toggle ─── */
#toc-toggle-btn {
  position: fixed; top: 10px; left: 10px; z-index: 10000;
  background: #1a1a2e; color: #a8d8ea;
  border: 1px solid #a8d8ea;
  padding: 5px 10px; cursor: pointer;
  border-radius: 4px; font-size: 12px;
  display: none;
}

/* ─── Typographie ─── */
h1 { color: #1a237e; border-bottom: 3px solid #2e7d32; padding-bottom: 10px;
     font-size: 1.8em; margin-top: 50px; margin-bottom: 16px; }
h2 { color: #2e7d32; margin-top: 44px; border-left: 5px solid #2e7d32;
     padding-left: 14px; font-size: 1.35em; }
h3 { color: #1b5e20; margin-top: 32px; font-size: 1.15em; }
h4 { color: #555; margin-top: 14px; font-size: 1.02em; }
p  { margin: 6px 0; }

/* ─── Cartes instrument ─── */
.instrument-card {
  background: white; border: 1px solid #ddd; border-radius: 8px;
  padding: 20px; margin: 20px 0; box-shadow: 0 2px 6px rgba(0,0,0,0.07);
  scroll-margin-top: 20px;
}
.formant-graph { max-width: 65%; border: 1px solid #eee; border-radius: 4px; display: block; }
.description {
  font-style: italic; color: #555; background: #e8f5e9;
  padding: 10px 14px; border-left: 4px solid #4caf50;
  margin: 10px 0; border-radius: 0 4px 4px 0;
}
.fp-note { color: #1b5e20; font-weight: bold; margin: 8px 0; }
.note-v4 {
  background: #fff3e0; border-left: 4px solid #ff9800;
  padding: 8px 12px; margin: 10px 0; font-size: 0.9em; color: #555;
}
.cluster-badge {
  display: inline-block; background: #e53935; color: white;
  padding: 2px 8px; border-radius: 10px; font-size: 0.78em;
  margin-left: 8px; font-weight: bold;
}
.yan-badge {
  display: inline-block; background: #ff9800; color: white;
  padding: 2px 7px; border-radius: 10px; font-size: 0.78em; margin-left: 8px;
}

/* ─── Introductions de section ─── */
.section-intro {
  padding: 14px 18px; border-radius: 6px; margin: 14px 0;
}
.section-intro.intro   { background: #e8eaf6; border-left: 5px solid #283593; }
.section-intro.bois    { background: #e8f5e9; border-left: 5px solid #2e7d32; }
.section-intro.cuivres { background: #fff3e0; border-left: 5px solid #e64a19; }
.section-intro.cordes  { background: #e3f2fd; border-left: 5px solid #1565c0; }
.section-intro.sax     { background: #fce4ec; border-left: 5px solid #ad1457; }
.section-intro.general { background: #f3e5f5; border-left: 5px solid #6a1b9a; }
.fp-explanation {
  background: #e8f5e9; border: 1px solid #a5d6a7;
  border-radius: 6px; padding: 12px 16px; margin: 14px 0; font-size: 0.93em;
}

/* ─── Tableaux ─── */
table { border-collapse: collapse; width: 100%; margin: 12px 0;
        font-size: 0.87em; word-break: break-word; }
th, td { border: 1px solid #ccc; padding: 5px 8px; text-align: center; }
.tech-table .header th,
.ref-table .header th  { background: #1a3a5c; color: white; font-size: 0.88em; }
.ref-table .header th  { background: #37474f; }
.vowel-table .header th { background: #4a148c; color: white; }
tr:nth-child(even) { background: #fafafa; }

/* ─── Doublures ─── */
.doublures-box {
  background: #fffde7; border: 1px solid #f9a825; border-radius: 6px;
  padding: 12px 16px; margin: 16px 0;
}
.doublures-box h4 { color: #f57f17; margin: 0 0 8px 0; }
.doublures-table th { background: #f57f17; color: white; }
.doublures-table td, .doublures-table th { border-color: #f9a825; padding: 4px 7px; }

/* ─── Matrices ─── */
.matrix-container { overflow-x: auto; margin: 14px 0; }

/* ─── Séparateur de section ─── */
.section-sep {
  border: none; border-top: 2px solid #e0e0e0;
  margin: 50px 0 30px 0;
}

/* ─── Source note ─── */
.source-note {
  font-size: 0.82em; color: #888; margin-top: 36px;
  border-top: 1px solid #ddd; padding-top: 10px;
}

/* ─── Doc header ─── */
.doc-header {
  background: linear-gradient(135deg, #1a237e 0%, #283593 60%, #1565c0 100%);
  color: white; padding: 30px 36px; border-radius: 10px; margin-bottom: 30px;
}
.doc-header h1 { color: white; border: none; margin: 0 0 10px 0; font-size: 1.9em; }
.doc-header .subtitle { color: #b3c5ff; font-size: 1.02em; margin: 3px 0; }
.doc-header .meta { color: #90a4c8; font-size: 0.86em; margin-top: 14px; }
"""

# ─── JavaScript sidebar ──────────────────────────────────────
JS_SIDEBAR = """
<script>
var tocVisible = true;
function toggleToc() {
  var toc = document.getElementById('toc-sidebar');
  var btn = document.getElementById('toc-toggle-btn');
  tocVisible = !tocVisible;
  if (tocVisible) {
    toc.style.transform = 'translateX(0)';
    btn.style.display = 'none';
    document.body.style.marginLeft = '340px';
  } else {
    toc.style.transform = 'translateX(-320px)';
    btn.style.display = 'block';
    document.body.style.marginLeft = '20px';
  }
}
// Highlight active section in sidebar
window.addEventListener('scroll', function() {
  var sections = document.querySelectorAll('[id]');
  var links = document.querySelectorAll('#toc-sidebar a');
  var scrollTop = window.scrollY + 80;
  var current = '';
  sections.forEach(function(s) {
    if (s.offsetTop <= scrollTop) current = s.id;
  });
  links.forEach(function(a) {
    var href = a.getAttribute('href');
    if (href && href === '#' + current) {
      a.style.background = '#2a2a4e';
    } else {
      a.style.background = '';
    }
  });
});
</script>
"""

# ═══════════════════════════════════════════════════════════
# BUILD HTML COMPLET
# ═══════════════════════════════════════════════════════════
def build_html_complet():
    print("\n=== BUILD HTML COMPLET ===")
    tmp_dir = tempfile.mkdtemp()

    sections_order = [
        'build_intro_html_docx.py',
        'build_bois_html_docx.py',
        'build_cuivres_html_docx.py',
        'build_sax_html_docx.py',
        'build_cordes_html_docx.py',
        'build_synthese_html_docx.py',
    ]

    # Générer chaque section
    bodies = {}
    for script in sections_order:
        bodies[script] = get_section_body(script, tmp_dir)

    # Assembler
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Référence Formantique de l'Orchestre — Document Complet</title>
<style>
{CSS_COMPLET}
</style>
</head>
<body>
"""
    # Sidebar
    html += build_sidebar()
    html += '\n<button id="toc-toggle-btn" onclick="toggleToc()">☰ Sommaire</button>\n'

    # Contenu
    for i, script in enumerate(sections_order):
        if i > 0:
            html += '\n<hr class="section-sep"/>\n'
        html += bodies[script]

    html += '\n' + JS_SIDEBAR
    html += '\n</body>\n</html>\n'

    with open(OUT_HTML, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f"\n✓ HTML complet : {OUT_HTML}")
    print(f"  Taille : {os.path.getsize(OUT_HTML):,} octets")

    shutil.rmtree(tmp_dir, ignore_errors=True)


# ═══════════════════════════════════════════════════════════
# BUILD DOCX COMPLET avec Table des matières
# ═══════════════════════════════════════════════════════════


# ─── Structure TDM statique ──────────────────────────────────
# (niveau, label, bookmark_name)
# bookmark_name doit correspondre exactement aux titres dans les sections DOCX
TOC_ENTRIES = [
    (1, "I. Introduction et Méthodologie",          "I__Introduction_et_Methodologie"),
    (2, "Sources des données",                       "Sources_des_donnees"),
    (2, "Méthodologie d'extraction formantique",     "Methodologie_extraction"),
    (2, "Le Formant Principal (Fp)",                 "Le_Formant_Principal_Fp"),
    (2, "Correspondance voyelles–fréquences",        "Correspondance_voyelles"),
    (1, "II. Les Bois",                              "II__Les_Bois"),
    (2, "Flûtes",                                    "Flutes"),
    (3, "Petite flûte",                              "Petite_flute"),
    (3, "Flûte",                                     "Flute"),
    (3, "Flûte basse",                               "Flute_basse"),
    (3, "Flûte contrebasse",                         "Flute_contrebasse"),
    (2, "Anches doubles",                            "Anches_doubles"),
    (3, "Hautbois",                                  "Hautbois"),
    (3, "Cor anglais",                               "Cor_anglais"),
    (3, "Basson",                                    "Basson"),
    (3, "Contrebasson",                              "Contrebasson"),
    (2, "Clarinettes",                               "Clarinettes"),
    (3, "Clarinette en Mib",                         "Clarinette_en_Mib"),
    (3, "Clarinette en Sib",                         "Clarinette_en_Sib"),
    (3, "Clarinette basse en Sib",                   "Clarinette_basse_en_Sib"),
    (3, "Clarinette contrebasse en Sib",             "Clarinette_contrebasse_en_Sib"),
    (1, "III. Les Cuivres",                          "III__Les_Cuivres"),
    (2, "Cuivres principaux",                        "Cuivres_principaux"),
    (3, "Cor en Fa",                                 "Cor_en_Fa"),
    (3, "Trompette en Ut",                           "Trompette_en_Ut"),
    (3, "Trombone ténor",                            "Trombone_tenor"),
    (3, "Trombone basse",                            "Trombone_basse"),
    (3, "Tuba basse",                                "Tuba_basse"),
    (3, "Tuba contrebasse",                          "Tuba_contrebasse"),
    (2, "Cuivres avec sourdine",                     "Cuivres_avec_sourdine"),
    (3, "Cor avec sourdine",                         "Cor_avec_sourdine"),
    (3, "Trombone avec sourdines",                   "Trombone_avec_sourdines"),
    (3, "Trompette avec sourdines",                  "Trompette_avec_sourdines"),
    (3, "Tuba avec sourdine",                        "Tuba_avec_sourdine"),
    (1, "IV. Les Saxophones",                        "IV__Les_Saxophones"),
    (2, "Saxophone alto",                            "Saxophone_alto"),
    (2, "Saxophones absents du corpus",              "Saxophones_absents"),
    (1, "V. Les Cordes",                             "V__Les_Cordes"),
    (2, "Cordes solistes",                           "Cordes_solistes"),
    (3, "Violon",                                    "Violon"),
    (3, "Alto",                                      "Alto"),
    (3, "Violoncelle",                               "Violoncelle"),
    (3, "Contrebasse",                               "Contrebasse"),
    (2, "Cordes d'ensemble",                         "Cordes_densemble"),
    (3, "Ensemble de violons",                       "Ensemble_de_violons"),
    (3, "Ensemble d'altos",                          "Ensemble_daltos"),
    (3, "Ensemble de violoncelles",                  "Ensemble_de_violoncelles"),
    (3, "Ensemble de contrebasses",                  "Ensemble_de_contrebasses"),
    (1, "VI. Synthèse — Convergences Formantiques",  "VI__Synthese"),
    (2, "Cluster de convergence 450–502 Hz",         "Cluster_de_convergence"),
    (2, "Positions F1/Fp",                           "Positions_F1_Fp"),
    (2, "Matrices de convergence",                   "Matrices_de_convergence"),
    (2, "Tableau des doublures vérifiées",           "Tableau_des_doublures"),
    (2, "6 Principes d'orchestration acoustique",    "Principes_dorchestration"),
    (2, "Concordance multi-sources",                 "Concordance_multi_sources"),
]




# ─── Niveaux de titres à inclure dans la TDM ─────────────────
TOC_HEADING_LEVELS = {'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3}

# Titres à exclure de la TDM (trop génériques, répétés pour chaque instrument)
TOC_EXCLUDE = {
    'Valeurs de référence (sources académiques)',
    'Valeurs de référence',
    'Analyse spectrale complète (toutes techniques)',
    'Analyse spectrale complète',
    'Doublures recommandées',
    'Caractéristiques principales',
    'Paramètres d\'analyse',
}


def _bookmark_name_from_text(text, seen):
    """Génère un nom de signet unique depuis un texte de titre."""
    import unicodedata as _ud, re as _re
    bm = _ud.normalize('NFD', text)
    bm = ''.join(c for c in bm if _ud.category(c) != 'Mn')
    bm = _re.sub(r'[^a-zA-Z0-9]+', '_', bm).strip('_')[:40]
    # Déduplication
    base, n = bm, 1
    while bm in seen:
        bm = f"{base}_{n}"
        n += 1
    seen.add(bm)
    return bm


def _inject_bookmarks(doc):
    """
    Parcourt tous les paragraphes Heading du document,
    injecte un bookmark Word sur chacun (sauf ceux exclus),
    et retourne la liste (level, text, bookmark_name) pour la TDM.
    """
    toc_items = []
    seen_bookmarks = set()
    bid_counter = [1000]  # démarrer à 1000 pour éviter les conflits avec bookmarks existants

    for p in doc.paragraphs:
        style = p.style.name
        level = TOC_HEADING_LEVELS.get(style)
        if level is None:
            continue
        text = p.text.strip()
        if not text or text in TOC_EXCLUDE:
            continue

        bm_name = _bookmark_name_from_text(text, seen_bookmarks)
        bid = str(bid_counter[0])
        bid_counter[0] += 1

        # Injecter bookmarkStart au début du paragraphe
        bm_start = OxmlElement('w:bookmarkStart')
        bm_start.set(qn('w:id'), bid)
        bm_start.set(qn('w:name'), bm_name)
        p._p.insert(0, bm_start)

        # Injecter bookmarkEnd à la fin
        bm_end = OxmlElement('w:bookmarkEnd')
        bm_end.set(qn('w:id'), bid)
        p._p.append(bm_end)

        toc_items.append((level, text, bm_name))

    return toc_items


def _add_toc_with_hyperlinks(doc, toc_items):
    """
    Insère la TDM cliquable avant le premier paragraphe non-vide du document.
    Chaque entrée est un hyperlien Word interne vers le bookmark du titre.
    """
    level_styles = {
        1: dict(size=12, bold=True,  indent=0.0, color='1A237E'),
        2: dict(size=10, bold=False, indent=0.5, color='1B5E20'),
        3: dict(size=9,  bold=False, indent=1.0, color='424242'),
    }

    # Trouver le premier paragraphe du body pour insérer avant lui
    body = doc.element.body
    # On va insérer les paragraphes TDM juste avant le premier paragraphe existant
    first_para = body[0]

    toc_paras = []

    # Titre "Table des matières"
    h = OxmlElement('w:p')
    hPr = OxmlElement('w:pPr')
    hStyle = OxmlElement('w:pStyle')
    hStyle.set(qn('w:val'), 'Heading1')
    hPr.append(hStyle)
    h.append(hPr)
    hR = OxmlElement('w:r')
    hT = OxmlElement('w:t')
    hT.text = 'Table des matières'
    hR.append(hT)
    h.append(hR)
    toc_paras.append(h)

    for level, text, bookmark in toc_items:
        s = level_styles[level]

        p_elem = OxmlElement('w:p')

        # Propriétés du paragraphe
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(int(s['indent'] * 720)))  # 720 twips = 1.27 cm ≈ 0.5 inch
        pPr.append(ind)
        spc = OxmlElement('w:spacing')
        spc.set(qn('w:before'), '40')
        spc.set(qn('w:after'), '20')
        pPr.append(spc)
        p_elem.append(pPr)

        # Hyperlien interne vers le bookmark
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark)

        rPr = OxmlElement('w:rPr')
        # Style hyperlien : souligné + couleur
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        clr = OxmlElement('w:color')
        clr.set(qn('w:val'), s['color'])
        rPr.append(clr)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(s['size'] * 2)))
        rPr.append(sz)
        if s['bold']:
            rPr.append(OxmlElement('w:b'))

        run_elem = OxmlElement('w:r')
        run_elem.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        run_elem.append(t)
        hyperlink.append(run_elem)
        p_elem.append(hyperlink)
        toc_paras.append(p_elem)

    # Saut de page après la TDM
    pb_p = OxmlElement('w:p')
    pb_r = OxmlElement('w:r')
    pb_br = OxmlElement('w:br')
    pb_br.set(qn('w:type'), 'page')
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    toc_paras.append(pb_p)

    # Insérer avant le premier élément du body
    for i, elem in enumerate(toc_paras):
        body.insert(i, elem)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def build_docx_complet():
    """
    Assemble les sections en un seul DOCX propre avec docxcompose.
    Stratégie TDM cliquable :
      1. docxcompose fusionne les 6 sections (les styles Heading survivent)
      2. On scanne tous les paragraphes Heading et on injecte des bookmarks
      3. On insère la TDM avec hyperliens cliquables vers ces bookmarks en tête de doc
    """
    from docxcompose.composer import Composer

    print("\n=== BUILD DOCX COMPLET ===")

    section_scripts = [
        ('build_intro_html_docx.py',    'section_intro_v4.docx'),
        ('build_bois_html_docx.py',     'section_bois_v4.docx'),
        ('build_cuivres_html_docx.py',  'section_cuivres_v4.docx'),
        ('build_sax_html_docx.py',      'section_sax_v4.docx'),
        ('build_cordes_html_docx.py',   'section_cordes_v4.docx'),
        ('build_synthese_html_docx.py', 'section_synthese_v4.docx'),
    ]

    # Régénérer toutes les sections (toujours, pour garantir la cohérence)
    section_docx_paths = []
    for script_name, out_filename in section_scripts:
        script_path = os.path.join(BASE, 'Scripts', 'v4-html-docx-enriched', script_name)
        out_path    = os.path.join(OUT_DIR, out_filename)
        print(f"  → Génération {script_name} ...", flush=True)
        result = subprocess.run(['python3', script_path], cwd=BASE,
                                capture_output=True, text=True)
        if result.returncode != 0:
            print(f"  ⚠ Erreur : {result.stderr[:400]}")
            continue
        section_docx_paths.append(out_path)
        print(f"  ✓ {out_filename}")

    # ── Document maître : page de garde seulement ────────────────
    master = new_docx()
    for sec in master.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)

    # Page de garde
    p = master.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Référence Formantique de l'Orchestre")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(26, 35, 126)

    master.add_paragraph()
    p2 = master.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Étude quantitative des formants spectraux instrumentaux")
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(70, 70, 70)

    master.add_paragraph()
    p3 = master.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(
        "5 914 échantillons · 30 instruments · Mars 2026\n"
        "SOL2020 IRCAM · Yan_Adds · Backus (1969) · Giesler (1985) · Meyer (2009)\n"
        "Pipeline v22 validé · 93 % de concordance multi-sources"
    )
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(100, 100, 100)

    add_page_break(master)

    # ── Fusionner les sections avec docxcompose ───────────────────
    composer = Composer(master)
    for path in section_docx_paths:
        print(f"  → Fusion : {os.path.basename(path)}", flush=True)
        try:
            composer.append(Document(path))
            print(f"  ✓ {os.path.basename(path)}")
        except Exception as e:
            print(f"  ⚠ Erreur fusion {os.path.basename(path)}: {e}")

    # Sauvegarder une version intermédiaire
    tmp_path = OUT_DOCX.replace('.docx', '_tmp.docx')
    composer.save(tmp_path)
    print("  ✓ Fusion terminée")

    # ── Post-processing : bookmarks + TDM cliquable ───────────────
    print("  → Injection bookmarks et TDM cliquable ...")
    merged = Document(tmp_path)

    # 1. Scanner les titres et injecter les bookmarks
    toc_items = _inject_bookmarks(merged)
    print(f"  ✓ {len(toc_items)} bookmarks injectés")

    # 2. Insérer la TDM avec hyperliens en tête du document
    _add_toc_with_hyperlinks(merged, toc_items)
    print(f"  ✓ TDM cliquable insérée ({len(toc_items)} entrées)")

    merged.save(OUT_DOCX)
    os.remove(tmp_path)

    size = os.path.getsize(OUT_DOCX)
    print(f"\n✓ DOCX complet : {OUT_DOCX}")
    print(f"  Taille : {size:,} octets")
    print(f"  TDM : {len(toc_items)} entrées cliquables")
    print("  Aucune manipulation requise à l'ouverture.")


# ═══════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════
if __name__ == '__main__':
    import sys
    mode = sys.argv[1] if len(sys.argv) > 1 else 'both'

    if mode in ('html', 'both'):
        build_html_complet()

    if mode in ('docx', 'both'):
        build_docx_complet()

    print(f"\n{'='*60}")
    if mode in ('html', 'both'):
        print(f"HTML : {OUT_HTML}")
    if mode in ('docx', 'both'):
        print(f"DOCX : {OUT_DOCX}")
    print("Usage : python3 build_document_complet.py [html|docx|both]")
