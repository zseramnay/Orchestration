#!/usr/bin/env python3
# Build Sections Saxophones + Cuivres — HTML + DOCX + graphs from CSV v22
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D
import matplotlib.ticker
import csv
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE, "Versions html and docx")
OUT_IMG = os.path.join(OUT_DIR, 'media')
os.makedirs(OUT_IMG, exist_ok=True)

def sf(v):
    try:
        return float(v)
    except Exception:
        return 0.0

DATA = {}
with open('Resultats/formants_all_techniques.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row
with open('Resultats/formants_yan_adds.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row

def get_f(inst, tech):
    r = DATA.get((inst, tech))
    if not r:
        return None
    return {'n': int(r['n_samples']), 'F': [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]}

VOWEL_ZONES = [
    (100,400,'#DCEEFB','u (oo)\nProfondeur'), (400,600,'#D5ECD5','o (oh)\nPlénitude'),
    (600,800,'#FDE8CE','å (aw)\nTransition'), (800,1250,'#F8D5D5','a (ah)\nPuissance'),
    (1250,2600,'#E8D5F0','e (eh)\nClarté'), (2600,6000,'#FFF8D0','i (ee)\nBrillance'),
]
FC = ['#D32F2F','#E64A19','#F57C00','#FFA000','#FBC02D','#CDDC39']
FA = [1.0, 0.85, 0.7, 0.55, 0.4, 0.3]

def make_graph(display, filename, n, formants, fp=None, family_color='#E65100'):
    valid = [(i, f) for i, f in enumerate(formants) if f > 0]
    if not valid:
        return None
    mf = min(max(max(f for _, f in valid) + 500, 3000), 6500)
    fig, ax = plt.subplots(figsize=(9.6, 4.8), dpi=150)
    for lo, hi, c, l in VOWEL_ZONES:
        if lo < mf:
            ax.axvspan(lo, min(hi, mf), alpha=0.35, color=c, zorder=0)
            mid = (lo + min(hi, mf)) / 2
            if mid < mf * 0.95:
                ax.text(mid, 0.97, l, ha='center', va='top', fontsize=7, color='#666', fontweight='bold', transform=ax.get_xaxis_transform())
    bw = mf * 0.012
    for i, freq in valid:
        bh = 1.0 * (1.0 - i * 0.12)
        ax.bar(freq, bh, width=bw * (1.2 - i * 0.05), color=FC[i], alpha=FA[i], edgecolor='#333', linewidth=0.8, zorder=3)
        ax.text(freq, bh + 0.03, f'F{i+1}\n{freq} Hz', ha='center', va='bottom', fontsize=8, fontweight='bold', color='#333', zorder=5)
    f1 = formants[0]
    if fp and abs(fp - f1) > 30:
        ax.plot(fp, 0.5, marker='D', markersize=14, color='#1B5E20', markeredgecolor='black', markeredgewidth=1.5, zorder=6)
        ax.annotate(f'Fp = {fp} Hz\n(centroïde)', xy=(fp, 0.5), xytext=(fp, 0.65), ha='center', fontsize=8, fontweight='bold', color='#1B5E20', arrowprops=dict(arrowstyle='->', color='#1B5E20', lw=1.5), zorder=7)
    ax.axvspan(420, 550, alpha=0.12, color='red', zorder=1, linestyle='--')
    ax.text(485, 0.02, 'cluster /o/', ha='center', va='bottom', fontsize=7, color='#C62828', fontstyle='italic', transform=ax.get_xaxis_transform())
    ax.set_xlim(100, mf)
    ax.set_ylim(0, 1.25)
    ax.set_xlabel('Fréquence (Hz)', fontsize=10, fontweight='bold')
    ax.set_ylabel('Importance relative du formant', fontsize=10, fontweight='bold')
    ax.set_title(f'{display} — Formants spectraux F1–F6 (ordinario, N={n})', fontsize=12, fontweight='bold', color=family_color, pad=12)
    ax.set_xscale('log')
    ticks = [t for t in [100,150,200,300,400,500,600,800,1000,1500,2000,3000,4000,5000,6000] if t <= mf]
    ax.set_xticks(ticks)
    ax.set_xticklabels([str(t) for t in ticks], fontsize=8)
    ax.get_xaxis().set_major_formatter(matplotlib.ticker.ScalarFormatter())
    ax.set_yticks([])
    for s in ['top', 'right', 'left']:
        ax.spines[s].set_visible(False)
    le = [mpatches.Patch(facecolor=FC[i], alpha=FA[i], edgecolor='#333', label=f'F{i+1} = {formants[i]} Hz') for i, _ in valid]
    if fp and abs(fp - f1) > 30:
        le.append(Line2D([0], [0], marker='D', color='w', markerfacecolor='#1B5E20', markeredgecolor='black', markersize=10, label=f'Fp centroïde = {fp} Hz'))
    ax.legend(handles=le, loc='upper right', fontsize=7, framealpha=0.9, edgecolor='#CCC')
    ax.text(0.01, -0.08, 'Source : CSV v22 (SOL2020 + Yan_Adds)', transform=ax.transAxes, fontsize=7, color='#888')
    plt.tight_layout()
    out = os.path.join(OUT_IMG, f'{filename}.png')
    fig.savefig(out, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return out

def fmt_hz(v):
    return '—' if v == 0 else f'{v:,}'.replace(',', ' ')

def tech_table_html(inst):
    techs = sorted([(t, DATA[(i, t)]) for (i, t) in DATA if i == inst], key=lambda x: x[0])
    if not techs:
        return ''
    rows = []
    for tech, r in techs:
        n = r['n_samples']
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        is_ord = 'ordinario' in tech.lower() and 'to_' not in tech
        bg = ' style="background-color:#dff0d8;"' if is_ord else ''
        fvals = ''.join(f'<td{bg}>{fmt_hz(f)}</td>' for f in fs)
        rows.append(f'<tr><td{bg}><b>{tech}</b></td><td{bg}>{n}</td>{fvals}</tr>')
    return '<table class="tech-table"><tr class="header"><th>Technique</th><th>N</th><th>F1</th><th>F2</th><th>F3</th><th>F4</th><th>F5</th><th>F6</th></tr>' + ''.join(rows) + '</table>'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

def tech_table_docx(doc, inst):
    techs = sorted([(t, DATA[(i, t)]) for (i, t) in DATA if i == inst], key=lambda x: x[0])
    if not techs:
        return
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Technique', 'N', 'F1', 'F2', 'F3', 'F4', 'F5', 'F6']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], '1A3A5C')
    for tech, r in techs:
        row = table.add_row().cells
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        vals = [tech, r['n_samples']] + [fmt_hz(f) for f in fs]
        is_ord = 'ordinario' in tech.lower() and 'to_' not in tech
        fill = 'DFF0D8' if is_ord else None
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
            if fill:
                set_cell_shading(row[idx], fill)
    widths_cm = [4.8, 1.1, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3]
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)

FP = {
    'Sax_Alto': 1440, 'Horn': 738, 'Trumpet_C': 1046, 'Trombone': 1218,
    'Bass_Tuba': 1239, 'Bass_Trombone': 1335, 'Contrabass_Tuba': 1182,
}
DESC = {
    ('Sax_Alto','ordinario'): 'Son chaud et expressif, grande flexibilité dynamique. F1=398 Hz dans la zone /o/ (plénitude). Le tuyau conique produit un spectre riche en harmoniques pairs et impairs. Fp=1440 Hz. Le saxophone alto est le seul saxophone de la base SOL2020.',
    ('Horn','ordinario'): 'Son rond et chaleureux, emblématique de la noblesse orchestrale. F1=388 Hz dans la zone /o/ (plénitude). Fp centroïde à 738 Hz, le plus bas des cuivres. Le cor possède le spectre le plus homogène de la famille.',
    ('Horn+sordina','ordinario'): 'Son voilé et lointain. F1 descend de 388 à 344 Hz. La sourdine déplace les formants vers le grave avec une légère nasalité. Utilisé pour les effets d'écho et de distance.',
    ('Trumpet_C','ordinario'): 'Son brillant et incisif, grande projection. F1=786 Hz dans la zone /a/ (puissance). Fp centroïde à 1046 Hz remarquablement stable malgré la très forte variabilité de F2.',
    ('Trumpet_C+sordina_straight','ordinario'): 'Son piquant et nasal, le plus utilisé en orchestre. F1=1098 Hz. Renforce la zone 1000–2000 Hz, caractère pointu et projeté.',
    ('Trumpet_C+sordina_cup','ordinario'): 'Son doux et arrondi, perd la brillance. F1=1443 Hz. Timbre proche du bugle, utile pour les passages lyriques.',
    ('Trumpet_C+sordina_wah','ordinario_open'): 'Position ouverte : son nasal et brillant. F1=560 Hz. Plus de projection que fermé, caractère expressif.',
    ('Trumpet_C+sordina_wah','ordinario_closed'): 'Position fermée : son très étouffé. F1=581 Hz. Spectre fortement filtré, caractère introverti.',
    ('Trumpet_C+sordina_harmon','ordinario'): 'Son miles-davisien. F1=2358 Hz, tout le spectre propulsé dans les aigus. Timbre très intime et concentré.',
    ('Trombone','ordinario'): 'Son plein et puissant, grande projection. F1=237 Hz (zone /u/), Fp centroïde à 1218 Hz. Le trombone couvre un large spectre, du grave profond au medium brillant.',
    ('Trombone+sordina_straight','ordinario'): 'Son nasal et métallique. La sourdine straight comprime le spectre et renforce la zone 800–1500 Hz. Caractère incisif.',
    ('Trombone+sordina_cup','ordinario'): 'Son voilé et sombre, projection réduite. La sourdine cup absorbe les harmoniques aigus. Timbre mat et feutré.',
    ('Trombone+sordina_wah','ordinario_open'): 'Position ouverte : son brillant et nasal. F1=226 Hz. Spectre riche, caractère expressif.',
    ('Trombone+sordina_wah','ordinario_closed'): 'Position fermée : son étouffé et nasal. F1 remonte à 398 Hz. Formants comprimés, timbre introverti.',
    ('Trombone+sordina_harmon','ordinario'): 'Son très concentré, quasi sinusoïdal (F1=162 Hz). Timbre jazz intime, projection directionnelle.',
    ('Bass_Tuba','ordinario'): 'Son profond et rond, fondamental de l'orchestre. F1=226 Hz (voyelle /u/). Le tuba basse ancre le registre grave. Fp=1239 Hz.',
    ('Bass_Tuba+sordina','ordinario'): 'Son assourdi et compact. F1 reste à 226 Hz mais projection réduite. Timbre plus mat et concentré.',
    ('Bass_Trombone','ordinario'): 'Son profond et puissant, plus sombre que le ténor. F1=258 Hz (zone /u/). Fp=1335 Hz. Élargit l'assise grave de la section cuivres.',
    ('Contrabass_Tuba','ordinario'): 'Son extrêmement grave et massif. F1=226 Hz identique au tuba basse. F2=463 Hz confirme sa place dans le cluster 450–502 Hz. Fp=1182 Hz.',
}
SAXOPHONES = [('Sax_Alto', 'Saxophone alto', 'sax_alto', 'ordinario')]
CUIVRES = [
    ('Horn', 'Cor en Fa', 'cuivres_horn', 'ordinario'), ('Horn+sordina', 'Cor — sourdine', 'cuivres_horn_sord', 'ordinario'),
    ('Trumpet_C', 'Trompette en Do', 'cuivres_trumpet_c', 'ordinario'), ('Trumpet_C+sordina_straight', 'Trompette — sourdine straight (sèche)', 'cuivres_tpt_sord_straight', 'ordinario'),
    ('Trumpet_C+sordina_cup', 'Trompette — sourdine cup (bol)', 'cuivres_tpt_sord_cup', 'ordinario'), ('Trumpet_C+sordina_wah', 'Trompette — sourdine wah open', 'cuivres_tpt_sord_wah_open', 'ordinario_open'),
    ('Trumpet_C+sordina_wah', 'Trompette — sourdine wah closed', 'cuivres_tpt_sord_wah_closed', 'ordinario_closed'), ('Trumpet_C+sordina_harmon', 'Trompette — sourdine harmon', 'cuivres_tpt_sord_harmon', 'ordinario'),
    ('Trombone', 'Trombone ténor', 'cuivres_trombone', 'ordinario'), ('Trombone+sordina_straight', 'Trombone — sourdine straight (sèche)', 'cuivres_trb_sord_straight', 'ordinario'),
    ('Trombone+sordina_cup', 'Trombone — sourdine cup (bol)', 'cuivres_trb_sord_cup', 'ordinario'), ('Trombone+sordina_wah', 'Trombone — sourdine wah open', 'cuivres_trb_sord_wah_open', 'ordinario_open'),
    ('Trombone+sordina_wah', 'Trombone — sourdine wah closed', 'cuivres_trb_sord_wah_closed', 'ordinario_closed'), ('Trombone+sordina_harmon', 'Trombone — sourdine harmon', 'cuivres_trb_sord_harmon', 'ordinario'),
    ('Bass_Tuba', 'Tuba basse', 'cuivres_bass_tuba', 'ordinario'), ('Bass_Tuba+sordina', 'Tuba basse — sourdine', 'cuivres_bass_tuba_sord', 'ordinario'),
    ('Bass_Trombone', 'Trombone basse', 'cuivres_bass_trombone', 'ordinario'), ('Contrabass_Tuba', 'Tuba contrebasse', 'cuivres_contrabass_tuba', 'ordinario'),
]
all_info = {}
for section_list, section_color in [(SAXOPHONES, '#E65100'), (CUIVRES, '#C62828')]:
    for csv_name, display, gfx, tech in section_list:
        d = get_f(csv_name, tech)
        if not d:
            print(f' ⚠ MANQUANT: {csv_name}/{tech}')
            continue
        fp = FP.get(csv_name)
        img = make_graph(display, gfx, d['n'], d['F'], fp, section_color)
        img_rel = os.path.relpath(img, OUT_DIR).replace(os.sep, '/') if img else None
        all_info[gfx] = {'csv': csv_name, 'display': display, 'tech': tech, 'data': d, 'fp': fp, 'img': img, 'img_rel': img_rel}
        print(f" ✓ {display:<45s} N={d['n']:>4d} F1={d['F'][0]:>5d} → {img_rel}")

def card_html(gfx, show_techs=True):
    info = all_info.get(gfx)
    if not info:
        return ''
    desc = DESC.get((info['csv'], info['tech']), '')
    is_yan = info['csv'] in ('Bass_Trombone','Contrabass_Tuba')
    badge = '<span class="yan-badge">Yan_Adds</span>' if is_yan else ''
    fp_html = f'<p class="fp-note">Fp centroïde = {info["fp"]} Hz</p>' if info['fp'] else ''
    desc_html = f'<p class="description">{desc}</p>' if desc else ''
    tt = tech_table_html(info['csv']) if show_techs and '+sordina' not in info['csv'] else ''
    return f'<div class="instrument-card"><h3>{info["display"]}{badge}</h3><img src="{info["img_rel"]}" alt="{info["display"]}" class="formant-graph"/>{desc_html}{fp_html}{tt}</div>'

def build_html(output_path):
    html = '''<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<title>Référence Formantique — Saxophones & Cuivres</title>
<style>
body { font-family: 'Segoe UI', Helvetica, Arial, sans-serif; max-width: 1100px; margin: 0 auto; padding: 20px; background: #fafafa; color: #333; }
h1 { color: #1a237e; border-bottom: 3px solid; padding-bottom: 10px; }
h1.sax { border-color: #e65100; }
h1.cuivres { border-color: #c62828; }
h2 { margin-top: 40px; border-left: 4px solid; padding-left: 12px; }
h2.sax { color: #e65100; border-color: #e65100; }
h2.cuivres { color: #c62828; border-color: #c62828; }
h3 { margin-top: 30px; }
.instrument-card { background: white; border: 1px solid #ddd; border-radius: 8px; padding: 20px; margin: 20px 0; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
.formant-graph { max-width: 100%; border: 1px solid #eee; border-radius: 4px; }
.description { font-style: italic; color: #555; background: #fff3e0; padding: 10px; border-left: 3px solid #ff9800; margin: 10px 0; }
.fp-note { color: #1b5e20; font-weight: bold; }
.tech-table { width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 0.9em; }
.tech-table th, .tech-table td { border: 1px solid #ccc; padding: 6px 10px; text-align: center; }
.tech-table .header th { background: #1a3a5c; color: white; }
.tech-table tr:nth-child(even) { background: #f8f8f8; }
.source-note { font-size: 0.85em; color: #888; margin-top: 30px; border-top: 1px solid #ddd; padding-top: 10px; }
.section-intro { padding: 15px; border-radius: 6px; margin: 15px 0; }
.section-intro.sax { background: #fff3e0; }
.section-intro.cuivres { background: #ffebee; }
.yan-badge { display: inline-block; background: #ff9800; color: white; padding: 2px 8px; border-radius: 10px; font-size: 0.8em; margin-left: 8px; }
.mute-summary { background: #fce4ec; padding: 15px; border-radius: 6px; margin: 20px 0; }
.mute-summary table { width: 100%; border-collapse: collapse; margin: 10px 0; }
.mute-summary th, .mute-summary td { border: 1px solid #ccc; padding: 6px 8px; text-align: center; font-size: 0.9em; }
.mute-summary th { background: #c62828; color: white; }
</style>
</head>
<body>
<h1 class="sax">IV. Les Saxophones</h1>
<div class="section-intro sax"><p><strong>Note :</strong> Seul le saxophone alto est présent dans la base SOL2020. Les saxophones ténor et baryton, mentionnés dans la littérature, ne disposent pas de données specenv dans le corpus actuel.</p></div>
'''
    html += card_html('sax_alto', True)
    html += '''<h1 class="cuivres">V. Les Cuivres</h1>
<div class="section-intro cuivres"><p><strong>Plage formantique :</strong> 162–2 358 Hz (voyelles /u/ → /i/).</p><p><strong>Caractéristique :</strong> formants bien définis, grande stabilité inter-dynamiques. Le <em>cluster de convergence</em> 450–502 Hz (zone /o/) rassemble Cor, Trombone, Basson, Violoncelle, Cor anglais et Saxophone ténor.</p></div>
<h2 class="cuivres">Cor</h2>'''
    html += card_html('cuivres_horn', True) + card_html('cuivres_horn_sord', False)
    html += '<h2 class="cuivres">Trompette</h2>' + card_html('cuivres_trumpet_c', True)
    for k in ['cuivres_tpt_sord_straight','cuivres_tpt_sord_cup','cuivres_tpt_sord_wah_open','cuivres_tpt_sord_wah_closed','cuivres_tpt_sord_harmon']:
        html += card_html(k, False)
    html += '<h2 class="cuivres">Trombone</h2>' + card_html('cuivres_trombone', True)
    for k in ['cuivres_trb_sord_straight','cuivres_trb_sord_cup','cuivres_trb_sord_wah_open','cuivres_trb_sord_wah_closed','cuivres_trb_sord_harmon']:
        html += card_html(k, False)
    html += '<h2 class="cuivres">Tubas</h2>'
    for k in ['cuivres_bass_tuba','cuivres_bass_tuba_sord','cuivres_bass_trombone','cuivres_contrabass_tuba']:
        html += card_html(k, k in ('cuivres_bass_tuba','cuivres_bass_trombone','cuivres_contrabass_tuba'))
    html += '<h2 class="cuivres">Effet des sourdines sur les formants</h2><div class="mute-summary"><p>Tableau comparatif : F1 ordinario sans sourdine vs avec chaque type de sourdine.</p><table><tr><th>Instrument</th><th>Sans sourdine</th><th>Straight</th><th>Cup</th><th>Wah open</th><th>Wah closed</th><th>Harmon</th></tr>'
    mute_data = [
        ('Trompette Do', all_info.get('cuivres_trumpet_c',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_tpt_sord_straight',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_tpt_sord_cup',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_tpt_sord_wah_open',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_tpt_sord_wah_closed',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_tpt_sord_harmon',{}).get('data',{}).get('F',[0])[0]),
        ('Trombone ténor', all_info.get('cuivres_trombone',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_trb_sord_straight',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_trb_sord_cup',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_trb_sord_wah_open',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_trb_sord_wah_closed',{}).get('data',{}).get('F',[0])[0], all_info.get('cuivres_trb_sord_harmon',{}).get('data',{}).get('F',[0])[0]),
    ]
    for name, sans, straight, cup, wah_o, wah_c, harmon in mute_data:
        html += f'<tr><td><b>{name}</b></td><td>{sans}</td><td>{straight}</td><td>{cup}</td><td>{wah_o}</td><td>{wah_c}</td><td>{harmon}</td></tr>'
    horn_sans = all_info.get('cuivres_horn',{}).get('data',{}).get('F',[0])[0]
    horn_sord = all_info.get('cuivres_horn_sord',{}).get('data',{}).get('F',[0])[0]
    tuba_sans = all_info.get('cuivres_bass_tuba',{}).get('data',{}).get('F',[0])[0]
    tuba_sord = all_info.get('cuivres_bass_tuba_sord',{}).get('data',{}).get('F',[0])[0]
    html += f"<tr><td><b>Cor en Fa</b></td><td>{horn_sans}</td><td colspan='5'>Sourdine unique : F1 = {horn_sord} Hz</td></tr>"
    html += f"<tr><td><b>Tuba basse</b></td><td>{tuba_sans}</td><td colspan='5'>Sourdine unique : F1 = {tuba_sord} Hz</td></tr>"
    html += '''</table><p><em>Observations : la sourdine harmon produit les déplacements les plus extrêmes. La sourdine wah fermée comprime les formants vers le medium. La sourdine cup atténue les aigus de manière homogène.</em></p></div>
<p class="source-note"><strong>Source des données :</strong> formants_all_techniques.csv — pipeline v22 validé (Δ=0 Hz).<br/><strong>Instruments Yan_Adds :</strong> Trombone basse, Tuba contrebasse.</p>
</body>
</html>'''
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)

def add_title(doc, text, color):
    p = doc.add_paragraph()
    p.style = doc.styles['Title'] if text.startswith(('IV.', 'V.')) else doc.styles['Heading 1']
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*color)

def add_instrument_docx(doc, info, show_techs=True):
    heading = doc.add_paragraph()
    heading.style = doc.styles['Heading 2']
    title = info['display']
    if info['csv'] in ('Bass_Trombone','Contrabass_Tuba'):
        title += ' [Yan_Adds]'
    run = heading.add_run(title)
    if 'cuivres' in info.get('img_rel',''):
        run.font.color.rgb = RGBColor(183, 28, 28)
    else:
        run.font.color.rgb = RGBColor(230, 81, 0)
    if info['img'] and os.path.exists(info['img']):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(info['img'], width=Inches(6.7))
    desc = DESC.get((info['csv'], info['tech']), '')
    if desc:
        p_desc = doc.add_paragraph()
        r = p_desc.add_run(desc)
        r.italic = True
        r.font.size = Pt(10)
    if info['fp']:
        p_fp = doc.add_paragraph()
        r = p_fp.add_run(f'Fp centroïde = {info["fp"]} Hz')
        r.bold = True
        r.font.color.rgb = RGBColor(27, 94, 32)
    if show_techs and '+sordina' not in info['csv']:
        tech_table_docx(doc, info['csv'])
    doc.add_paragraph()

def build_docx(output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    add_title(doc, 'IV. Les Saxophones', (26, 35, 126))
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run('Note : ')
    r1.bold = True
    p.add_run('Seul le saxophone alto est présent dans la base SOL2020. Les saxophones ténor et baryton, mentionnés dans la littérature, ne disposent pas de données specenv dans le corpus actuel.')
    for run in p.runs:
        run.font.size = Pt(10)
    if 'sax_alto' in all_info:
        add_instrument_docx(doc, all_info['sax_alto'], True)
    add_title(doc, 'V. Les Cuivres', (26, 35, 126))
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run('Plage formantique : ')
    r1.bold = True
    p.add_run('162–2 358 Hz (voyelles /u/ → /i/).')
    for run in p.runs:
        run.font.size = Pt(10)
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run('Caractéristique : ')
    r1.bold = True
    p.add_run('formants bien définis, grande stabilité inter-dynamiques. Le cluster de convergence 450–502 Hz (zone /o/) rassemble Cor, Trombone, Basson, Violoncelle, Cor anglais et Saxophone ténor.')
    for run in p.runs:
        run.font.size = Pt(10)
    add_title(doc, 'Cor', (198, 40, 40))
    for k in ['cuivres_horn','cuivres_horn_sord']:
        if k in all_info:
            add_instrument_docx(doc, all_info[k], k == 'cuivres_horn')
    add_title(doc, 'Trompette', (198, 40, 40))
    for idx, k in enumerate(['cuivres_trumpet_c','cuivres_tpt_sord_straight','cuivres_tpt_sord_cup','cuivres_tpt_sord_wah_open','cuivres_tpt_sord_wah_closed','cuivres_tpt_sord_harmon']):
        if k in all_info:
            add_instrument_docx(doc, all_info[k], idx == 0)
    add_title(doc, 'Trombone', (198, 40, 40))
    for idx, k in enumerate(['cuivres_trombone','cuivres_trb_sord_straight','cuivres_trb_sord_cup','cuivres_trb_sord_wah_open','cuivres_trb_sord_wah_closed','cuivres_trb_sord_harmon']):
        if k in all_info:
            add_instrument_docx(doc, all_info[k], idx == 0)
    add_title(doc, 'Tubas', (198, 40, 40))
    for k in ['cuivres_bass_tuba','cuivres_bass_tuba_sord','cuivres_bass_trombone','cuivres_contrabass_tuba']:
        if k in all_info:
            add_instrument_docx(doc, all_info[k], k in ('cuivres_bass_tuba','cuivres_bass_trombone','cuivres_contrabass_tuba'))
    p = doc.add_paragraph()
    notes = ['Source des données : formants_all_techniques.csv — pipeline v22 validé (Δ=0 Hz).', 'Instruments Yan_Adds : Trombone basse, Tuba contrebasse.']
    for i, line in enumerate(notes):
        run = p.add_run(line + ('\n' if i < len(notes) - 1 else ''))
        run.font.size = Pt(9)
    doc.save(output_path)

html_path = os.path.join(OUT_DIR, 'section_saxophones_cuivres.html')
docx_path = os.path.join(OUT_DIR, 'section_saxophones_cuivres.docx')
build_html(html_path)
build_docx(docx_path)
print(f"\n{'=' * 60}")
print(f'HTML: {html_path}')
print(f'DOCX: {docx_path}')
print(f'Graphiques: {len(all_info)}')
