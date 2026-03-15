#!/usr/bin/env python3
"""
Build Section Bois — HTML + DOCX + graphs from CSV v22 (single source of truth)
"""
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D
import matplotlib.ticker
import csv, os

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE, "Versions html and docx")
OUT_IMG = os.path.join(OUT_DIR, "media")
os.makedirs(OUT_IMG, exist_ok=True)


def sf(v):
    try:
        return float(v)
    except Exception:
        return 0.0


DATA = {}
with open('Resultats/formants_all_techniques.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row
with open('Resultats/formants_yan_adds.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row


def get_f(inst, tech):
    r = DATA.get((inst, tech))
    if not r:
        return None
    return {'n': int(r['n_samples']), 'F': [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]}


# ─── Fp from audit ───
FP = {
    'Piccolo': 876, 'Flute': 1535, 'Bass_Flute': 1338, 'Contrabass_Flute': 1092,
    'Oboe': 1485, 'English_Horn': 1135, 'Clarinet_Eb': 1266, 'Clarinet_Bb': 1412,
    'Bass_Clarinet_Bb': 1204, 'Contrabass_Clarinet_Bb': 1090, 'Bassoon': 1079,
    'Contrabassoon': 1279,
}


# ─── Descriptions ───
DESC = {
    'Piccolo': "Son extrêmement brillant et perçant. F1=1109 Hz, déjà dans la zone /a/ (puissance). Le piccolo est l'instrument le plus aigu de l'orchestre — ses formants se situent intégralement au-dessus de 1000 Hz. Fp centroïde à 876 Hz capture la zone de résonance de l'embouchure.",
    'Flute': "Son pur et aérien, riche en fondamentale. F1=743 Hz dans la zone /å/ (transition). Le Fp centroïde à 1535 Hz reflète la zone de résonance du tuyau. Grande variabilité de F2 par registre (σ=375 Hz).",
    'Bass_Flute': "Son chaud et soufflé, plus riche en harmoniques que la flûte standard. F1=301 Hz dans la zone /u/ (profondeur). Timbre velouté avec une composante d'air caractéristique.",
    'Contrabass_Flute': "Son très grave et breathy, à la limite du souffle perceptible. F1=334 Hz (zone /u/). Timbre profond et enveloppant, souvent utilisé pour ses qualités texturales.",
    'Oboe': "Son nasal et pénétrant, grande projection. F1=743 Hz (zone /å/), Fp=1485 Hz. L'anche double produit un spectre riche en harmoniques pairs et impairs, avec une zone formantique caractéristique autour de 1000-1200 Hz (Meyer).",
    'English_Horn': "Son plus sombre et mélancolique que le hautbois. F1=452 Hz tombe dans le cluster /o/ (450-502 Hz) — ce qui explique sa fusion naturelle avec cor, basson et violoncelle. Fp=1135 Hz.",
    'Clarinet_Eb': "Son brillant et incisif, plus perçant que la clarinette Sib. F1=678 Hz (zone /å/). La clarinette Mib possède le spectre le plus aigu de la famille des clarinettes, avec F2=1540 Hz.",
    'Clarinet_Bb': "Son chaud dans le chalumeau, brillant dans le clairon. F2 extrêmement variable (σ=795 Hz) : de 549 Hz (chalumeau) à 2638 Hz (aigu). Le Fp centroïde à 1412 Hz est 4.7× plus stable que F2. Spectre dominé par les harmoniques impairs (tuyau cylindrique).",
    'Bass_Clarinet_Bb': "Son profond et velouté, riche et soyeux dans le grave. F1=323 Hz (zone /u/). Le registre chalumeau possède une couleur très distinctive, utilisée pour ses qualités expressives et mystérieuses.",
    'Contrabass_Clarinet_Bb': "Son extrêmement grave, puissant et bourdonnant. F1=323 Hz identique à la clarinette basse, mais F2=937 Hz montre une résonance plus développée dans le medium. Timbre massif et enveloppant.",
    'Bassoon': "Son chaud et expressif, grande flexibilité timbrale. F1=495 Hz au cœur du cluster /o/ (450-502 Hz). Le basson est le pivot timbral de l'orchestre : Δ=3 Hz avec le violoncelle, Δ=45 Hz avec le cor. Fp=1079 Hz.",
    'Contrabassoon': "Son très grave et bourdonnant, fondation des bois graves. F1=226 Hz (zone /u/), identique au tuba basse. Fp=1279 Hz. Technique analysée : non-vibrato (pas d'ordinario disponible dans la base).",
}


# ─── Graph generation ───
VOWEL_ZONES = [
    (100, 400, '#DCEEFB', 'u (oo)\nProfondeur'),
    (400, 600, '#D5ECD5', 'o (oh)\nPlénitude'),
    (600, 800, '#FDE8CE', 'å (aw)\nTransition'),
    (800, 1250, '#F8D5D5', 'a (ah)\nPuissance'),
    (1250, 2600, '#E8D5F0', 'e (eh)\nClarté'),
    (2600, 6000, '#FFF8D0', 'i (ee)\nBrillance'),
]
FC = ['#D32F2F', '#E64A19', '#F57C00', '#FFA000', '#FBC02D', '#CDDC39']
FA = [1.0, 0.85, 0.7, 0.55, 0.4, 0.3]


def make_graph(display, filename, n, formants, fp=None):
    valid = [(i, f) for i, f in enumerate(formants) if f > 0]
    if not valid:
        return None

    mf = max(f for _, f in valid) + 500
    mf = min(max(mf, 3000), 6500)

    fig, ax = plt.subplots(figsize=(9.6, 4.8), dpi=150)

    for lo, hi, c, l in VOWEL_ZONES:
        if lo < mf:
            ax.axvspan(lo, min(hi, mf), alpha=0.35, color=c, zorder=0)
            mid = (lo + min(hi, mf)) / 2
            if mid < mf * 0.95:
                ax.text(mid, 0.97, l, ha='center', va='top', fontsize=7,
                        color='#666', fontweight='bold', transform=ax.get_xaxis_transform())

    bw = mf * 0.012
    for i, freq in valid:
        bh = 1.0 * (1.0 - i * 0.12)
        ax.bar(freq, bh, width=bw * (1.2 - i * 0.05), color=FC[i], alpha=FA[i],
               edgecolor='#333', linewidth=0.8, zorder=3)
        ax.text(freq, bh + 0.03, f"F{i+1}\n{freq} Hz", ha='center', va='bottom',
                fontsize=8, fontweight='bold', color='#333', zorder=5)

    f1 = formants[0]
    if fp and abs(fp - f1) > 30:
        ax.plot(fp, 0.5, marker='D', markersize=14, color='#1B5E20',
                markeredgecolor='black', markeredgewidth=1.5, zorder=6)
        ax.annotate(f"Fp = {fp} Hz\n(centroïde)", xy=(fp, 0.5), xytext=(fp, 0.65),
                    ha='center', fontsize=8, fontweight='bold', color='#1B5E20',
                    arrowprops=dict(arrowstyle='->', color='#1B5E20', lw=1.5), zorder=7)

    ax.axvspan(420, 550, alpha=0.12, color='red', zorder=1, linestyle='--')
    ax.text(485, 0.02, 'cluster /o/', ha='center', va='bottom', fontsize=7,
            color='#C62828', fontstyle='italic', transform=ax.get_xaxis_transform())
    ax.set_xlim(100, mf)
    ax.set_ylim(0, 1.25)
    ax.set_xlabel("Fréquence (Hz)", fontsize=10, fontweight='bold')
    ax.set_ylabel("Importance relative du formant", fontsize=10, fontweight='bold')
    ax.set_title(f"{display} — Formants spectraux F1–F6 (ordinario, N={n})",
                 fontsize=12, fontweight='bold', color='#2E7D32', pad=12)
    ax.set_xscale('log')
    ticks = [t for t in [100,150,200,300,400,500,600,800,1000,1500,2000,3000,4000,5000,6000] if t <= mf]
    ax.set_xticks(ticks)
    ax.set_xticklabels([str(t) for t in ticks], fontsize=8)
    ax.get_xaxis().set_major_formatter(matplotlib.ticker.ScalarFormatter())
    ax.set_yticks([])
    for s in ['top', 'right', 'left']:
        ax.spines[s].set_visible(False)

    le = [mpatches.Patch(facecolor=FC[i], alpha=FA[i], edgecolor='#333', label=f'F{i+1} = {formants[i]} Hz') for i, _ in valid]
    if fp and abs(fp - f1) > 30:
        le.append(Line2D([0], [0], marker='D', color='w', markerfacecolor='#1B5E20',
                         markeredgecolor='black', markersize=10, label=f'Fp centroïde = {fp} Hz'))
    ax.legend(handles=le, loc='upper right', fontsize=7, framealpha=0.9, edgecolor='#CCC')
    ax.text(0.01, -0.08, "Famille : Bois\nSource : CSV v22 (SOL2020 + Yan_Adds)",
            transform=ax.transAxes, fontsize=7, color='#888')
    plt.tight_layout()

    out = os.path.join(OUT_IMG, f"{filename}.png")
    fig.savefig(out, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return out


def fmt_hz(v):
    return '—' if v == 0 else f"{v:,}".replace(',', ' ')


def tech_table_html(inst_csv):
    techs = sorted([(t, DATA[(i, t)]) for (i, t) in DATA if i == inst_csv], key=lambda x: x[0])
    if not techs:
        return ""
    rows = []
    for tech, r in techs:
        n = r['n_samples']
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        is_ord = 'ordinario' in tech.lower() and 'to_' not in tech
        bg = ' style="background-color:#dff0d8;"' if is_ord else ''
        fvals = ''.join(f'<td{bg}>{fmt_hz(f)}</td>' for f in fs)
        rows.append(f'<tr><td{bg}><b>{tech}</b></td><td{bg}>{n}</td>{fvals}</tr>')
    return f"""<table class=\"tech-table\">
<tr class=\"header\"><th>Technique</th><th>N</th><th>F1</th><th>F2</th><th>F3</th><th>F4</th><th>F5</th><th>F6</th></tr>
{''.join(rows)}</table>"""


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def tech_table_docx(doc, inst_csv):
    techs = sorted([(t, DATA[(i, t)]) for (i, t) in DATA if i == inst_csv], key=lambda x: x[0])
    if not techs:
        return

    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ['Technique', 'N', 'F1', 'F2', 'F3', 'F4', 'F5', 'F6']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], '1A3A5C')

    for tech, r in techs:
        row = table.add_row().cells
        n = r['n_samples']
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        vals = [tech, n] + [fmt_hz(f) for f in fs]
        is_ord = 'ordinario' in tech.lower() and 'to_' not in tech
        fill = 'DFF0D8' if is_ord else None
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
            if fill:
                set_cell_shading(row[idx], fill)

    widths_cm = [4.8, 1.1, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3]
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)


# ═══════════════════════════════════════════════════════════
# INSTRUMENTS — dans l'ordre demandé
# ═══════════════════════════════════════════════════════════
BOIS = [
    ('Piccolo', 'Piccolo', 'bois_piccolo', 'ordinario'),
    ('Flute', 'Flûte traversière', 'bois_flute', 'ordinario'),
    ('Bass_Flute', 'Flûte basse', 'bois_bass_flute', 'ordinario'),
    ('Contrabass_Flute', 'Flûte contrebasse', 'bois_contrabass_flute', 'ordinario'),
    ('Oboe', 'Hautbois', 'bois_oboe', 'ordinario'),
    ('English_Horn', 'Cor anglais', 'bois_english_horn', 'ordinario'),
    ('Clarinet_Eb', 'Clarinette en Mib', 'bois_clarinet_eb', 'ordinario'),
    ('Clarinet_Bb', 'Clarinette en Sib', 'bois_clarinet_bb', 'ordinario'),
    ('Bass_Clarinet_Bb', 'Clarinette basse en Sib', 'bois_bass_clarinet', 'ordinario'),
    ('Contrabass_Clarinet_Bb', 'Clarinette contrebasse Sib', 'bois_contrabass_clarinet', 'ordinario'),
    ('Bassoon', 'Basson', 'bois_bassoon', 'ordinario'),
    ('Contrabassoon', 'Contrebasson', 'bois_contrabassoon', 'non-vibrato'),
]

all_info = {}
for csv_name, display, gfx, tech in BOIS:
    d = get_f(csv_name, tech)
    if not d:
        print(f" ⚠ MANQUANT: {csv_name}/{tech}")
        continue
    fp = FP.get(csv_name)
    img = make_graph(display, gfx, d['n'], d['F'], fp)
    img_rel = os.path.relpath(img, OUT_DIR).replace(os.sep, '/') if img else None
    all_info[gfx] = {
        'csv': csv_name,
        'display': display,
        'tech': tech,
        'data': d,
        'fp': fp,
        'img': img,
        'img_rel': img_rel,
    }
    print(f" ✓ {display:<35s} N={d['n']:>4d} F1={d['F'][0]:>5d} → {img_rel}")


def build_html(output_path):
    html = """<!DOCTYPE html>
<html lang=\"fr\">
<head>
<meta charset=\"UTF-8\">
<title>Référence Formantique — Section Bois</title>
<style>
 body { font-family: 'Segoe UI', Helvetica, Arial, sans-serif; max-width: 1100px; margin: 0 auto; padding: 20px; background: #fafafa; color: #333; }
 h1 { color: #1a237e; border-bottom: 3px solid #2e7d32; padding-bottom: 10px; }
 h2 { color: #2e7d32; margin-top: 40px; border-left: 4px solid #2e7d32; padding-left: 12px; }
 h3 { color: #1b5e20; margin-top: 30px; }
 .instrument-card { background: white; border: 1px solid #ddd; border-radius: 8px; padding: 20px; margin: 20px 0; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
 .formant-graph { max-width: 100%; border: 1px solid #eee; border-radius: 4px; }
 .description { font-style: italic; color: #555; background: #e8f5e9; padding: 10px; border-left: 3px solid #4caf50; margin: 10px 0; }
 .fp-note { color: #1b5e20; font-weight: bold; }
 .tech-table { width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 0.9em; }
 .tech-table th, .tech-table td { border: 1px solid #ccc; padding: 6px 10px; text-align: center; }
 .tech-table .header th { background: #1a3a5c; color: white; }
 .tech-table tr:nth-child(even) { background: #f8f8f8; }
 .source-note { font-size: 0.85em; color: #888; margin-top: 30px; border-top: 1px solid #ddd; padding-top: 10px; }
 .section-intro { background: #e8f5e9; padding: 15px; border-radius: 6px; margin: 15px 0; }
 .yan-badge { display: inline-block; background: #ff9800; color: white; padding: 2px 8px; border-radius: 10px; font-size: 0.8em; margin-left: 8px; }
</style>
</head>
<body>
<h1>III. Les Bois</h1>
<div class=\"section-intro\">
 <p><strong>Plage formantique :</strong> 226–2 638 Hz (voyelles /u/ → /e/).</p>
 <p><strong>Caractéristique :</strong> grande diversité timbrale selon le type d'anche et le profil du tuyau.
 Les instruments à anche double (hautbois, cor anglais, basson, contrebasson) partagent une zone formantique
 commune autour de 900–1200 Hz. Les clarinettes (tuyau cylindrique, harmoniques impairs) ont un comportement
 spectral très différent des flûtes (tuyau ouvert, spectre harmonique complet).</p>
 <p><strong>Découverte clé :</strong> le cor anglais (F1=452 Hz) et le basson (F1=495 Hz) tombent dans le
 <em>cluster de convergence</em> 450–502 Hz, aux côtés du cor, du trombone et du violoncelle.</p>
</div>
<h2>Flûtes</h2>
"""

    for key in ['bois_piccolo', 'bois_flute', 'bois_bass_flute', 'bois_contrabass_flute']:
        info = all_info.get(key)
        if not info:
            continue
        desc = DESC.get(info['csv'], '')
        is_yan = info['csv'] in ('Piccolo', 'Bass_Flute', 'Contrabass_Flute')
        badge = '<span class="yan-badge">Yan_Adds</span>' if is_yan else ''
        fp_html = f'<p class="fp-note">Fp centroïde = {info["fp"]} Hz</p>' if info['fp'] else ''
        html += f"""
<div class=\"instrument-card\">
<h3>{info['display']}{badge}</h3>
<img src=\"{info['img_rel']}\" alt=\"{info['display']}\" class=\"formant-graph\"/>
<p class=\"description\">{desc}</p>
{fp_html}
{tech_table_html(info['csv'])}
</div>"""

    html += "\n<h2>Anches doubles</h2>\n"
    for key in ['bois_oboe', 'bois_english_horn', 'bois_bassoon', 'bois_contrabassoon']:
        info = all_info.get(key)
        if not info:
            continue
        desc = DESC.get(info['csv'], '')
        is_yan = info['csv'] in ('English_Horn', 'Contrabassoon')
        badge = '<span class="yan-badge">Yan_Adds</span>' if is_yan else ''
        fp_html = f'<p class="fp-note">Fp centroïde = {info["fp"]} Hz</p>' if info['fp'] else ''
        tech_note = "<p><em>⚠ Contrebasson : technique analysée = non-vibrato (pas d'ordinario dans la base).</em></p>" if info['csv'] == 'Contrabassoon' else ''
        html += f"""
<div class=\"instrument-card\">
<h3>{info['display']}{badge}</h3>
<img src=\"{info['img_rel']}\" alt=\"{info['display']}\" class=\"formant-graph\"/>
<p class=\"description\">{desc}</p>
{fp_html}
{tech_note}
{tech_table_html(info['csv'])}
</div>"""

    html += "\n<h2>Clarinettes</h2>\n"
    for key in ['bois_clarinet_eb', 'bois_clarinet_bb', 'bois_bass_clarinet', 'bois_contrabass_clarinet']:
        info = all_info.get(key)
        if not info:
            continue
        desc = DESC.get(info['csv'], '')
        is_yan = info['csv'] in ('Clarinet_Eb', 'Bass_Clarinet_Bb', 'Contrabass_Clarinet_Bb')
        badge = '<span class="yan-badge">Yan_Adds</span>' if is_yan else ''
        fp_html = f'<p class="fp-note">Fp centroïde = {info["fp"]} Hz</p>' if info['fp'] else ''
        html += f"""
<div class=\"instrument-card\">
<h3>{info['display']}{badge}</h3>
<img src=\"{info['img_rel']}\" alt=\"{info['display']}\" class=\"formant-graph\"/>
<p class=\"description\">{desc}</p>
{fp_html}
{tech_table_html(info['csv'])}
</div>"""

    html += """
<p class=\"source-note\">
 <strong>Source des données :</strong> formants_all_techniques.csv — pipeline v22 validé (Δ=0 Hz pour 16/16 instruments CSV).<br/>
 <strong>Méthode :</strong> peak-picking sur enveloppes spectrales (seuil -30 dB, distance min 70 Hz, 6 formants max), agrégation par médiane.<br/>
 <strong>Fp :</strong> centroïde spectral pondéré en énergie dans une bande optimisée par instrument.<br/>
 <strong>Instruments Yan_Adds :</strong> Piccolo, Flûte basse, Flûte contrebasse, Cor anglais, Clarinette Mib, Clarinette basse, Clarinette contrebasse, Contrebasson.
</p>
</body>
</html>"""

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)


def add_intro(doc):
    p = doc.add_paragraph()
    p.style = doc.styles['Title']
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('III. Les Bois')
    run.font.color.rgb = RGBColor(26, 35, 126)

    bullets = [
        ('Plage formantique : ', '226–2 638 Hz (voyelles /u/ → /e/).'),
        ('Caractéristique : ', "grande diversité timbrale selon le type d'anche et le profil du tuyau. Les instruments à anche double (hautbois, cor anglais, basson, contrebasson) partagent une zone formantique commune autour de 900–1200 Hz. Les clarinettes (tuyau cylindrique, harmoniques impairs) ont un comportement spectral très différent des flûtes (tuyau ouvert, spectre harmonique complet)."),
        ('Découverte clé : ', 'le cor anglais (F1=452 Hz) et le basson (F1=495 Hz) tombent dans le cluster de convergence 450–502 Hz, aux côtés du cor, du trombone et du violoncelle.'),
    ]
    for label, text in bullets:
        para = doc.add_paragraph(style='List Bullet')
        r1 = para.add_run(label)
        r1.bold = True
        r2 = para.add_run(text)
        for run in para.runs:
            run.font.size = Pt(10)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(46, 125, 50)


def add_instrument(doc, info):
    heading = doc.add_paragraph()
    heading.style = doc.styles['Heading 2']
    title = info['display']
    if info['csv'] in ('Piccolo', 'Bass_Flute', 'Contrabass_Flute', 'English_Horn', 'Clarinet_Eb', 'Bass_Clarinet_Bb', 'Contrabass_Clarinet_Bb', 'Contrabassoon'):
        title += ' [Yan_Adds]'
    run = heading.add_run(title)
    run.font.color.rgb = RGBColor(27, 94, 32)

    if info['img'] and os.path.exists(info['img']):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(info['img'], width=Inches(6.7))

    desc = DESC.get(info['csv'], '')
    p_desc = doc.add_paragraph()
    r = p_desc.add_run(desc)
    r.italic = True
    r.font.size = Pt(10)

    if info['fp']:
        p_fp = doc.add_paragraph()
        r = p_fp.add_run(f"Fp centroïde = {info['fp']} Hz")
        r.bold = True
        r.font.color.rgb = RGBColor(27, 94, 32)

    if info['csv'] == 'Contrabassoon':
        p_note = doc.add_paragraph()
        r = p_note.add_run("Contrebasson : technique analysée = non-vibrato (pas d'ordinario dans la base).")
        r.italic = True

    tech_table_docx(doc, info['csv'])
    doc.add_paragraph()


def build_docx(output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = doc.styles
    styles['Normal'].font.name = 'Calibri'
    styles['Normal'].font.size = Pt(10)

    add_intro(doc)

    add_section_heading(doc, 'Flûtes')
    for key in ['bois_piccolo', 'bois_flute', 'bois_bass_flute', 'bois_contrabass_flute']:
        info = all_info.get(key)
        if info:
            add_instrument(doc, info)

    add_section_heading(doc, 'Anches doubles')
    for key in ['bois_oboe', 'bois_english_horn', 'bois_bassoon', 'bois_contrabassoon']:
        info = all_info.get(key)
        if info:
            add_instrument(doc, info)

    add_section_heading(doc, 'Clarinettes')
    for key in ['bois_clarinet_eb', 'bois_clarinet_bb', 'bois_bass_clarinet', 'bois_contrabass_clarinet']:
        info = all_info.get(key)
        if info:
            add_instrument(doc, info)

    p = doc.add_paragraph()
    notes = [
        'Source des données : formants_all_techniques.csv — pipeline v22 validé (Δ=0 Hz pour 16/16 instruments CSV).',
        'Méthode : peak-picking sur enveloppes spectrales (seuil -30 dB, distance min 70 Hz, 6 formants max), agrégation par médiane.',
        'Fp : centroïde spectral pondéré en énergie dans une bande optimisée par instrument.',
        'Instruments Yan_Adds : Piccolo, Flûte basse, Flûte contrebasse, Cor anglais, Clarinette Mib, Clarinette basse, Clarinette contrebasse, Contrebasson.',
    ]
    for i, line in enumerate(notes):
        run = p.add_run(line + ('\n' if i < len(notes) - 1 else ''))
        run.font.size = Pt(9)

    doc.save(output_path)


html_path = os.path.join(OUT_DIR, 'section_bois.html')
docx_path = os.path.join(OUT_DIR, 'section_bois.docx')

build_html(html_path)
build_docx(docx_path)

print(f"\n{'=' * 60}")
print(f"HTML: {html_path}")
print(f"DOCX: {docx_path}")
print(f"Images: {len(all_info)} graphiques")
