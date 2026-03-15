#!/usr/bin/env python3
'''
Build Section Cordes — HTML + DOCX + graphs from CSV v22 (single source of truth)
Solistes + sourdines + sourdines lourdes + ensembles
'''
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D
import matplotlib.ticker
import csv
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE, "Versions html and docx")
OUT_IMG = os.path.join(OUT_DIR, "media")
os.makedirs(OUT_IMG, exist_ok=True)


def sf(v):
    try:
        return float(v)
    except Exception:
        return 0.0


DATA = {}
with open('Resultats/formants_all_techniques.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row
with open('Resultats/formants_yan_adds.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row


def get_f(inst, tech):
    r = DATA.get((inst, tech))
    if not r:
        return None
    return {
        'n': int(r['n_samples']),
        'F': [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
    }


VOWEL_ZONES = [
    (100, 400, '#DCEEFB', 'u (oo)\nProfondeur'),
    (400, 600, '#D5ECD5', 'o (oh)\nPlénitude'),
    (600, 800, '#FDE8CE', 'å (aw)\nTransition'),
    (800, 1250, '#F8D5D5', 'a (ah)\nPuissance'),
    (1250, 2600, '#E8D5F0', 'e (eh)\nClarté'),
    (2600, 6000, '#FFF8D0', 'i (ee)\nBrillance'),
]
FC = ['#D32F2F', '#E64A19', '#F57C00', '#FFA000', '#FBC02D', '#CDDC39']
FA = [1.0, 0.85, 0.7, 0.55, 0.4, 0.3]


def make_graph(display, filename, n, formants, fp=None):
    valid = [(i, f) for i, f in enumerate(formants) if f > 0]
    if not valid:
        return None

    mf = max(f for _, f in valid) + 500
    mf = min(max(mf, 3000), 6500)

    fig, ax = plt.subplots(figsize=(9.6, 4.8), dpi=150)

    for lo, hi, c, l in VOWEL_ZONES:
        if lo < mf:
            ax.axvspan(lo, min(hi, mf), alpha=0.35, color=c, zorder=0)
            mid = (lo + min(hi, mf)) / 2
            if mid < mf * 0.95:
                ax.text(mid, 0.97, l, ha='center', va='top', fontsize=7,
                        color='#666', fontweight='bold', transform=ax.get_xaxis_transform())

    bw = mf * 0.012
    for i, freq in valid:
        bh = 1.0 * (1.0 - i * 0.12)
        ax.bar(freq, bh, width=bw * (1.2 - i * 0.05), color=FC[i], alpha=FA[i],
               edgecolor='#333', linewidth=0.8, zorder=3)
        ax.text(freq, bh + 0.03, f"F{i+1}\n{freq} Hz", ha='center', va='bottom',
                fontsize=8, fontweight='bold', color='#333', zorder=5)

    f1 = formants[0]
    if fp and abs(fp - f1) > 30:
        ax.plot(fp, 0.5, marker='D', markersize=14, color='#1B5E20',
                markeredgecolor='black', markeredgewidth=1.5, zorder=6)
        ax.annotate(f"Fp = {fp} Hz\n(centroïde)", xy=(fp, 0.5), xytext=(fp, 0.65),
                    ha='center', fontsize=8, fontweight='bold', color='#1B5E20',
                    arrowprops=dict(arrowstyle='->', color='#1B5E20', lw=1.5), zorder=7)

    ax.axvspan(420, 550, alpha=0.12, color='red', zorder=1, linestyle='--')
    ax.text(485, 0.02, 'cluster /o/', ha='center', va='bottom', fontsize=7,
            color='#C62828', fontstyle='italic', transform=ax.get_xaxis_transform())

    tech_label = 'non-vibrato' if 'Contrebasse' in display and 'Ensemble' in display else 'ordinario'
    ax.set_xlim(100, mf)
    ax.set_ylim(0, 1.25)
    ax.set_xlabel("Fréquence (Hz)", fontsize=10, fontweight='bold')
    ax.set_ylabel("Importance relative du formant", fontsize=10, fontweight='bold')
    ax.set_title(f"{display} — Formants spectraux F1–F6 ({tech_label}, N={n})",
                 fontsize=12, fontweight='bold', color='#1565C0', pad=12)
    ax.set_xscale('log')

    ticks = [t for t in [100, 150, 200, 300, 400, 500, 600, 800, 1000, 1500, 2000, 3000, 4000, 5000, 6000] if t <= mf]
    ax.set_xticks(ticks)
    ax.set_xticklabels([str(t) for t in ticks], fontsize=8)
    ax.get_xaxis().set_major_formatter(matplotlib.ticker.ScalarFormatter())
    ax.set_yticks([])

    for s in ['top', 'right', 'left']:
        ax.spines[s].set_visible(False)

    legend_elems = [
        mpatches.Patch(facecolor=FC[i], alpha=FA[i], edgecolor='#333', label=f'F{i+1} = {formants[i]} Hz')
        for i, _ in valid
    ]
    if fp and abs(fp - f1) > 30:
        legend_elems.append(
            Line2D([0], [0], marker='D', color='w', markerfacecolor='#1B5E20',
                   markeredgecolor='black', markersize=10, label=f'Fp centroïde = {fp} Hz')
        )
    ax.legend(handles=legend_elems, loc='upper right', fontsize=7, framealpha=0.9, edgecolor='#CCC')
    ax.text(0.01, -0.08, "Famille : Cordes\nSource : CSV v22 (SOL2020 + Yan_Adds)",
            transform=ax.transAxes, fontsize=7, color='#888')

    plt.tight_layout()
    out = os.path.join(OUT_IMG, f"{filename}.png")
    fig.savefig(out, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return out


def fmt_hz(v):
    return '—' if v == 0 else f"{v:,}".replace(',', ' ')


def tech_table_html(inst_csv):
    techs = sorted([(t, DATA[(i, t)]) for (i, t) in DATA if i == inst_csv], key=lambda x: x[0])
    if not techs:
        return ""

    rows = []
    for tech, r in techs:
        n = r['n_samples']
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        is_ord = 'ordinario' in tech.lower() and 'to_' not in tech
        bg = ' style="background-color:#dff0d8;"' if is_ord else ''
        fvals = ''.join(f'<td{bg}>{fmt_hz(f)}</td>' for f in fs)
        rows.append(f'<tr><td{bg}><b>{tech}</b></td><td{bg}>{n}</td>{fvals}</tr>')

    return f'''<table class="tech-table">
<tr class="header"><th>Technique</th><th>N</th><th>F1</th><th>F2</th><th>F3</th><th>F4</th><th>F5</th><th>F6</th></tr>
{''.join(rows)}</table>'''


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def tech_table_docx(doc, inst_csv):
    techs = sorted([(t, DATA[(i, t)]) for (i, t) in DATA if i == inst_csv], key=lambda x: x[0])
    if not techs:
        return

    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ['Technique', 'N', 'F1', 'F2', 'F3', 'F4', 'F5', 'F6']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], '1A3A5C')

    for tech, r in techs:
        row = table.add_row().cells
        n = r['n_samples']
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        vals = [tech, n] + [fmt_hz(f) for f in fs]
        is_ord = 'ordinario' in tech.lower() and 'to_' not in tech
        fill = 'DFF0D8' if is_ord else None
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
            if fill:
                set_cell_shading(row[idx], fill)

    widths_cm = [4.8, 1.1, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3]
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)


def add_compare_table_docx(doc):
    comparisons = [
        ('Violon', 'Violin', 'Violin_Ensemble', 'ordinario'),
        ('Alto', 'Viola', 'Viola_Ensemble', 'ordinario'),
        ('Violoncelle', 'Violoncello', 'Violoncello_Ensemble', 'ordinario'),
        ('Contrebasse', 'Contrabass', 'Contrabass_Ensemble', 'non-vibrato'),
    ]

    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.add_run('Comparaison soliste vs. ensemble (F1–F2)')
    run.font.color.rgb = RGBColor(13, 71, 161)

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ['Instrument', 'F1 solo', 'F1 ens.', 'Δ', 'F2 solo', 'F2 ens.', 'Δ']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], '1565C0')

    for name, solo_csv, ens_csv, ens_tech in comparisons:
        s = get_f(solo_csv, 'ordinario')
        e = get_f(ens_csv, ens_tech)
        if not (s and e):
            continue
        d1 = abs(s['F'][0] - e['F'][0])
        d2 = abs(s['F'][1] - e['F'][1])
        vals = [name, s['F'][0], e['F'][0], d1, s['F'][1], e['F'][1], d2]
        row = table.add_row().cells
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)

    widths_cm = [4.0, 1.7, 1.7, 1.2, 1.7, 1.7, 1.2]
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)


FP = {
    'Violin': 893, 'Violin+sordina': 853, 'Violin+sordina_piombo': None,
    'Viola': 1221, 'Viola+sordina': 1102, 'Viola+sordina_piombo': None,
    'Violoncello': 1025, 'Violoncello+sordina': 909, 'Violoncello+sordina_piombo': None,
    'Contrabass': 1324, 'Contrabass+sordina': None,
    'Violin_Ensemble': 993, 'Violin_Ensemble+sordina': None,
    'Viola_Ensemble': 1004, 'Viola_Ensemble+sordina': None,
    'Violoncello_Ensemble': 1469, 'Violoncello_Ensemble+sordina': None,
    'Contrabass_Ensemble': 1149,
}

DESC = {
    'Violin': "Son brillant et expressif, grande tessiture. F1=506 Hz dans le cluster /o/. Le Hauptformant de Meyer (800–1200 Hz) se manifeste comme un plateau spectral large : Fp=893 Hz (σ=92), remarquablement stable sur 4 octaves. F2=1518 Hz (zone nasale) est un artefact statistique du peak-picking.",
    'Violin+sordina': "La sourdine abaisse F1 de 506 à 366 Hz et atténue les aigus. Le timbre devient plus doux et voilé, avec moins de projection. Fp descend de 893 à 853 Hz.",
    'Violin+sordina_piombo': "La sourdine lourde (piombo) produit un effet plus radical : F1=344 Hz, son très étouffé et distant. Le spectre est fortement comprimé, timbre quasi irréel.",
    'Viola': "Son chaud et mélancolique, intermédiaire entre violon et violoncelle. F1=377 Hz (zone /o/). Fp=1221 Hz capture la résonance de caisse caractéristique de l'alto. Tessiture expressive dans le médium.",
    'Viola+sordina': "Sourdine : F1 descend de 377 à 344 Hz. Son plus intime et voilé, perte de projection. Fp=1102 Hz.",
    'Viola+sordina_piombo': "Sourdine lourde : F1 chute à 226 Hz. Effet très prononcé, son distant et éthéré, quasi fantomatique.",
    'Violoncello': "Son profond et lyrique, grande richesse harmonique. F1=205 Hz (zone /u/). Fp=1025 Hz (σ=70). Le violoncelle converge remarquablement avec le basson (Δ=3 Hz sur F2) et le cor (Δ=42 Hz).",
    'Violoncello+sordina': "Sourdine : F1 reste à 205 Hz mais F2 monte légèrement (538 vs 506). Son plus mat et contenu. Fp=909 Hz.",
    'Violoncello+sordina_piombo': "Sourdine lourde : F1 descend à 172 Hz. Forte atténuation des harmoniques aigus, son très assourdi et lointain.",
    'Contrabass': "Son grave et puissant, fondation des cordes. F1=172 Hz (zone /u/). Fp=1324 Hz. Faible variabilité de F2 (σ=89 Hz) — les formants de la contrebasse sont très stables.",
    'Contrabass+sordina': "Sourdine : F1 descend à 162 Hz. Son plus mat, perte de résonance. Utilisé pour les passages chambristes.",
    'Violin_Ensemble': "L'ensemble de violons comprime légèrement les formants par rapport au soliste (F1: 495 vs 506 Hz). Le spectre est plus homogène, les irrégularités individuelles s'annulent. Fp=993 Hz.",
    'Violin_Ensemble+sordina': "Ensemble avec sourdines : F1=355 Hz. L'effet de la sourdine est similaire au soliste mais encore plus homogène grâce à la fusion d'ensemble.",
    'Viola_Ensemble': "Ensemble d'altos : F1=366 Hz, proche du soliste (377 Hz). La section d'altos produit un son riche et velouté. Fp=1004 Hz.",
    'Viola_Ensemble+sordina': "Ensemble d'altos avec sourdines : F1=291 Hz. Son très voilé et intime.",
    'Violoncello_Ensemble': "Ensemble de violoncelles : F1=205 Hz, identique au soliste. La section produit un son profond et enveloppant. Fp=1469 Hz.",
    'Violoncello_Ensemble+sordina': "Ensemble de violoncelles avec sourdines : F1 reste à 205 Hz, F2 descend à 463 Hz. Son contenu et mat.",
    'Contrabass_Ensemble': "Ensemble de contrebasses (non-vibrato) : F1=172 Hz. La section produit une fondation grave massive et stable. Fp=1149 Hz.",
}

SOLISTES = [
    ('Violin', 'Violon', 'cordes_violin', 'ordinario'),
    ('Violin+sordina', 'Violon — sourdine', 'cordes_violin_sord', 'ordinario'),
    ('Violin+sordina_piombo', 'Violon — sourdine lourde', 'cordes_violin_sord_piombo', 'ordinario'),
    ('Viola', 'Alto', 'cordes_viola', 'ordinario'),
    ('Viola+sordina', 'Alto — sourdine', 'cordes_viola_sord', 'ordinario'),
    ('Viola+sordina_piombo', 'Alto — sourdine lourde', 'cordes_viola_sord_piombo', 'ordinario'),
    ('Violoncello', 'Violoncelle', 'cordes_violoncello', 'ordinario'),
    ('Violoncello+sordina', 'Violoncelle — sourdine', 'cordes_violoncello_sord', 'ordinario'),
    ('Violoncello+sordina_piombo', 'Violoncelle — sourdine lourde', 'cordes_violoncello_sord_piombo', 'ordinario'),
    ('Contrabass', 'Contrebasse', 'cordes_contrabass', 'ordinario'),
    ('Contrabass+sordina', 'Contrebasse — sourdine', 'cordes_contrabass_sord', 'ordinario'),
]
ENSEMBLES = [
    ('Violin_Ensemble', 'Ensemble de violons', 'cordes_violin_ens', 'ordinario'),
    ('Violin_Ensemble+sordina', 'Ensemble de violons — sourdine', 'cordes_violin_ens_sord', 'ordinario'),
    ('Viola_Ensemble', "Ensemble d'altos", 'cordes_viola_ens', 'ordinario'),
    ('Viola_Ensemble+sordina', "Ensemble d'altos — sourdine", 'cordes_viola_ens_sord', 'ordinario'),
    ('Violoncello_Ensemble', 'Ensemble de violoncelles', 'cordes_violoncello_ens', 'ordinario'),
    ('Violoncello_Ensemble+sordina', 'Ensemble de violoncelles — sourdine', 'cordes_violoncello_ens_sord', 'ordinario'),
    ('Contrabass_Ensemble', 'Ensemble de contrebasses', 'cordes_contrabass_ens', 'non-vibrato'),
]

GROUPS_SOLISTES = [
    ('Violon', ['cordes_violin', 'cordes_violin_sord', 'cordes_violin_sord_piombo']),
    ('Alto', ['cordes_viola', 'cordes_viola_sord', 'cordes_viola_sord_piombo']),
    ('Violoncelle', ['cordes_violoncello', 'cordes_violoncello_sord', 'cordes_violoncello_sord_piombo']),
    ('Contrebasse', ['cordes_contrabass', 'cordes_contrabass_sord']),
]
GROUPS_ENSEMBLES = [
    ('Ensembles de cordes', [
        'cordes_violin_ens', 'cordes_violin_ens_sord',
        'cordes_viola_ens', 'cordes_viola_ens_sord',
        'cordes_violoncello_ens', 'cordes_violoncello_ens_sord',
        'cordes_contrabass_ens',
    ])
]

all_info = {}
for section in [SOLISTES, ENSEMBLES]:
    for csv_name, display, gfx, tech in section:
        d = get_f(csv_name, tech)
        if not d:
            print(f" ⚠ MANQUANT: {csv_name}/{tech}")
            continue
        fp = FP.get(csv_name)
        img = make_graph(display, gfx, d['n'], d['F'], fp)
        img_rel = os.path.relpath(img, OUT_DIR).replace(os.sep, '/') if img else None
        all_info[gfx] = {
            'csv': csv_name,
            'display': display,
            'tech': tech,
            'data': d,
            'fp': fp,
            'img': img,
            'img_rel': img_rel,
        }
        print(f" ✓ {display:<45s} N={d['n']:>4d} F1={d['F'][0]:>5d} → {img_rel}")


def card_html(gfx, show_techs=True):
    info = all_info.get(gfx)
    if not info:
        return ""
    desc = DESC.get(info['csv'], '')
    is_yan = 'Ensemble' in info['csv']
    badge = '<span class="yan-badge">Yan_Adds</span>' if is_yan else ''
    fp_html = f'<p class="fp-note">Fp centroïde = {info["fp"]} Hz</p>' if info['fp'] else ''
    desc_html = f'<p class="description">{desc}</p>' if desc else ''
    tt = tech_table_html(info['csv']) if show_techs and '+sordina' not in info['csv'] else ''
    note = ''
    if info['csv'] == 'Contrabass_Ensemble':
        note = "<p><em>Technique analysée : non-vibrato (pas d'ordinario disponible dans la base).</em></p>"
    return f'''
<div class="instrument-card">
<h3>{info['display']}{badge}</h3>
<img src="{info['img_rel']}" alt="{info['display']}" class="formant-graph"/>
{desc_html}
{fp_html}
{note}
{tt}
</div>'''


def build_html(output_path):
    html = '''<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<title>Référence Formantique — Section Cordes</title>
<style>
body { font-family: 'Segoe UI', Helvetica, Arial, sans-serif; max-width: 1100px; margin: 0 auto; padding: 20px; background: #fafafa; color: #333; }
h1 { color: #1a237e; border-bottom: 3px solid #1565c0; padding-bottom: 10px; }
h2 { color: #1565c0; margin-top: 40px; border-left: 4px solid #1565c0; padding-left: 12px; }
h3 { color: #0d47a1; margin-top: 30px; }
.instrument-card { background: white; border: 1px solid #ddd; border-radius: 8px; padding: 20px; margin: 20px 0; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
.formant-graph { max-width: 100%; border: 1px solid #eee; border-radius: 4px; }
.description { font-style: italic; color: #555; background: #e3f2fd; padding: 10px; border-left: 3px solid #42a5f5; margin: 10px 0; }
.fp-note { color: #1b5e20; font-weight: bold; }
.tech-table { width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 0.9em; }
.tech-table th, .tech-table td { border: 1px solid #ccc; padding: 6px 10px; text-align: center; }
.tech-table .header th { background: #1a3a5c; color: white; }
.tech-table tr:nth-child(even) { background: #f8f8f8; }
.source-note { font-size: 0.85em; color: #888; margin-top: 30px; border-top: 1px solid #ddd; padding-top: 10px; }
.section-intro { background: #e3f2fd; padding: 15px; border-radius: 6px; margin: 15px 0; }
.yan-badge { display: inline-block; background: #ff9800; color: white; padding: 2px 8px; border-radius: 10px; font-size: 0.8em; margin-left: 8px; }
.compare-table { width: 100%; border-collapse: collapse; margin: 15px 0; }
.compare-table th, .compare-table td { border: 1px solid #ccc; padding: 8px 12px; text-align: center; }
.compare-table th { background: #1565c0; color: white; }
.compare-table tr:nth-child(even) { background: #e3f2fd; }
</style>
</head>
<body>
<h1>VI. Les Cordes</h1>
<div class="section-intro">
<p><strong>Plage formantique :</strong> 162–1 518 Hz (voyelles /u/ → /e/).</p>
<p><strong>Caractéristique :</strong> les cordes frottées possèdent des résonances de caisse larges
(plateaux spectraux) que le peak-picking ne capture pas toujours correctement. C'est pour les cordes
que le Fp centroïde apporte le gain de stabilité le plus significatif — notamment pour le violon
(Fp=893 Hz, σ=92, vs F2=1518, σ=651).</p>
<p><strong>Cluster /o/ :</strong> le violoncelle (F2=506 Hz) et l'alto (F1=377 Hz) sont proches du
cluster 450–502 Hz, facilitant leur fusion avec les cuivres (cor, trombone) et les bois (basson).</p>
</div>
<h2>Cordes solistes</h2>
'''

    for group_name, keys in GROUPS_SOLISTES:
        html += f"<h3 style='color:#1565c0; font-size:1.3em; margin-top:30px;'>{group_name}</h3>"
        for idx, k in enumerate(keys):
            html += card_html(k, show_techs=(idx == 0))

    html += '''
<h2>Ensembles de cordes</h2>
<div class="section-intro">
<p><strong>Comparaison soliste vs. ensemble :</strong> l'effet de section comprime légèrement les
formants et homogénéise le spectre. Les irrégularités individuelles s'annulent, produisant un
timbre plus lisse et plus fondu.</p>
</div>
<h3>Comparaison soliste vs. ensemble (F1-F2, ordinario / non-vibrato pour contrebasses)</h3>
<table class="compare-table">
<tr><th>Instrument</th><th>F1 solo</th><th>F1 ens.</th><th>Δ</th><th>F2 solo</th><th>F2 ens.</th><th>Δ</th></tr>
'''

    comparisons = [
        ('Violon', 'Violin', 'Violin_Ensemble', 'ordinario'),
        ('Alto', 'Viola', 'Viola_Ensemble', 'ordinario'),
        ('Violoncelle', 'Violoncello', 'Violoncello_Ensemble', 'ordinario'),
        ('Contrebasse', 'Contrabass', 'Contrabass_Ensemble', 'non-vibrato'),
    ]
    for name, solo_csv, ens_csv, ens_tech in comparisons:
        s = get_f(solo_csv, 'ordinario')
        e = get_f(ens_csv, ens_tech)
        if s and e:
            d1 = abs(s['F'][0] - e['F'][0])
            d2 = abs(s['F'][1] - e['F'][1])
            html += f"<tr><td><b>{name}</b></td><td>{s['F'][0]}</td><td>{e['F'][0]}</td><td>{d1}</td><td>{s['F'][1]}</td><td>{e['F'][1]}</td><td>{d2}</td></tr>\n"
    html += '</table>\n'

    for _, keys in GROUPS_ENSEMBLES:
        for k in keys:
            html += card_html(k, show_techs=('sord' not in k))

    html += '''
<p class="source-note">
<strong>Source des données :</strong> formants_all_techniques.csv (SOL2020) + formants_yan_adds.csv (ensembles) — pipeline v22 validé.<br/>
<strong>Ensembles :</strong> Yan_Adds (Violin-Ensemble, Viola-Ensemble, Violoncello-Ensemble, Contrabass-Ensemble).<br/>
<strong>Sourdines lourdes (piombo) :</strong> SOL2020 — effet plus prononcé que la sourdine standard.
</p>
</body>
</html>'''

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)


def add_intro(doc):
    p = doc.add_paragraph()
    p.style = doc.styles['Title']
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('VI. Les Cordes')
    run.font.color.rgb = RGBColor(26, 35, 126)

    bullets = [
        ('Plage formantique : ', '162–1 518 Hz (voyelles /u/ → /e/).'),
        ('Caractéristique : ', "les cordes frottées possèdent des résonances de caisse larges (plateaux spectraux) que le peak-picking ne capture pas toujours correctement. C'est pour les cordes que le Fp centroïde apporte le gain de stabilité le plus significatif — notamment pour le violon (Fp=893 Hz, σ=92, vs F2=1518, σ=651)."),
        ('Cluster /o/ : ', "le violoncelle (F2=506 Hz) et l'alto (F1=377 Hz) sont proches du cluster 450–502 Hz, facilitant leur fusion avec les cuivres (cor, trombone) et les bois (basson)."),
    ]
    for label, text in bullets:
        para = doc.add_paragraph(style='List Bullet')
        r1 = para.add_run(label)
        r1.bold = True
        r2 = para.add_run(text)
        for run in para.runs:
            run.font.size = Pt(10)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(21, 101, 192)


def add_subgroup_heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(13, 71, 161)


def add_instrument_docx(doc, info, show_techs=True):
    heading = doc.add_paragraph()
    heading.style = doc.styles['Heading 2']
    title = info['display']
    if 'Ensemble' in info['csv']:
        title += ' [Yan_Adds]'
    run = heading.add_run(title)
    run.font.color.rgb = RGBColor(13, 71, 161)

    if info['img'] and os.path.exists(info['img']):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(info['img'], width=Inches(6.7))

    desc = DESC.get(info['csv'], '')
    if desc:
        p_desc = doc.add_paragraph()
        r = p_desc.add_run(desc)
        r.italic = True
        r.font.size = Pt(10)

    if info['fp']:
        p_fp = doc.add_paragraph()
        r = p_fp.add_run(f"Fp centroïde = {info['fp']} Hz")
        r.bold = True
        r.font.color.rgb = RGBColor(27, 94, 32)

    if info['csv'] == 'Contrabass_Ensemble':
        p_note = doc.add_paragraph()
        r = p_note.add_run("Technique analysée : non-vibrato (pas d'ordinario disponible dans la base).")
        r.italic = True
        r.font.size = Pt(9)

    if show_techs and '+sordina' not in info['csv']:
        tech_table_docx(doc, info['csv'])

    doc.add_paragraph()


def build_docx(output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    styles = doc.styles
    styles['Normal'].font.name = 'Calibri'
    styles['Normal'].font.size = Pt(10)

    add_intro(doc)
    add_section_heading(doc, 'Cordes solistes')

    for group_name, keys in GROUPS_SOLISTES:
        add_subgroup_heading(doc, group_name)
        for idx, key in enumerate(keys):
            info = all_info.get(key)
            if info:
                add_instrument_docx(doc, info, show_techs=(idx == 0))

    add_section_heading(doc, 'Ensembles de cordes')
    p = doc.add_paragraph()
    r = p.add_run("Comparaison soliste vs. ensemble : l'effet de section comprime légèrement les formants et homogénéise le spectre. Les irrégularités individuelles s'annulent, produisant un timbre plus lisse et plus fondu.")
    r.font.size = Pt(10)
    add_compare_table_docx(doc)
    doc.add_paragraph()

    for _, keys in GROUPS_ENSEMBLES:
        for key in keys:
            info = all_info.get(key)
            if info:
                add_instrument_docx(doc, info, show_techs=('sord' not in key))

    p = doc.add_paragraph()
    notes = [
        'Source des données : formants_all_techniques.csv (SOL2020) + formants_yan_adds.csv (ensembles) — pipeline v22 validé.',
        'Ensembles : Yan_Adds (Violin-Ensemble, Viola-Ensemble, Violoncello-Ensemble, Contrabass-Ensemble).',
        'Sourdines lourdes (piombo) : SOL2020 — effet plus prononcé que la sourdine standard.',
    ]
    for i, line in enumerate(notes):
        run = p.add_run(line + ('\n' if i < len(notes) - 1 else ''))
        run.font.size = Pt(9)

    doc.save(output_path)


html_path = os.path.join(OUT_DIR, 'section_cordes.html')
docx_path = os.path.join(OUT_DIR, 'section_cordes.docx')

build_html(html_path)
build_docx(docx_path)

print(f"\n{'=' * 60}")
print(f"HTML: {html_path}")
print(f"DOCX: {docx_path}")
print(f"Images: {len(all_info)} graphiques")
