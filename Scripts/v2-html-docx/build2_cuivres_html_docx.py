#!/usr/bin/env python3
# Build Section Cuivres — HTML + DOCX + graphs from CSV v22 (single source of truth)
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D
import matplotlib.ticker
import csv
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE, "Versions html and docx")
OUT_IMG = os.path.join(OUT_DIR, 'media')
os.makedirs(OUT_IMG, exist_ok=True)

def safe_float(v):
    try:
        return float(v)
    except Exception:
        return 0.0

DATA = {}
with open('Resultats/formants_all_techniques.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row
with open('Resultats/formants_yan_adds.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row

def get_formants(inst, tech):
    r = DATA.get((inst, tech))
    if not r:
        return None
    return {
        'n': int(r['n_samples']),
        'F': [round(safe_float(r[f'F{i}_hz'])) for i in range(1, 7)],
        'reliability': safe_float(r.get('reliability_score', 0)),
    }

FP_VALUES = {
    'Bass_Tuba': 1239, 'Horn': 738, 'Trombone': 1218, 'Trumpet_C': 1046,
    'Bass_Trombone': 1335, 'Contrabass_Tuba': 1182,
}

VOWEL_ZONES = [
    (100, 400, '#DCEEFB', 'u (oo)\nProfondeur'),
    (400, 600, '#D5ECD5', 'o (oh)\nPlénitude'),
    (600, 800, '#FDE8CE', 'å (aw)\nTransition'),
    (800, 1250, '#F8D5D5', 'a (ah)\nPuissance'),
    (1250, 2600, '#E8D5F0', 'e (eh)\nClarté'),
    (2600, 6000, '#FFF8D0', 'i (ee)\nBrillance'),
]
FORMANT_COLORS = ['#D32F2F', '#E64A19', '#F57C00', '#FFA000', '#FBC02D', '#CDDC39']
FORMANT_ALPHAS = [1.0, 0.85, 0.7, 0.55, 0.4, 0.3]

def generate_bar_chart(display_name, filename, n, formants, fp=None):
    valid = [(i, f) for i, f in enumerate(formants) if f > 0]
    if not valid:
        return None
    max_freq = max(f for _, f in valid) + 500
    max_freq = min(max(max_freq, 3000), 6500)
    fig, ax = plt.subplots(figsize=(9.6, 4.8), dpi=150)
    for lo, hi, color, label in VOWEL_ZONES:
        if lo < max_freq:
            ax.axvspan(lo, min(hi, max_freq), alpha=0.35, color=color, zorder=0)
            mid = (lo + min(hi, max_freq)) / 2
            if mid < max_freq * 0.95:
                ax.text(mid, 0.97, label, ha='center', va='top', fontsize=7,
                        color='#666', fontweight='bold', transform=ax.get_xaxis_transform())
    bw = max_freq * 0.012
    for i, freq in valid:
        bar_h = 1.0 * (1.0 - i * 0.12)
        ax.bar(freq, bar_h, width=bw * (1.2 - i * 0.05),
               color=FORMANT_COLORS[i], alpha=FORMANT_ALPHAS[i], edgecolor='#333', linewidth=0.8, zorder=3)
        ax.text(freq, bar_h + 0.03, f'F{i+1}\n{freq} Hz', ha='center', va='bottom', fontsize=8,
                fontweight='bold', color='#333', zorder=5)
    f1 = formants[0]
    if fp and abs(fp - f1) > 30:
        ax.plot(fp, 0.5, marker='D', markersize=14, color='#1B5E20', markeredgecolor='black', markeredgewidth=1.5, zorder=6)
        ax.annotate(f'Fp = {fp} Hz\n(centroïde)', xy=(fp, 0.5), xytext=(fp, 0.65), ha='center', fontsize=8,
                    fontweight='bold', color='#1B5E20', arrowprops=dict(arrowstyle='->', color='#1B5E20', lw=1.5), zorder=7)
    ax.axvspan(420, 550, alpha=0.12, color='red', zorder=1, linestyle='--')
    ax.text(485, 0.02, 'cluster /o/', ha='center', va='bottom', fontsize=7, color='#C62828', fontstyle='italic', transform=ax.get_xaxis_transform())
    ax.set_xlim(100, max_freq)
    ax.set_ylim(0, 1.25)
    ax.set_xlabel('Fréquence (Hz)', fontsize=10, fontweight='bold')
    ax.set_ylabel('Importance relative du formant', fontsize=10, fontweight='bold')
    ax.set_title(f'{display_name} — Formants spectraux F1–F6 (ordinario, N={n})', fontsize=12, fontweight='bold', color='#C62828', pad=12)
    ax.set_xscale('log')
    ticks = [t for t in [100,150,200,300,400,500,600,800,1000,1500,2000,3000,4000,5000,6000] if t <= max_freq]
    ax.set_xticks(ticks)
    ax.set_xticklabels([str(t) for t in ticks], fontsize=8)
    ax.get_xaxis().set_major_formatter(matplotlib.ticker.ScalarFormatter())
    ax.set_yticks([])
    for s in ['top', 'right', 'left']:
        ax.spines[s].set_visible(False)
    legend_elements = [mpatches.Patch(facecolor=FORMANT_COLORS[i], alpha=FORMANT_ALPHAS[i], edgecolor='#333', label=f'F{i+1} = {formants[i]} Hz') for i, _ in valid]
    if fp and abs(fp - f1) > 30:
        legend_elements.append(Line2D([0], [0], marker='D', color='w', markerfacecolor='#1B5E20', markeredgecolor='black', markersize=10, label=f'Fp centroïde = {fp} Hz'))
    ax.legend(handles=legend_elements, loc='upper right', fontsize=7, framealpha=0.9, edgecolor='#CCC')
    ax.text(0.01, -0.08, 'Famille : Cuivres\nSource : formants_all_techniques.csv (SOL2020 v22)', transform=ax.transAxes, fontsize=7, color='#888')
    plt.tight_layout()
    outpath = os.path.join(OUT_IMG, f'{filename}.png')
    fig.savefig(outpath, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return outpath

DESCRIPTIONS = {
    ('Bass_Tuba', 'ordinario'): "Son profond et rond, fondamental de l'orchestre. Formants très graves (F1=226 Hz, voyelle /u/), le tuba basse ancre le registre grave avec une couleur sombre et enveloppante.",
    ('Horn', 'ordinario'): "Son rond et chaleureux, emblématique de la noblesse orchestrale. F1=388 Hz dans la zone /o/ (plénitude). Le cor possède le spectre le plus homogène des cuivres, avec un Fp centroïde à 738 Hz.",
    ('Trombone', 'ordinario'): "Son plein et puissant, grande projection. F1=237 Hz (zone /u/) avec un Fp centroïde à 1218 Hz. Le trombone couvre un large spectre, du grave profond au medium brillant.",
    ('Trumpet_C', 'ordinario'): "Son brillant et incisif, grande projection. F1=786 Hz dans la zone /a/ (puissance). Le Fp centroïde à 1046 Hz est remarquablement stable (σ=98 Hz) malgré la très forte variabilité de F2 (σ=1018 Hz).",
    ('Trombone+sordina_cup', 'ordinario'): 'Son voilé et sombre, projection réduite. La sourdine cup absorbe les harmoniques aigus, produisant un timbre mat et feutré.',
    ('Trombone+sordina_straight', 'ordinario'): 'Son nasal et métallique. La sourdine straight comprime le spectre et renforce la zone 800–1500 Hz, donnant un caractère incisif.',
    ('Trombone+sordina_harmon', 'ordinario'): 'Son très concentré, quasi sinusoïdal dans le grave (F1=162 Hz). Timbre jazz intime, projection très directionnelle.',
    ('Trombone+sordina_wah', 'ordinario_open'): 'Position ouverte : son brillant et nasal, spectre riche. F1=226 Hz, caractère expressif.',
    ('Trombone+sordina_wah', 'ordinario_closed'): 'Position fermée : son très étouffé et nasal. F1 remonte à 398 Hz, les formants sont comprimés.',
    ('Trumpet_C+sordina_cup', 'ordinario'): 'Son doux et arrondi, perd la brillance caractéristique. F1=1443 Hz, timbre proche du bugle.',
    ('Trumpet_C+sordina_straight', 'ordinario'): 'Son piquant et nasal, le plus utilisé en orchestre. F1=1098 Hz, caractère pointu et projeté.',
    ('Trumpet_C+sordina_harmon', 'ordinario'): 'Son miles-davisien. F1=2358 Hz, tout le spectre propulsé dans les aigus. Timbre très intime.',
    ('Trumpet_C+sordina_wah', 'ordinario_open'): 'Tige insérée, position ouverte : son nasal et brillant. F1=560 Hz, plus de projection que fermé.',
    ('Trumpet_C+sordina_wah', 'ordinario_closed'): 'Tige insérée, position fermée : son très étouffé. F1=581 Hz, spectre fortement filtré.',
    ('Horn+sordina', 'ordinario'): 'Son voilé et lointain. F1 descend de 388 à 344 Hz, la sourdine déplace les formants vers le grave avec une légère nasalité.',
    ('Bass_Tuba+sordina', 'ordinario'): 'Son assourdi et compact. F1 reste à 226 Hz mais la projection est réduite. Timbre plus mat.',
    ('Bass_Trombone', 'ordinario'): "Son profond et puissant, plus sombre que le ténor. F1=258 Hz (zone /u/), Fp=1335 Hz. Le trombone basse élargit l'assise grave de la section cuivres.",
    ('Contrabass_Tuba', 'ordinario'): 'Son extrêmement grave et massif. F1=226 Hz identique au tuba basse, mais F2=463 Hz confirme sa place dans le cluster 450-502 Hz.',
}

INSTRUMENTS = [
    ('Bass_Tuba', 'Tuba basse', 'cuivres_bass_tuba_ord', 1239),
    ('Horn', 'Cor en Fa', 'cuivres_horn_ord', 738),
    ('Trombone', 'Trombone ténor', 'cuivres_trombone_ord', 1218),
    ('Trumpet_C', 'Trompette en Do', 'cuivres_trumpet_c_ord', 1046),
    ('Bass_Trombone', 'Trombone basse', 'cuivres_bass_trombone_ord', 1335),
    ('Contrabass_Tuba', 'Tuba contrebasse', 'cuivres_contrabass_tuba_ord', 1182),
    ('Trombone+sordina_cup', 'Trombone + sourd. cup', 'cuivres_trb_sord_cup', None),
    ('Trombone+sordina_straight', 'Trombone + sourd. straight', 'cuivres_trb_sord_straight', None),
    ('Trombone+sordina_harmon', 'Trombone + sourd. harmon', 'cuivres_trb_sord_harmon', None),
    ('Trombone+sordina_wah', 'Trombone + sourd. wah (ouvert)', 'cuivres_trb_sord_wah_open', None),
    ('Trombone+sordina_wah', 'Trombone + sourd. wah (fermé)', 'cuivres_trb_sord_wah_closed', None),
    ('Trumpet_C+sordina_cup', 'Trompette + sourd. cup', 'cuivres_tpt_sord_cup', None),
    ('Trumpet_C+sordina_straight', 'Trompette + sourd. straight', 'cuivres_tpt_sord_straight', None),
    ('Trumpet_C+sordina_harmon', 'Trompette + sourd. harmon', 'cuivres_tpt_sord_harmon', None),
    ('Trumpet_C+sordina_wah', 'Trompette + sourd. wah (ouvert)', 'cuivres_tpt_sord_wah_open', None),
    ('Trumpet_C+sordina_wah', 'Trompette + sourd. wah (fermé)', 'cuivres_tpt_sord_wah_closed', None),
    ('Horn+sordina', 'Cor + sourdine', 'cuivres_horn_sord', None),
    ('Bass_Tuba+sordina', 'Tuba basse + sourdine', 'cuivres_bass_tuba_sord', None),
]
TECH_OVERRIDE = {
    'cuivres_trb_sord_wah_open': 'ordinario_open',
    'cuivres_trb_sord_wah_closed': 'ordinario_closed',
    'cuivres_tpt_sord_wah_open': 'ordinario_open',
    'cuivres_tpt_sord_wah_closed': 'ordinario_closed',
}
all_data = {}
for csv_name, display, gfx_name, fp in INSTRUMENTS:
    tech = TECH_OVERRIDE.get(gfx_name, 'ordinario')
    d = get_formants(csv_name, tech)
    if not d:
        print(f' ⚠ PAS DE DONNÉES: {csv_name} / {tech}')
        continue
    img_path = generate_bar_chart(display, gfx_name, d['n'], d['F'], fp)
    img_rel = os.path.relpath(img_path, OUT_DIR).replace(os.sep, '/') if img_path else None
    all_data[gfx_name] = {'name': display, 'csv': csv_name, 'tech': tech, 'data': d, 'fp': fp, 'img': img_path, 'img_rel': img_rel, 'description': DESCRIPTIONS.get((csv_name, tech), '')}
    print(f" ✓ {display:<35s} N={d['n']:>4d} F1={d['F'][0]:>5d} → {img_rel}")

def fmt_hz(v):
    return '—' if v == 0 else f'{v:,}'.replace(',', ' ')

def make_technique_table_html(inst_csv_name):
    techs = sorted([(t, DATA[(i, t)]) for (i, t) in DATA if i == inst_csv_name], key=lambda x: x[0])
    if not techs:
        return ''
    rows = []
    for tech, r in techs:
        n = r['n_samples']
        fs = [round(safe_float(r[f'F{i}_hz'])) for i in range(1, 7)]
        is_ord = 'ordinario' in tech.lower() and 'to_' not in tech
        bg = ' style="background-color:#dff0d8;"' if is_ord else ''
        fvals = ''.join(f'<td{bg}>{fmt_hz(f)}</td>' for f in fs)
        rows.append(f'<tr><td{bg}><b>{tech}</b></td><td{bg}>{n}</td>{fvals}</tr>')
    return '<table class="tech-table"><tr class="header"><th>Technique</th><th>N</th><th>F1</th><th>F2</th><th>F3</th><th>F4</th><th>F5</th><th>F6</th></tr>' + ''.join(rows) + '</table>'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

def make_technique_table_docx(doc, inst_csv_name):
    techs = sorted([(t, DATA[(i, t)]) for (i, t) in DATA if i == inst_csv_name], key=lambda x: x[0])
    if not techs:
        return
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Technique', 'N', 'F1', 'F2', 'F3', 'F4', 'F5', 'F6']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], '1A3A5C')
    for tech, r in techs:
        row = table.add_row().cells
        fs = [round(safe_float(r[f'F{i}_hz'])) for i in range(1, 7)]
        vals = [tech, r['n_samples']] + [fmt_hz(f) for f in fs]
        is_ord = 'ordinario' in tech.lower() and 'to_' not in tech
        fill = 'DFF0D8' if is_ord else None
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
            if fill:
                set_cell_shading(row[idx], fill)
    widths_cm = [4.8, 1.1, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3]
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)

def instrument_section_html(gfx_key, show_all_techniques=True):
    info = all_data[gfx_key]
    fp_html = f'<p class="fp-note">Fp centroïde = {info["fp"]} Hz</p>' if info['fp'] else ''
    desc_html = f'<p class="description">{info["description"]}</p>' if info['description'] else ''
    tech_table = make_technique_table_html(info['csv']) if show_all_techniques else ''
    return f'<div class="instrument-card"><h3>{info["name"]}</h3><img src="{info["img_rel"]}" alt="{info["name"]}" class="formant-graph"/>{desc_html}{fp_html}{tech_table}</div>'

def build_html(output_path):
    html = '''<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<title>Référence Formantique — Section Cuivres</title>
<style>
body { font-family: 'Segoe UI', Helvetica, Arial, sans-serif; max-width: 1100px; margin: 0 auto; padding: 20px; background: #fafafa; color: #333; }
h1 { color: #1a237e; border-bottom: 3px solid #c62828; padding-bottom: 10px; }
h2 { color: #c62828; margin-top: 40px; border-left: 4px solid #c62828; padding-left: 12px; }
h3 { color: #b71c1c; margin-top: 30px; }
.instrument-card { background: white; border: 1px solid #ddd; border-radius: 8px; padding: 20px; margin: 20px 0; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
.formant-graph { max-width: 100%; border: 1px solid #eee; border-radius: 4px; }
.description { font-style: italic; color: #555; background: #fff8e1; padding: 10px; border-left: 3px solid #ffc107; margin: 10px 0; }
.fp-note { color: #1b5e20; font-weight: bold; }
.tech-table { width: 100%; border-collapse: collapse; margin: 15px 0; font-size: 0.9em; }
.tech-table th, .tech-table td { border: 1px solid #ccc; padding: 6px 10px; text-align: center; }
.tech-table .header th { background: #1a3a5c; color: white; }
.tech-table tr:nth-child(even) { background: #f8f8f8; }
.source-note { font-size: 0.85em; color: #888; margin-top: 30px; border-top: 1px solid #ddd; padding-top: 10px; }
.section-intro { background: #e3f2fd; padding: 15px; border-radius: 6px; margin: 15px 0; }
</style>
</head>
<body>
<h1>II. Les Cuivres</h1>
<div class="section-intro">
<p><strong>Plage formantique :</strong> 162–2 358 Hz (voyelles /u/ → /i/).</p>
<p><strong>Caractéristique :</strong> formants bien définis, grande stabilité inter-dynamiques (sauf trompette). Le <em>cluster de convergence</em> 450–502 Hz (zone /o/) rassemble Cor, Trombone, Basson, Violoncelle, Cor anglais et Saxophone ténor — fondement acoustique des doublures classiques.</p>
<p><strong>Source unique :</strong> toutes les valeurs proviennent du CSV <code>formants_all_techniques.csv</code>, extrait par le script v22 à partir des specenv SOL2020/Yan_Adds.</p>
</div>
<h2>Instruments principaux</h2>
'''
    for key in ['cuivres_bass_tuba_ord', 'cuivres_horn_ord', 'cuivres_trombone_ord', 'cuivres_trumpet_c_ord']:
        if key in all_data:
            html += instrument_section_html(key, True)
    html += '<h2>Instruments additionnels (Yan_Adds)</h2>'
    for key in ['cuivres_bass_trombone_ord', 'cuivres_contrabass_tuba_ord']:
        if key in all_data:
            html += instrument_section_html(key, True)
    html += '<h2>Cuivres avec sourdine</h2><div class="section-intro"><p>Les sourdines modifient profondément le spectre formantique des cuivres. L'effet varie selon le type de sourdine : cup, straight, harmon et wah.</p></div>'
    sordina_order = ['cuivres_trb_sord_cup', 'cuivres_trb_sord_straight', 'cuivres_trb_sord_harmon', 'cuivres_trb_sord_wah_open', 'cuivres_trb_sord_wah_closed', 'cuivres_tpt_sord_cup', 'cuivres_tpt_sord_straight', 'cuivres_tpt_sord_harmon', 'cuivres_tpt_sord_wah_open', 'cuivres_tpt_sord_wah_closed', 'cuivres_horn_sord', 'cuivres_bass_tuba_sord']
    for key in sordina_order:
        if key in all_data:
            html += instrument_section_html(key, False)
    html += '''<p class="source-note"><strong>Source des données :</strong> formants_all_techniques.csv — pipeline v22 validé.<br/>
<strong>Méthode :</strong> peak-picking sur enveloppes spectrales (seuil -30 dB, distance min 70 Hz, 6 formants max), agrégation par médiane.<br/>
<strong>Fp :</strong> centroïde spectral pondéré en énergie dans une bande optimisée par instrument.<br/>
<strong>Échantillons :</strong> SOL2020 (IRCAM) + Yan_Adds — technique ordinario uniquement sauf mention contraire.</p>
</body>
</html>'''
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)

def add_intro(doc):
    p = doc.add_paragraph()
    p.style = doc.styles['Title']
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('II. Les Cuivres')
    run.font.color.rgb = RGBColor(26, 35, 126)
    bullets = [
        ('Plage formantique : ', '162–2 358 Hz (voyelles /u/ → /i/).'),
        ('Caractéristique : ', 'formants bien définis, grande stabilité inter-dynamiques (sauf trompette). Le cluster de convergence 450–502 Hz (zone /o/) rassemble Cor, Trombone, Basson, Violoncelle, Cor anglais et Saxophone ténor.'),
        ('Source unique : ', 'toutes les valeurs proviennent du CSV formants_all_techniques.csv, extrait par le script v22 à partir des specenv SOL2020/Yan_Adds.'),
    ]
    for label, text in bullets:
        para = doc.add_paragraph(style='List Bullet')
        r1 = para.add_run(label)
        r1.bold = True
        para.add_run(text)
        for run in para.runs:
            run.font.size = Pt(10)

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 1']
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(198, 40, 40)

def add_instrument_docx(doc, info, show_all_techniques=True):
    heading = doc.add_paragraph()
    heading.style = doc.styles['Heading 2']
    run = heading.add_run(info['name'])
    run.font.color.rgb = RGBColor(183, 28, 28)
    if info['img'] and os.path.exists(info['img']):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(info['img'], width=Inches(6.7))
    if info['description']:
        p_desc = doc.add_paragraph()
        r = p_desc.add_run(info['description'])
        r.italic = True
        r.font.size = Pt(10)
    if info['fp']:
        p_fp = doc.add_paragraph()
        r = p_fp.add_run(f'Fp centroïde = {info["fp"]} Hz')
        r.bold = True
        r.font.color.rgb = RGBColor(27, 94, 32)
    if show_all_techniques:
        make_technique_table_docx(doc, info['csv'])
    doc.add_paragraph()

def build_docx(output_path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    add_intro(doc)
    add_section_heading(doc, 'Instruments principaux')
    for key in ['cuivres_bass_tuba_ord', 'cuivres_horn_ord', 'cuivres_trombone_ord', 'cuivres_trumpet_c_ord']:
        if key in all_data:
            add_instrument_docx(doc, all_data[key], True)
    add_section_heading(doc, 'Instruments additionnels (Yan_Adds)')
    for key in ['cuivres_bass_trombone_ord', 'cuivres_contrabass_tuba_ord']:
        if key in all_data:
            add_instrument_docx(doc, all_data[key], True)
    add_section_heading(doc, 'Cuivres avec sourdine')
    p = doc.add_paragraph()
    r = p.add_run('Les sourdines modifient profondément le spectre formantique des cuivres. Effets principaux : cup, straight, harmon et wah.')
    r.font.size = Pt(10)
    sordina_order = ['cuivres_trb_sord_cup', 'cuivres_trb_sord_straight', 'cuivres_trb_sord_harmon', 'cuivres_trb_sord_wah_open', 'cuivres_trb_sord_wah_closed', 'cuivres_tpt_sord_cup', 'cuivres_tpt_sord_straight', 'cuivres_tpt_sord_harmon', 'cuivres_tpt_sord_wah_open', 'cuivres_tpt_sord_wah_closed', 'cuivres_horn_sord', 'cuivres_bass_tuba_sord']
    for key in sordina_order:
        if key in all_data:
            add_instrument_docx(doc, all_data[key], False)
    p = doc.add_paragraph()
    notes = [
        'Source des données : formants_all_techniques.csv — pipeline v22 validé.',
        'Méthode : peak-picking sur enveloppes spectrales (seuil -30 dB, distance min 70 Hz, 6 formants max), agrégation par médiane.',
        'Fp : centroïde spectral pondéré en énergie dans une bande optimisée par instrument.',
        'Échantillons : SOL2020 (IRCAM) + Yan_Adds — technique ordinario uniquement sauf mention contraire.',
    ]
    for i, line in enumerate(notes):
        run = p.add_run(line + ('\n' if i < len(notes) - 1 else ''))
        run.font.size = Pt(9)
    doc.save(output_path)

html_path = os.path.join(OUT_DIR, 'section_cuivres.html')
docx_path = os.path.join(OUT_DIR, 'section_cuivres.docx')
build_html(html_path)
build_docx(docx_path)
print(f"\n{'=' * 60}")
print(f'HTML: {html_path}')
print(f'DOCX: {docx_path}')
print(f'Images: {len(all_data)} graphiques')
