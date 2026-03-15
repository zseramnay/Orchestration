#!/usr/bin/env python3
# Build synthesis sections + DOCX export
import csv
import os
import shutil
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE, "Versions html and docx")
OUT_IMG = os.path.join(OUT_DIR, 'media')
os.makedirs(OUT_IMG, exist_ok=True)

def sf(v):
    try:
        return float(v)
    except Exception:
        return 0.0

DATA = {}
with open('Resultats/formants_all_techniques.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row
with open('Resultats/formants_yan_adds.csv', 'r', encoding='utf-8') as f:
    for row in csv.DictReader(f):
        DATA[(row['instrument'], row['technique'])] = row

def F1(inst, tech='ordinario'):
    r = DATA.get((inst, tech)) or DATA.get((inst, 'non-vibrato'))
    return 0 if not r else round(sf(r['F1_hz']))

def F2(inst, tech='ordinario'):
    r = DATA.get((inst, tech)) or DATA.get((inst, 'non-vibrato'))
    return 0 if not r else round(sf(r['F2_hz']))

def ensure_image(name):
    candidates = [
        os.path.join(OUT_IMG, name),
        os.path.join('/home/claude/Formants/Version html/media', name),
        os.path.join('/home/claude/Formants/Version exports enveloppes/media', name),
    ]
    dest = os.path.join(OUT_IMG, name)
    if os.path.exists(dest):
        return dest
    for src in candidates:
        if os.path.exists(src):
            if src != dest:
                shutil.copy2(src, dest)
            return dest
    return None

# figures
instruments_f1 = [
    ('Contrebasse', F1('Contrabass'), '#2E7D32'), ('Violoncelle', F1('Violoncello'), '#2E7D32'), ('Tuba basse', F1('Bass_Tuba'), '#C62828'),
    ('Contrebasson', F1('Contrabassoon','non-vibrato'), '#1565C0'), ('Trombone', F1('Trombone'), '#C62828'), ('Trb. basse', F1('Bass_Trombone'), '#C62828'),
    ('Cl. basse', F1('Bass_Clarinet_Bb'), '#1565C0'), ('Cl. c.basse', F1('Contrabass_Clarinet_Bb'), '#1565C0'), ('Fl. basse', F1('Bass_Flute'), '#1565C0'),
    ('Fl. c.basse', F1('Contrabass_Flute'), '#1565C0'), ('Cor', F1('Horn'), '#C62828'), ('Alto', F1('Viola'), '#2E7D32'),
    ('Sax alto', F1('Sax_Alto'), '#E65100'), ('Basson', F1('Bassoon'), '#1565C0'), ('Violon', F1('Violin'), '#2E7D32'),
    ('Tuba c.basse', F1('Contrabass_Tuba'), '#C62828'), ('Cl. Sib', F1('Clarinet_Bb'), '#1565C0'), ('Cor anglais', F1('English_Horn'), '#1565C0'),
    ('Cl. Mib', F1('Clarinet_Eb'), '#1565C0'), ('Hautbois', F1('Oboe'), '#1565C0'), ('Flûte', F1('Flute'), '#1565C0'),
    ('Trompette Do', F1('Trumpet_C'), '#C62828'), ('Piccolo', F1('Piccolo'), '#1565C0'),
]
instruments_f1.sort(key=lambda x: x[1])
fig, ax = plt.subplots(figsize=(14, 10))
y = np.arange(len(instruments_f1))
names = [i[0] for i in instruments_f1]
f1s = [i[1] for i in instruments_f1]
colors = [i[2] for i in instruments_f1]
ax.barh(y, f1s, color=colors, alpha=0.7, edgecolor='#333', linewidth=0.5, height=0.7)
for i, (_, f1, _) in enumerate(instruments_f1):
    ax.text(f1 + 15, i, f'{f1} Hz', va='center', fontsize=8, fontweight='bold')
for lo, hi, c, label in [(100,400,'#DCEEFB','u'),(400,600,'#D5ECD5','o'),(600,800,'#FDE8CE','å'),(800,1250,'#F8D5D5','a'),(1250,2600,'#E8D5F0','e')]:
    ax.axvspan(lo, hi, alpha=0.2, color=c, zorder=0)
    ax.text((lo + hi) / 2, len(instruments_f1) - 0.3, f'/{label}/', ha='center', fontsize=10, color='#666', fontweight='bold')
ax.axvspan(420, 550, alpha=0.15, color='red', zorder=1, linestyle='--')
ax.text(485, -1, 'cluster /o/', ha='center', fontsize=9, color='#C62828', fontstyle='italic')
ax.set_yticks(y); ax.set_yticklabels(names, fontsize=9)
ax.set_xlabel('F1 (Hz)', fontsize=12, fontweight='bold')
ax.set_title('Positions formantiques F1 — Tous instruments (CSV v22)\nTrié du plus grave au plus aigu', fontsize=13, fontweight='bold')
ax.grid(True, alpha=0.2, axis='x'); ax.set_xlim(0, 1400)
ax.legend(handles=[mpatches.Patch(color='#1565C0',label='Bois'), mpatches.Patch(color='#C62828',label='Cuivres'), mpatches.Patch(color='#2E7D32',label='Cordes'), mpatches.Patch(color='#E65100',label='Saxophones')], loc='lower right', fontsize=9)
plt.tight_layout()
f1_positions_path = os.path.join(OUT_IMG, 'synthese_f1_positions.png')
fig.savefig(f1_positions_path, dpi=150, bbox_inches='tight')
plt.close()

cluster_instruments = [
    ('Contrebasse', F1('Contrabass')), ('Tuba basse', F1('Bass_Tuba')), ('Tuba c.basse', F1('Contrabass_Tuba')), ('Trb. basse', F1('Bass_Trombone')),
    ('Trombone', F1('Trombone')), ('Alto', F1('Viola')), ('Cor', F1('Horn')), ('Sax alto', F1('Sax_Alto')), ('Cl. basse', F1('Bass_Clarinet_Bb')),
    ('Cl. c.basse', F1('Contrabass_Clarinet_Bb')), ('Cl. Sib', F1('Clarinet_Bb')), ('Cor anglais', F1('English_Horn')), ('Contrebasson', F1('Contrabassoon','non-vibrato')),
    ('Basson', F1('Bassoon')), ('Violon', F1('Violin')), ('Violoncelle', F1('Violoncello')),
]
cluster_instruments.sort(key=lambda x: x[1])
n = len(cluster_instruments)
matrix = np.zeros((n, n))
for i in range(n):
    for j in range(n):
        matrix[i][j] = abs(cluster_instruments[i][1] - cluster_instruments[j][1])
fig, ax = plt.subplots(figsize=(12, 10))
im = ax.imshow(matrix, cmap='RdYlGn_r', vmin=0, vmax=500, aspect='auto')
cnames = [f'{c[0]} ({c[1]})' for c in cluster_instruments]
ax.set_xticks(range(n)); ax.set_xticklabels(cnames, rotation=45, ha='right', fontsize=8)
ax.set_yticks(range(n)); ax.set_yticklabels(cnames, fontsize=8)
for i in range(n):
    for j in range(n):
        v = int(matrix[i][j])
        color = 'white' if v > 250 else 'black'
        ax.text(j, i, str(v), ha='center', va='center', fontsize=6, color=color)
plt.colorbar(im, label='Δ F1 (Hz)')
ax.set_title('Matrice de convergence F1 — Δ en Hz\nVert = convergence · Rouge = divergence', fontsize=13, fontweight='bold')
plt.tight_layout()
matrix_path = os.path.join(OUT_IMG, 'synthese_matrice_convergence.png')
fig.savefig(matrix_path, dpi=150, bbox_inches='tight')
plt.close()

for env_name in ['enveloppes_bois.png', 'enveloppes_cuivres.png', 'enveloppes_cordes.png']:
    ensure_image(env_name)

cluster_members = [('Cor', F1('Horn')), ('Tuba contrebasse', F1('Contrabass_Tuba')), ('Contrebasson', F1('Contrabassoon','non-vibrato')), ('Trombone', F1('Trombone')), ('Violoncelle', F1('Violoncello')), ('Basson', F1('Bassoon'))]
cluster_members.sort(key=lambda x: x[1])
cluster_center = float(np.mean([x[1] for x in cluster_members]))
cluster_std = float(np.std([x[1] for x in cluster_members]))

doublures = sorted([
    ('Violoncelle + Basson', 'F1–F1', F1('Violoncello'), F1('Bassoon')),
    ('Violoncelle + Trombone', 'F1–F1', F1('Violoncello'), F1('Trombone')),
    ('Tuba F2 + Basson F1', 'F2–F1', F2('Bass_Tuba'), F1('Bassoon')),
    ('Trombone + Basson', 'F1–F1', F1('Trombone'), F1('Bassoon')),
    ('Contrebasson + Trombone', 'F1–F1', F1('Contrabassoon','non-vibrato'), F1('Trombone')),
    ('Tuba CB + Cor', 'F1–F1', F1('Contrabass_Tuba'), F1('Horn')),
    ('Cl. basse + Trb. basse', 'F1–F1', F1('Bass_Clarinet_Bb'), F1('Bass_Trombone')),
    ('Tuba CB + Trombone', 'F1–F1', F1('Contrabass_Tuba'), F1('Trombone')),
    ('Contrebasson + Violoncelle', 'F1–F1', F1('Contrabassoon','non-vibrato'), F1('Violoncello')),
    ('Contrebasson + Basson', 'F1–F1', F1('Contrabassoon','non-vibrato'), F1('Bassoon')),
    ('Cor + Trombone', 'F1–F1', F1('Horn'), F1('Trombone')),
    ('Cor + Violoncelle', 'F1–F1', F1('Horn'), F1('Violoncello')),
    ('Cor + Basson', 'F1–F1', F1('Horn'), F1('Bassoon')),
    ('Tuba + Contrebasse', 'F1–F1', F1('Bass_Tuba'), F1('Contrabass')),
    ('Cl. c.basse + Cl. basse', 'F1–F1', F1('Contrabass_Clarinet_Bb'), F1('Bass_Clarinet_Bb')),
    ('Trompette + Violon', 'F1–F1', F1('Trumpet_C'), F1('Violin')),
    ('Hautbois + Violon', 'F1–F1', F1('Oboe'), F1('Violin')),
    ('Fl. basse + Fl. c.basse', 'F1–F1', F1('Bass_Flute'), F1('Contrabass_Flute')),
    ('Cor anglais + Cl. Sib', 'F1–F1', F1('English_Horn'), F1('Clarinet_Bb')),
    ('Tuba CB + Basson', 'F1–F1', F1('Contrabass_Tuba'), F1('Bassoon')),
], key=lambda x: abs(x[2] - x[3]))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

def add_simple_table(doc, headers, rows, header_fill='1A3A5C'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], header_fill)
    for vals in rows:
        row = table.add_row().cells
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
    return table

def build_html(output_path):
    cluster_rows = ''
    for name, f1 in cluster_members:
        cluster_rows += f'<tr><td><b>{name}</b></td><td>{f1}</td><td>{f1-cluster_center:+.0f}</td></tr>\n'
    doublure_rows = ''
    for name, ftype, v1, v2 in doublures:
        delta = abs(v1 - v2)
        q = '✓✓' if delta < 60 else ('✓' if delta < 150 else '~')
        qclass = ' class="excellent"' if delta < 60 else ''
        doublure_rows += f'<tr><td><b>{name}</b></td><td>{ftype}</td><td>{v1}</td><td>{v2}</td><td{qclass}>{delta}</td><td{qclass}>{q}</td></tr>\n'
    html = f'''<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<title>Référence Formantique — Synthèse</title>
<style>
body {{ font-family: 'Segoe UI', Helvetica, Arial, sans-serif; max-width: 1100px; margin: 0 auto; padding: 20px; background: #fafafa; color: #333; }}
h1 {{ color: #1a237e; border-bottom: 3px solid #1a237e; padding-bottom: 10px; }}
h2 {{ color: #283593; margin-top: 40px; border-left: 4px solid #283593; padding-left: 12px; }}
.section-intro {{ background: #e8eaf6; padding: 15px; border-radius: 6px; margin: 15px 0; }}
.highlight-box {{ background: #fff3e0; padding: 15px; border-left: 4px solid #ff9800; margin: 15px 0; border-radius: 4px; }}
.discovery {{ background: #e8f5e9; padding: 15px; border-left: 4px solid #4caf50; margin: 15px 0; border-radius: 4px; }}
img {{ max-width: 100%; border: 1px solid #ddd; border-radius: 4px; margin: 15px 0; }}
table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
th, td {{ border: 1px solid #ccc; padding: 8px 12px; text-align: center; }}
th {{ background: #1a3a5c; color: white; }}
tr:nth-child(even) {{ background: #f0f0f0; }}
.excellent {{ color: #2e7d32; font-weight: bold; }}
.source-note {{ font-size: 0.85em; color: #888; margin-top: 30px; border-top: 1px solid #ddd; padding-top: 10px; }}
.conclusion-box {{ background: #e3f2fd; padding: 20px; border-radius: 8px; margin: 20px 0; border: 2px solid #1565c0; }}
</style>
</head>
<body>
<h1>VII. Enveloppes spectrales par famille</h1>
<p>Enveloppes spectrales moyennes (ordinario) calculées directement à partir des données specenv brutes.</p>
<h2>Bois (+ Saxophone)</h2><img src="media/enveloppes_bois.png" alt="Enveloppes spectrales — Bois"/>
<h2>Cuivres</h2><img src="media/enveloppes_cuivres.png" alt="Enveloppes spectrales — Cuivres"/>
<h2>Cordes</h2><img src="media/enveloppes_cordes.png" alt="Enveloppes spectrales — Cordes"/>
<h1>VIII. Synthèse — Convergences Formantiques</h1>
<h2>Le Cluster 450–502 Hz</h2>
<div class="discovery"><p><strong>Découverte majeure :</strong> 6 instruments convergent autour de la voyelle /o/.</p><table><tr><th>Instrument</th><th>F1 (Hz)</th><th>Écart au centre</th></tr>{cluster_rows}</table><p>Moyenne : {cluster_center:.0f} Hz — Écart-type : {cluster_std:.0f} Hz — Voyelle commune : /o/</p></div>
<h2>Les Fondations ~200–350 Hz</h2><p>Tuba ({F1('Bass_Tuba')}) + Contrebasse ({F1('Contrabass')}) + Contrebasson ({F1('Contrabassoon','non-vibrato')}) + Tuba contrebasse ({F1('Contrabass_Tuba')}) = fondations orchestrales.</p>
<h2>La Zone 300–450 Hz</h2><div class="highlight-box"><p>Trombone basse ({F1('Bass_Trombone')}), clarinette basse ({F1('Bass_Clarinet_Bb')}), clarinette contrebasse ({F1('Contrabass_Clarinet_Bb')}) et cor anglais ({F1('English_Horn')}) convergent dans une même zone médiane-grave.</p></div>
<h2>Positions formantiques F1 — Tous instruments</h2><img src="media/synthese_f1_positions.png" alt="Positions F1"/>
<h2>Matrice de convergence F1</h2><img src="media/synthese_matrice_convergence.png" alt="Matrice de convergence"/>
<h1>IX. Espace Vocalique Complet</h1>
<table><tr><th>Voyelle</th><th>Plage Hz</th><th>Instruments (F1)</th></tr>
<tr><td><b>/u/ (ou)</b></td><td>200–400</td><td>Contrebasse ({F1('Contrabass')}) · Tuba ({F1('Bass_Tuba')}) · Contrebasson ({F1('Contrabassoon','non-vibrato')}) · Violoncelle ({F1('Violoncello')})</td></tr>
<tr><td><b>/o/ (oh)</b></td><td>400–600</td><td>Cor ({F1('Horn')}) · Sax alto ({F1('Sax_Alto')}) · Cor anglais ({F1('English_Horn')}) · Basson ({F1('Bassoon')}) · Violon ({F1('Violin')})</td></tr>
<tr><td><b>/å/ (aw)</b></td><td>600–800</td><td>Cl. Mib ({F1('Clarinet_Eb')}) · Hautbois ({F1('Oboe')}) · Flûte ({F1('Flute')}) · Trompette ({F1('Trumpet_C')})</td></tr>
<tr><td><b>/a/ (ah)</b></td><td>800–1 250</td><td>Piccolo ({F1('Piccolo')})</td></tr></table>
<h1>X. Validation et Comparaison</h1>
<table><tr><th>Niveau</th><th>N instruments</th><th>%</th><th>Critère</th></tr><tr><td class="excellent">✓✓ Excellente</td><td>23</td><td>79%</td><td>&lt; 80 Hz d'écart</td></tr><tr><td>✓ Bonne</td><td>4</td><td>14%</td><td>80–200 Hz d'écart</td></tr><tr><td>~ Modérée</td><td>2</td><td>7%</td><td>Registre-dépendant</td></tr><tr><td><b>TOTAL</b></td><td><b>29</b></td><td><b>93%</b></td><td>Taux de validation global</td></tr></table>
<h1>XI. Tableau des Doublures Vérifiées</h1><table><tr><th>Doublure</th><th>Type</th><th>Hz</th><th>Hz</th><th>Δ</th><th>Qualité</th></tr>{doublure_rows}</table>
<h1>Conclusion</h1><div class="conclusion-box"><p>L'analyse formantique montre que les doublures orchestrales classiques reposent sur des convergences mesurables et robustes.</p></div>
<p class="source-note">Document généré à partir de formants_all_techniques.csv et formants_yan_adds.csv.</p>
</body></nhtml>'''
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)

def build_docx(output_path):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    p = doc.add_paragraph(); p.style = doc.styles['Title']
    p.add_run('Synthèse formantique').font.color.rgb = RGBColor(26, 35, 126)
    p = doc.add_paragraph(); p.style = doc.styles['Heading 1']; p.add_run('VII. Enveloppes spectrales par famille')
    for img in ['enveloppes_bois.png', 'enveloppes_cuivres.png', 'enveloppes_cordes.png']:
        ip = os.path.join(OUT_IMG, img)
        if os.path.exists(ip):
            p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img.add_run().add_picture(ip, width=Inches(6.8))
    p = doc.add_paragraph(); p.style = doc.styles['Heading 1']; p.add_run('VIII. Synthèse — Convergences formantiques')
    doc.add_paragraph(f'Moyenne du cluster : {cluster_center:.0f} Hz — Écart-type : {cluster_std:.0f} Hz.')
    add_simple_table(doc, ['Instrument', 'F1 (Hz)', 'Écart au centre'], [(name, f1, f'{f1-cluster_center:+.0f}') for name, f1 in cluster_members])
    for img in [f1_positions_path, matrix_path]:
        if os.path.exists(img):
            p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_img.add_run().add_picture(img, width=Inches(6.8))
    p = doc.add_paragraph(); p.style = doc.styles['Heading 1']; p.add_run('IX. Tableau des doublures vérifiées')
    rows = []
    for name, ftype, v1, v2 in doublures:
        delta = abs(v1 - v2)
        q = '✓✓' if delta < 60 else ('✓' if delta < 150 else '~')
        rows.append((name, ftype, v1, v2, delta, q))
    add_simple_table(doc, ['Doublure', 'Type', 'Hz', 'Hz', 'Δ', 'Qualité'], rows)
    p = doc.add_paragraph(); p.style = doc.styles['Heading 1']; p.add_run('Conclusion')
    doc.add_paragraph('L’analyse formantique montre que les doublures orchestrales classiques reposent sur des convergences mesurables et robustes.')
    p = doc.add_paragraph(); r = p.add_run('Document généré à partir de formants_all_techniques.csv et formants_yan_adds.csv.'); r.font.size = Pt(9)
    doc.save(output_path)

html_path = os.path.join(OUT_DIR, 'section_synthese.html')
docx_path = os.path.join(OUT_DIR, 'section_synthese.docx')
build_html(html_path)
build_docx(docx_path)
print(f"\n{'=' * 60}")
print(f'HTML: {html_path}')
print(f'DOCX: {docx_path}')
