#!/usr/bin/env python3
"""Spectral envelope images — one image per instrument, including all mutes.
Generates section_enveloppes_v6.html + .docx"""
import os, unicodedata, re
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SAMPLE_RATE = 44100
FFT_SIZE = 4096
FREQ_RES = SAMPLE_RATE / FFT_SIZE
FORMANT_MIN_BIN = int(80 / FREQ_RES)
FORMANT_MAX_BIN = int(6000 / FREQ_RES)
MIN_PEAK_DIST_HZ = 70
PEAK_THRESHOLD_DB = 30
MAX_FORMANTS = 6
freqs = np.arange(1024) * FREQ_RES

BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
SOL = os.path.join(BASE, 'Data', 'FullSOL2020_specenv par instrument')
YAN = os.path.join(BASE, 'Data', 'Yan_Adds-Divers_specenv par instrument')
OUT_DIR = os.path.join(BASE, "Etude-Formants", "Versions-html-and-docx")
OUT = os.path.join(OUT_DIR, 'media')
os.makedirs(OUT, exist_ok=True)

def find_formant_peaks_v22(envelope):
    env_db = list(envelope)
    lo = max(FORMANT_MIN_BIN, 1); hi = min(FORMANT_MAX_BIN, len(env_db) - 1)
    if hi <= lo: return []
    region = env_db[lo:hi]; mx = max(region); th = mx - PEAK_THRESHOLD_DB
    peaks = []
    for i in range(1, len(region) - 1):
        v = region[i]
        if v < th: continue
        if v >= region[i-1] and v >= region[i+1]:
            peaks.append(((lo+i)*FREQ_RES, v))
    peaks.sort(key=lambda x: x[1], reverse=True)
    sel = []
    for f, a in peaks:
        if not any(abs(f-sf) < MIN_PEAK_DIST_HZ for sf,_ in sel):
            sel.append((f, a))
        if len(sel) >= MAX_FORMANTS: break
    sel.sort(key=lambda x: x[0])
    return sel

def compute_band_centroid(envelope, lo_hz, hi_hz):
    lo_bin = max(0, min(int(lo_hz/FREQ_RES), len(envelope)-1))
    hi_bin = max(0, min(int(hi_hz/FREQ_RES), len(envelope)-1))
    region = envelope[lo_bin:hi_bin+1]
    linear = np.array([10**(v/10) for v in region])
    total = linear.sum()
    if total <= 0: return None
    return sum(linear[i]*(lo_bin+i)*FREQ_RES for i in range(len(linear)))/total

def load_ordinario(filepath, techs=('ordinario',)):
    samples = []
    with open(filepath, encoding='utf-8') as f:
        f.readline()
        for line in f:
            parts = line.strip().split(';')
            if len(parts) < 10: continue
            path = parts[0]
            if not path.startswith('/'): continue
            pp = path.split('/')
            if len(pp) < 5 or pp[3] not in techs: continue
            try: vals = [float(v) for v in parts[1:]]
            except: continue
            if len(vals) < 100: continue
            formants = find_formant_peaks_v22(vals)
            if len(formants) >= 1:
                samples.append({'formants': formants, 'envelope': np.array(vals)})
    return samples

def make_slug(display):
    s = unicodedata.normalize('NFD', display)
    s = ''.join(c for c in s if unicodedata.category(c) != 'Mn')
    s = s.lower()
    s = re.sub(r'[^a-z0-9]+', '_', s)
    return s.strip('_')

# ─── Complete instrument list with all mutes ───
# (display, dir, filename, techs, family, fp_band)
INSTRUMENTS = [
    # Bois
    ('Piccolo',                  YAN, 'Yan_Adds-Divers_specenv.db_Piccolo.txt',               ('ordinario',), 'Bois', (400,1000)),
    ('Flûte',                    SOL, 'FullSOL2020_specenv.db_Flute.txt',                     ('ordinario',), 'Bois', (1000,2000)),
    ('Flûte basse',              YAN, 'Yan_Adds-Divers_specenv.db_Bass-Flute.txt',            ('ordinario',), 'Bois', (1000,2000)),
    ('Flûte contrebasse',        YAN, 'Yan_Adds-Divers_specenv.db_Contrabass-Flute.txt',      ('ordinario',), 'Bois', (800,1600)),
    ('Hautbois',                 SOL, 'FullSOL2020_specenv.db_Oboe.txt',                      ('ordinario',), 'Bois', (1000,2000)),
    ('Cor anglais',              YAN, 'Yan_Adds-Divers_specenv.db_EnglishHorn.txt',           ('ordinario',), 'Bois', (800,1600)),
    ('Clarinette en Mib',        YAN, 'Yan_Adds-Divers_specenv.db_Clarinet-Eb.txt',           ('ordinario',), 'Bois', (800,1600)),
    ('Clarinette en Sib',        SOL, 'FullSOL2020_specenv.db_Clarinet_Bb.txt',               ('ordinario',), 'Bois', (1000,2000)),
    ('Clarinette basse',         YAN, 'Yan_Adds-Divers_specenv.db_Bass-Clarinet-Bb.txt',      ('ordinario',), 'Bois', (800,1600)),
    ('Clarinette contrebasse',   YAN, 'Yan_Adds-Divers_specenv.db_Contrabass-Clarinet-Bb.txt',('ordinario',), 'Bois', (600,1400)),
    ('Basson',                   SOL, 'FullSOL2020_specenv.db_Bassoon.txt',                   ('ordinario',), 'Bois', (800,1600)),
    ('Contrebasson',             YAN, 'Yan_Adds-Divers_specenv.db_Contrebasson.txt',          ('non-vibrato',),'Bois',(1000,2000)),
    ('Saxophone alto',           SOL, 'FullSOL2020_specenv.db_Sax_Alto.txt',                  ('ordinario',), 'Bois', (1000,2000)),
    # Cuivres
    ('Cor en Fa',                SOL, 'FullSOL2020_specenv.db_Horn.txt',                      ('ordinario',), 'Cuivres', (600,1400)),
    ('Cor+sourdine',             SOL, 'FullSOL2020_specenv.db_Horn+sordina.txt',              ('ordinario',), 'Cuivres', (600,1400)),
    ('Trompette en Ut',          SOL, 'FullSOL2020_specenv.db_Trumpet_C.txt',                 ('ordinario',), 'Cuivres', (600,1400)),
    ('Trompette+sourd. cup',     SOL, 'FullSOL2020_specenv.db_Trumpet_C+sordina_cup.txt',     ('ordinario',), 'Cuivres', (800,1800)),
    ('Trompette+sourd. sèche',   SOL, 'FullSOL2020_specenv.db_Trumpet_C+sordina_straight.txt',('ordinario',), 'Cuivres', (800,1800)),
    ('Trompette+sourd. harmon',  SOL, 'FullSOL2020_specenv.db_Trumpet_C+sordina_harmon.txt',  ('ordinario',), 'Cuivres', (800,1800)),
    ('Trompette+sourd. wah',     SOL, 'FullSOL2020_specenv.db_Trumpet_C+sordina_wah.txt',     ('ordinario_open',),'Cuivres',(600,1400)),
    ('Trombone ténor',           SOL, 'FullSOL2020_specenv.db_Trombone.txt',                  ('ordinario',), 'Cuivres', (1000,2000)),
    ('Trombone+sourd. cup',      SOL, 'FullSOL2020_specenv.db_Trombone+sordina_cup.txt',      ('ordinario',), 'Cuivres', (800,1800)),
    ('Trombone+sourd. sèche',    SOL, 'FullSOL2020_specenv.db_Trombone+sordina_straight.txt', ('ordinario',), 'Cuivres', (800,1800)),
    ('Trombone+sourd. harmon',   SOL, 'FullSOL2020_specenv.db_Trombone+sordina_harmon.txt',   ('ordinario',), 'Cuivres', (800,1800)),
    ('Trombone+sourd. wah',      SOL, 'FullSOL2020_specenv.db_Trombone+sordina_wah.txt',      ('ordinario_open',),'Cuivres',(800,1800)),
    ('Trombone basse',           YAN, 'Yan_Adds-Divers_specenv.db_Bass-Trombone.txt',         ('ordinario',), 'Cuivres', (1000,2000)),
    ('Tuba basse',               SOL, 'FullSOL2020_specenv.db_Bass_Tuba.txt',                 ('ordinario',), 'Cuivres', (1000,2000)),
    ('Tuba basse+sourdine',      SOL, 'FullSOL2020_specenv.db_Bass_Tuba+sordina.txt',         ('ordinario',), 'Cuivres', (1000,2000)),
    ('Tuba contrebasse',         YAN, 'Yan_Adds-Divers_specenv.db_Contrabass-Tuba.txt',       ('ordinario',), 'Cuivres', (1000,2000)),
    # Cordes solistes
    ('Violon',                   SOL, 'FullSOL2020_specenv.db_Violin.txt',                    ('ordinario',), 'Cordes', (600,1400)),
    ('Violon+sourdine',          SOL, 'FullSOL2020_specenv.db_Violin+sordina.txt',            ('ordinario',), 'Cordes', (600,1400)),
    ('Violon+sourd. piombo',     SOL, 'FullSOL2020_specenv.db_Violin+sordina_piombo.txt',     ('ordinario',), 'Cordes', (600,1400)),
    ('Alto',                     SOL, 'FullSOL2020_specenv.db_Viola.txt',                     ('ordinario',), 'Cordes', (800,1600)),
    ('Alto+sourdine',            SOL, 'FullSOL2020_specenv.db_Viola+sordina.txt',             ('ordinario',), 'Cordes', (800,1600)),
    ('Alto+sourd. piombo',       SOL, 'FullSOL2020_specenv.db_Viola+sordina_piombo.txt',      ('ordinario',), 'Cordes', (800,1600)),
    ('Violoncelle',              SOL, 'FullSOL2020_specenv.db_Violoncello.txt',               ('ordinario',), 'Cordes', (600,1400)),
    ('Violoncelle+sourdine',     SOL, 'FullSOL2020_specenv.db_Violoncello+sordina.txt',       ('ordinario',), 'Cordes', (600,1400)),
    ('Violoncelle+sourd. piombo',SOL, 'FullSOL2020_specenv.db_Violoncello+sordina_piombo.txt',('ordinario',), 'Cordes', (600,1400)),
    ('Contrebasse',              SOL, 'FullSOL2020_specenv.db_Contrabass.txt',                ('ordinario',), 'Cordes', (1000,2000)),
    ('Contrebasse+sourdine',     SOL, 'FullSOL2020_specenv.db_Contrabass+sordina.txt',        ('ordinario',), 'Cordes', (1000,2000)),
    # Cordes ensembles
    ('Ens. violons',             YAN, 'Yan_Adds-Divers_specenv.db_Violin-Ensemble.txt',       ('ordinario',), 'Cordes', (600,1400)),
    ('Ens. violons+sourdine',    YAN, 'Yan_Adds-Divers_specenv.db_Violin-Ensemble-sordina.txt',('ordinario',),'Cordes',(600,1400)),
    ('Ens. altos',               YAN, 'Yan_Adds-Divers_specenv.db_Viola-Ensemble.txt',        ('ordinario',), 'Cordes', (600,1400)),
    ('Ens. altos+sourdine',      YAN, 'Yan_Adds-Divers_specenv.db_Viola-Ensemble-sordina.txt',('ordinario',), 'Cordes', (600,1400)),
    ('Ens. violoncelles',        YAN, 'Yan_Adds-Divers_specenv.db_Violoncello-Ensemble.txt',  ('ordinario',), 'Cordes', (1000,2000)),
    ('Ens. violoncelles+sourdine',YAN,'Yan_Adds-Divers_specenv.db_Violoncello-Ensemble-sordina.txt',('ordinario',),'Cordes',(1000,2000)),
    ('Ens. contrebasses',        YAN, 'Yan_Adds-Divers_specenv.db_Contrabass-Ensemble.txt',   ('non-vibrato',),'Cordes',(800,1600)),
]

FAMILY_COLORS = {'Bois': '#1565C0', 'Cuivres': '#C62828', 'Cordes': '#2E7D32'}

# ─── Generate individual images ───
summaries = []
for display, directory, fname, techs, family, fp_band in INSTRUMENTS:
    filepath = os.path.join(directory, fname)
    if not os.path.exists(filepath):
        print(f"  ⚠ Fichier manquant: {display} -> {filepath}")
        continue
    samples = load_ordinario(filepath, techs)
    if not samples:
        print(f"  ⚠ Pas de donnees: {display}")
        continue
    color = FAMILY_COLORS[family]
    slug = make_slug(display)
    img_name = f'specenv_{slug}.png'
    fig, ax = plt.subplots(figsize=(8, 3.5), dpi=150)
    envs = [s['envelope'] for s in samples]
    mean_env = np.mean(envs, axis=0)
    mask = freqs <= 5500
    ax.plot(freqs[mask], mean_env[mask], color=color, linewidth=2)
    if len(envs) > 3:
        std_env = np.std(envs, axis=0)
        ax.fill_between(freqs[mask], (mean_env-std_env)[mask], (mean_env+std_env)[mask], color=color, alpha=0.1)
    for fi in range(min(4, max(len(s['formants']) for s in samples))):
        fi_freqs = sorted([s['formants'][fi][0] for s in samples if fi < len(s['formants'])])
        if fi_freqs:
            fval = fi_freqs[len(fi_freqs)//2]
            bin_idx = min(int(fval/FREQ_RES), len(mean_env)-1)
            ax.plot(fval, mean_env[bin_idx], 'v', color='#D32F2F', markersize=6, zorder=5)
            yo = 3 if fi%2==0 else -5
            ax.annotate(f'F{fi+1}\n{fval:.0f}', (fval, mean_env[bin_idx]+yo), fontsize=6, ha='center', color='#D32F2F', fontweight='bold')
    fps = [compute_band_centroid(s['envelope'], fp_band[0], fp_band[1]) for s in samples]
    fps = [f for f in fps if f]
    fp_med = None
    if fps:
        fp_med = np.median(fps)
        fp_std_val = np.std(fps)
        ax.axvline(fp_med, color='#2E7D32', linewidth=2, alpha=0.7)
        bin_fp = min(int(fp_med/FREQ_RES), len(mean_env)-1)
        ax.text(fp_med, mean_env[bin_fp]+3, f'Fp={fp_med:.0f}\n(σ={fp_std_val:.0f})', fontsize=7, ha='center',
                color='#2E7D32', fontweight='bold', bbox=dict(boxstyle='round,pad=0.2', facecolor='white', alpha=0.8, edgecolor='#2E7D32'))
    f2_all = [s['formants'][1][0] for s in samples if len(s['formants']) > 1]
    f2_std = np.std(f2_all) if f2_all else 0
    ax.set_title(f'{display}  |  n={len(samples)}  σ(F2)={f2_std:.0f} Hz', fontsize=10, fontweight='bold', color=color)
    ax.set_xlabel('Hz', fontsize=8); ax.set_ylabel('dB', fontsize=8)
    ax.grid(True, alpha=0.2); ax.set_xlim(50, 5500); ax.tick_params(labelsize=7)
    plt.tight_layout()
    outpath = os.path.join(OUT, img_name)
    fig.savefig(outpath, dpi=150, bbox_inches='tight'); plt.close(fig)
    fp_str = f'{fp_med:.0f}' if fp_med else '---'
    summaries.append({'display': display, 'family': family, 'slug': slug, 'img_name': img_name, 'img_rel': f'media/{img_name}', 'n': len(samples), 'fp': fp_med, 'f2_std': f2_std})
    print(f"  ok {display:40s} n={len(samples):>4d}  Fp={fp_str:>6s}")

# ─── Build HTML ───
def build_html(output_path):
    current_family = None; parts = []
    family_intros = {
        'Bois': "Enveloppes spectrales moyennes des bois, du piccolo à la flûte contrebasse.",
        'Cuivres': "Enveloppes spectrales des cuivres, y compris toutes les sourdines.",
        'Cordes': "Enveloppes spectrales des cordes solistes, ensembles, et toutes les sourdines (classiques et piombo).",
    }
    family_slugs = {'Bois': 'env-bois', 'Cuivres': 'env-cuivres', 'Cordes': 'env-cordes'}
    for s in summaries:
        if s['family'] != current_family:
            current_family = s['family']
            sid = family_slugs.get(current_family, current_family.lower())
            parts.append(f'<h2 id="{sid}">{current_family}</h2>')
            parts.append(f'<p>{family_intros.get(current_family, "")}</p>')
        parts.append(f'<h3>{s["display"]}</h3>')
        parts.append(f'<img src="{s["img_rel"]}" alt="Enveloppe {s["display"]}"/>')
    html = f'''<!DOCTYPE html>
<html lang="fr"><head><meta charset="UTF-8">
<title>Reference Formantique - Enveloppes spectrales</title>
<style>
body {{ font-family: 'Segoe UI', Helvetica, Arial, sans-serif; max-width: 1100px; margin: 0 auto; padding: 20px; background: #fafafa; color: #333; }}
h1 {{ color: #1a237e; border-bottom: 3px solid #1a237e; padding-bottom: 10px; }}
h2 {{ color: #283593; margin-top: 40px; border-left: 4px solid #283593; padding-left: 12px; }}
h3 {{ color: #444; margin-top: 20px; }}
.section-intro {{ background: #e8eaf6; padding: 15px; border-radius: 6px; margin: 15px 0; }}
img {{ max-width: 100%; height: auto; border: 1px solid #ddd; border-radius: 4px; margin: 8px 0; }}
.source-note {{ font-size: 0.85em; color: #888; margin-top: 30px; border-top: 1px solid #ddd; padding-top: 10px; }}
</style></head><body>
<h1 id="vii-enveloppes">VII. Enveloppes spectrales par instrument</h1>
<div class="section-intro">
<p>Enveloppes spectrales moyennes (ordinario) calculées directement à partir des données specenv Orchidea.
Marqueurs F1–F4 (rouge), Fp centroïde (vert), bande ±1σ.
Inclut toutes les sourdines cuivres et cordes.</p>
</div>
{''.join(parts)}
<p class="source-note">Sources : exports specenv bruts SOL2020 + Yan_Adds-Divers.</p>
</body></html>'''
    with open(output_path, 'w', encoding='utf-8') as f: f.write(html)

def build_docx(output_path):
    doc = Document()
    sec = doc.sections[0]; sec.top_margin = Cm(2); sec.bottom_margin = Cm(2); sec.left_margin = Cm(2); sec.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(10)
    p = doc.add_paragraph(); p.style = doc.styles['Title']
    r = p.add_run('VII. Enveloppes spectrales par instrument'); r.font.color.rgb = RGBColor(26, 35, 126)
    doc.add_paragraph('Enveloppes spectrales moyennes (ordinario). Marqueurs F1-F4, Fp centroïde, bande ±1σ. Inclut toutes les sourdines.')
    current_family = None
    for s in summaries:
        if s['family'] != current_family:
            current_family = s['family']
            p = doc.add_paragraph(); p.style = doc.styles['Heading 1']; p.add_run(current_family)
        p2 = doc.add_paragraph(); p2.style = doc.styles['Heading 2']; p2.add_run(s['display'])
        img_path = os.path.join(OUT, s['img_name'])
        if os.path.exists(img_path):
            p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(img_path, width=Inches(6.0))
    doc.save(output_path)

html_path = os.path.join(OUT_DIR, 'section_enveloppes_v6.html')
docx_path = os.path.join(OUT_DIR, 'section_enveloppes_v6.docx')
build_html(html_path)
build_docx(docx_path)
print(f"\n{'='*60}")
print(f'HTML: {html_path}')
print(f'DOCX: {docx_path}')
print(f'Images: {len(summaries)} instruments')
