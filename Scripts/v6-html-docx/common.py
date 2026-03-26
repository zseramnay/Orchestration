#!/usr/bin/env python3
"""
common.py — Module partagé pour tous les scripts build v4
Contient : chargement CSV, graphiques, utilitaires HTML/DOCX, CSS global.
"""
import csv
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D
import matplotlib.ticker
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─── Chemins ────────────────────────────────────────────────
# Ce module est dans Formants/Scripts/v4-html-docx-enriched/
# BASE = Formants/
BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUT_DIR = os.path.join(BASE, "Versions-html-and-docx")
OUT_IMG = os.path.join(OUT_DIR, "media")
os.makedirs(OUT_IMG, exist_ok=True)

# ─── Chargement CSV ─────────────────────────────────────────
NAME_MAPPING = {
    'Bass_Clarinet_Bb': 'Clarinette basse en Sib',
    'Bass_Flute': 'Flûte basse',
    'Bass_Trombone': 'Trombone basse',
    'Bass_Tuba': 'Tuba basse',
    'Bass_Tuba+sordina': 'Tuba basse+sourdine',
    'Bassoon': 'Basson',
    'Bassoon+sordina': 'Basson+sourdine',
    'Clarinet_Bb': 'Clarinette en Sib',
    'Clarinet_Eb': 'Clarinette en Mib',
    'Contrabass': 'Contrebasse',
    'Contrabass+sordina': 'Contrebasse+sourdine',
    'Contrabass_Clarinet_Bb': 'Clarinette contrebasse en Sib',
    'Contrabass_Ensemble': 'Ensemble de contrebasses',
    'Contrabass_Flute': 'Flûte contrebasse',
    'Contrabass_Tuba': 'Tuba contrebasse',
    'Contrabassoon': 'Contrebasson',
    'English_Horn': 'Cor anglais',
    'Flute': 'Flûte',
    'Horn': 'Cor en Fa',
    'Horn+sordina': 'Cor en Fa+sourdine',
    'Oboe': 'Hautbois',
    'Oboe+sordina': 'Hautbois+sourdine',
    'Piccolo': 'Petite flûte',
    'Sax_Alto': 'Saxophone alto',
    'Trombone': 'Trombone ténor',
    'Trombone+sordina_cup': 'Trombone+sourdine cup',
    'Trombone+sordina_harmon': 'Trombone+sourdine harmon',
    'Trombone+sordina_straight': 'Trombone+sourdine sèche',
    'Trombone+sordina_wah': 'Trombone+sourdine wah',
    'Trumpet_C': 'Trompette en Ut',
    'Trumpet_C+sordina_cup': 'Trompette+sourdine cup',
    'Trumpet_C+sordina_harmon': 'Trompette+sourdine harmon',
    'Trumpet_C+sordina_straight': 'Trompette+sourdine sèche',
    'Trumpet_C+sordina_wah': 'Trompette+sourdine wah',
    'Viola': 'Alto',
    'Viola+sordina': 'Alto+sourdine',
    'Viola+sordina_piombo': 'Alto+sourdine piombo',
    'Viola_Ensemble': "Ensemble d'altos",
    'Viola_Ensemble+sordina': "Ensemble d'altos+sourdine",
    'Violin': 'Violon',
    'Violin+sordina': 'Violon+sourdine',
    'Violin+sordina_piombo': 'Violon+sourdine piombo',
    'Violin_Ensemble': 'Ensemble de violons',
    'Violin_Ensemble+sordina': 'Ensemble de violons+sourdine',
    'Violoncello': 'Violoncelle',
    'Violoncello+sordina': 'Violoncelle+sourdine',
    'Violoncello+sordina_piombo': 'Violoncelle+sourdine piombo',
    'Violoncello_Ensemble': 'Ensemble de violoncelles',
    'Violoncello_Ensemble+sordina': 'Ensemble de violoncelles+sourdine',
}
V2_TO_LEGACY = {v: k for k, v in NAME_MAPPING.items()}

DATA = {}
BW_DATA = {}  # (inst, tech) -> [F1_bw, ..., F6_bw]

def sf(v):
    try:
        return float(v)
    except Exception:
        return 0.0

def load_data(path):
    if not os.path.exists(path):
        print(f"  ⚠ CSV introuvable : {path}")
        return
    with open(path, 'r', encoding='utf-8') as f:
        for row in csv.DictReader(f):
            inst = row['instrument']
            tech = row['technique']
            DATA[(inst, tech)] = row
            if inst in NAME_MAPPING:
                DATA[(NAME_MAPPING[inst], tech)] = row
            if inst in V2_TO_LEGACY:
                DATA[(V2_TO_LEGACY[inst], tech)] = row

def load_bandwidths():
    bw_path = os.path.join(BASE, 'Resultats/bandwidths_3db.csv')
    if not os.path.exists(bw_path):
        print(f"  ⚠ Bandwidths CSV introuvable : {bw_path}")
        return
    with open(bw_path, 'r', encoding='utf-8') as f:
        for row in csv.DictReader(f):
            inst = row['instrument']
            tech = row['technique']
            bws = [sf(row.get(f'F{i}_bw', '')) for i in range(1, 7)]
            BW_DATA[(inst, tech)] = bws
            if inst in NAME_MAPPING:
                BW_DATA[(NAME_MAPPING[inst], tech)] = bws
            if inst in V2_TO_LEGACY:
                BW_DATA[(V2_TO_LEGACY[inst], tech)] = bws

def load_all_csvs():
    csv_all = os.path.join(BASE, 'Resultats/formants_all_techniques_v3.csv')
    csv_yan = os.path.join(BASE, 'Resultats/formants_yan_adds_v3.csv')
    if not os.path.exists(csv_all):
        csv_all = os.path.join(BASE, 'Resultats/formants_all_techniques_v2.csv')
    if not os.path.exists(csv_yan):
        csv_yan = os.path.join(BASE, 'Resultats/formants_yan_adds_v2.csv')
    load_data(csv_all)
    load_data(csv_yan)
    load_bandwidths()

def get_f(inst, tech):
    r = DATA.get((inst, tech))
    if not r:
        return None
    bw = BW_DATA.get((inst, tech), [0]*6)
    return {
        'n': int(r['n_samples']),
        'F': [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)],
        'dB': [sf(r.get(f'F{i}_db', '')) for i in range(1, 7)],
        'q25': [sf(r.get(f'F{i}_q25', '')) for i in range(1, 4)] + [0, 0, 0],
        'q75': [sf(r.get(f'F{i}_q75', '')) for i in range(1, 4)] + [0, 0, 0],
        'std': [sf(r.get(f'F{i}_std', '')) for i in range(1, 7)],
        'bw': bw,
    }

def fmt_hz(v):
    return '—' if v == 0 else f"{int(v):,}".replace(',', '\u202f')

# ─── Zones vocaliques ───────────────────────────────────────
VOWEL_ZONES = [
    (100,  400,  '#DCEEFB', 'u (oo)\nProfondeur'),
    (400,  600,  '#D5ECD5', 'o (oh)\nPlénitude'),
    (600,  800,  '#FDE8CE', 'å (aw)\nTransition'),
    (800,  1250, '#F8D5D5', 'a (ah)\nPuissance'),
    (1250, 2600, '#E8D5F0', 'e (eh)\nClarté'),
    (2600, 6000, '#FFF8D0', 'i (ee)\nBrillance'),
]
FC = ['#D32F2F', '#E64A19', '#F57C00', '#FFA000', '#FBC02D', '#CDDC39']
FA = [1.0, 0.85, 0.7, 0.55, 0.4, 0.3]

# ─── Génération graphiques ───────────────────────────────────
def make_graph(display, filename, n, formants, fp=None, family_color='#2E7D32', family_label='',
               amplitudes=None, bandwidths=None, **_ignored):
    """
    Gaussian bell curves with real -3dB bandwidths and automatic anti-collision labels.
    """
    valid = [(i, f) for i, f in enumerate(formants) if f > 0]
    if not valid:
        return None

    mf = max(f for _, f in valid) + 500
    mf = min(max(mf, 3000), 6500)

    fig, ax = plt.subplots(figsize=(9.6, 4.8), dpi=150)

    for lo, hi, c, l in VOWEL_ZONES:
        if lo < mf:
            ax.axvspan(lo, min(hi, mf), alpha=0.30, color=c, zorder=0)
            mid = (lo + min(hi, mf)) / 2
            if mid < mf * 0.95:
                ax.text(mid, 0.97, l, ha='center', va='top', fontsize=7,
                        color='#666', fontweight='bold',
                        transform=ax.get_xaxis_transform())

    # Heights from dB
    use_real_amp = amplitudes is not None and any(amplitudes[i] != 0 for i, _ in valid)
    if use_real_amp:
        max_db = max(amplitudes[i] for i, _ in valid)
        heights = {i: max(0.05, 10**((amplitudes[i]-max_db)/20.0)) for i, _ in valid}
    else:
        heights = {i: 1.0*(1.0-i*0.12) for i, _ in valid}

    # Bell curves
    x = np.linspace(80, mf, 2000)
    has_bw = bandwidths is not None and any(bandwidths[i] > 0 for i, _ in valid)
    for i, freq in valid:
        h = heights[i]
        sigma = bandwidths[i]/2.355 if has_bw and bandwidths[i] > 0 else max(40, freq*0.08)
        y = h * np.exp(-0.5*((x-freq)/sigma)**2)
        ax.fill_between(x, y, alpha=0.25, color=FC[i], zorder=2)
        ax.plot(x, y, color=FC[i], lw=1.8, alpha=0.85, zorder=3)
        ax.plot(freq, h, 'o', color=FC[i], markersize=6,
                markeredgecolor='#333', markeredgewidth=0.8, zorder=5)

    # Cluster
    ax.axvspan(420, 550, alpha=0.12, color='red', zorder=1)
    ax.text(485, 0.02, 'cluster /o/', ha='center', va='bottom', fontsize=7,
            color='#C62828', fontstyle='italic',
            transform=ax.get_xaxis_transform())

    # Fp diamond — ALWAYS shown when available, no arrow
    f1 = formants[0]
    fp_h = None
    if fp:
        fp_h = heights.get(0, 0.5) * 0.55 if use_real_amp else 0.5
        ax.plot(fp, fp_h, marker='D', markersize=7, color='#1B5E20',
                markeredgecolor='black', markeredgewidth=1.0, zorder=7)

    # ── Anti-collision v6: vertical-first, small horizontal fallback ──
    log_range = np.log10(mf) - np.log10(100)
    LABEL_W = 0.065 * log_range
    LABEL_H = 0.075
    ylim_top = 1.55
    MAX_LABEL_Y = ylim_top - 0.10

    # Detect dense spectra: if all 6 formants span less than 1 octave in log space
    freqs_valid = [f for _, f in valid]
    log_span = np.log10(max(freqs_valid)) - np.log10(min(freqs_valid)) if len(freqs_valid) > 1 else 1.0
    dense = log_span < 0.5  # less than ~1.5 octaves for all formants
    if dense:
        LABEL_W *= 0.75
        LABEL_H *= 0.75

    labels = []
    for i, freq in valid:
        h = heights[i]
        if use_real_amp:
            if i >= 3 or dense:
                txt = f"F{i+1} {freq} Hz\n({amplitudes[i]:.0f} dB)"
                nlines = 2; fs = 5.5 if dense else 6.5
            else:
                txt = f"F{i+1}\n{freq} Hz\n({amplitudes[i]:.0f} dB)"
                nlines = 3; fs = 6 if dense else 7
        else:
            txt = f"F{i+1}\n{freq} Hz"
            nlines = 2; fs = 6 if dense else 7
        labels.append({'x': freq, 'y': h, 'text': txt, 'color': '#333',
                       'fs': fs, 'priority': 6-i, 'type': 'formant', 'nlines': nlines})

    if fp and fp_h is not None:
        labels.append({'x': fp, 'y': fp_h, 'text': f"Fp = {fp} Hz\n(centroïde)",
                       'color': '#1B5E20', 'fs': 7.5, 'priority': 7, 'type': 'fp', 'nlines': 2})

    labels.sort(key=lambda l: -l['priority'])
    placed = []

    def _collides(lx, ty, lh):
        for px, py_bot, py_top, pw in placed:
            if abs(lx - px) < (LABEL_W + pw) * 0.5:
                if ty < py_top + 0.008 and (ty + lh) > py_bot - 0.008:
                    return True
        return False

    for lab in labels:
        lx = np.log10(lab['x'])
        lh = lab['nlines'] * LABEL_H
        base_y = lab['y'] + 0.04
        best = None

        # PASS 1: vertical only (center x) — strongly preferred
        for yp in [0, 0.06, 0.12, 0.18, 0.24, 0.30, 0.36, 0.42, 0.48, 0.54]:
            ty = base_y + yp
            if ty + lh > MAX_LABEL_Y or ty < 0.04:
                continue
            if not _collides(lx, ty, lh):
                best = (lx, ty)
                break

        # PASS 2: small horizontal shifts only if vertical-center failed
        if best is None:
            for ox in [-LABEL_W*0.6, LABEL_W*0.6]:
                tx = lx + ox
                if tx < np.log10(100) or tx > np.log10(mf):
                    continue
                for yp in [0, 0.06, 0.12, 0.18, 0.24, 0.30, 0.36, 0.42]:
                    ty = base_y + yp
                    if ty + lh > MAX_LABEL_Y or ty < 0.04:
                        continue
                    if not _collides(tx, ty, lh):
                        best = (tx, ty)
                        break
                if best:
                    break

        if best is None:
            best = (lx, min(base_y, MAX_LABEL_Y - lh))

        best_lx, best_y = best
        best_x = 10 ** best_lx
        placed.append((best_lx, best_y, best_y + lh, LABEL_W))

        dx = abs(best_lx - lx)
        dy = abs(best_y - base_y)
        clr = lab['color']
        if dx > 0.01 or dy > 0.03:
            ax.annotate(lab['text'], xy=(lab['x'], lab['y']),
                        xytext=(best_x, best_y),
                        ha='center', va='bottom', fontsize=lab['fs'],
                        fontweight='bold', color=clr,
                        arrowprops=dict(arrowstyle='->', color='#999', lw=0.8, ls='--'),
                        zorder=6)
        else:
            ax.text(best_x, best_y, lab['text'],
                    ha='center', va='bottom', fontsize=lab['fs'],
                    fontweight='bold', color=clr, zorder=6)

    # Axes
    ax.set_xlim(100, mf); ax.set_ylim(0, ylim_top)
    ax.set_xlabel("Fréquence (Hz)", fontsize=10, fontweight='bold')
    ax.set_ylabel("Amplitude relative (normalisée)" if use_real_amp else "Importance relative",
                  fontsize=10, fontweight='bold')
    ax.set_title(f"{display} — Profil formantique moyen (ordinario, N={n})",
                 fontsize=12, fontweight='bold', color=family_color, pad=12)
    ax.set_xscale('log')
    ticks = [t for t in [100,150,200,300,400,500,600,800,1000,1500,2000,3000,4000,5000,6000] if t<=mf]
    ax.set_xticks(ticks)
    ax.set_xticklabels([str(t) for t in ticks], fontsize=8)
    ax.get_xaxis().set_major_formatter(matplotlib.ticker.ScalarFormatter())
    ax.set_yticks([])
    for s in ['top','right','left']: ax.spines[s].set_visible(False)

    # Legend — lower left
    if use_real_amp and has_bw:
        le = [mpatches.Patch(facecolor=FC[i], alpha=0.4, edgecolor=FC[i],
              label=f'F{i+1} = {formants[i]} Hz ({amplitudes[i]:.0f} dB) BW={bandwidths[i]:.0f}')
              for i, _ in valid]
    elif use_real_amp:
        le = [mpatches.Patch(facecolor=FC[i], alpha=0.4, edgecolor=FC[i],
              label=f'F{i+1} = {formants[i]} Hz ({amplitudes[i]:.0f} dB)') for i, _ in valid]
    else:
        le = [mpatches.Patch(facecolor=FC[i], alpha=0.4, edgecolor=FC[i],
              label=f'F{i+1} = {formants[i]} Hz') for i, _ in valid]
    if fp:
        le.append(Line2D([0],[0],marker='D',color='w',markerfacecolor='#1B5E20',
                         markeredgecolor='black',markersize=5,label=f'Fp centroïde = {fp} Hz'))
    ax.legend(handles=le, loc='lower left', fontsize=6.5, framealpha=0.92, edgecolor='#CCC')
    ax.text(0.99, -0.06,
            f"Famille : {family_label or 'Orchestre'} · Source : CSV v3 + specenv Orchidea (SOL2020 + Yan_Adds)",
            transform=ax.transAxes, fontsize=7, color='#888', ha='right')
    plt.tight_layout()

    out = os.path.join(OUT_IMG, f"{filename}.png")
    fig.savefig(out, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return out


# ─── Calcul Fp depuis specenv brut ───────────────────────────
SAMPLE_RATE = 44100
FFT_SIZE    = 4096
FREQ_RES    = SAMPLE_RATE / FFT_SIZE   # ~10.77 Hz/bin

# Bandes de centroïde par instrument (même que build2v2_envelopes)
FP_BANDS_SPECENV = {
    'Horn':                  (600,  1400),
    'Horn+sordina':          (600,  1400),
    'Trumpet_C':             (600,  1400),
    'Trumpet_C+sordina_cup':     (600, 1400),
    'Trumpet_C+sordina_straight':(600, 1400),
    'Trumpet_C+sordina_harmon':  (800, 2000),
    'Trumpet_C+sordina_wah':     (600, 1400),
    'Trombone':              (1000, 2000),
    'Trombone+sordina_cup':      (1000, 2000),
    'Trombone+sordina_straight': (1000, 2000),
    'Trombone+sordina_harmon':   (800,  2000),
    'Trombone+sordina_wah':      (1000, 2000),
    'Bass_Tuba':             (1000, 2000),
    'Bass_Tuba+sordina':     (1000, 2000),
    'Bass_Trombone':         (1000, 2000),
    'Contrabass_Tuba':       (1000, 2000),
}

SOL_DIR = os.path.join(BASE, 'Data', 'FullSOL2020_specenv par instrument')
YAN_DIR = os.path.join(BASE, 'Data', 'Yan_Adds-Divers_specenv par instrument')

def _band_centroid(env_vals, lo_hz, hi_hz):
    """Centroïde spectral pondéré en énergie dans une bande Hz."""
    lo_bin = max(0, int(lo_hz / FREQ_RES))
    hi_bin = min(len(env_vals) - 1, int(hi_hz / FREQ_RES))
    region = env_vals[lo_bin:hi_bin + 1]
    if not region:
        return None
    linear = [10 ** (v / 10) for v in region]
    total  = sum(linear)
    if total <= 0:
        return None
    return sum(linear[i] * (lo_bin + i) * FREQ_RES for i in range(len(linear))) / total

def compute_fp_from_specenv(instrument_key, techs=('ordinario',),
                             specenv_dir=None, specenv_filename=None):
    """
    Calcule le Fp (centroïde spectral) directement depuis le fichier specenv brut.
    
    instrument_key  : ex. 'Trombone+sordina_cup', 'Horn+sordina'
    techs           : tuple de techniques à inclure, ex. ('ordinario',) ou ('ordinario_open',)
    specenv_dir     : répertoire du fichier (défaut : SOL_DIR)
    specenv_filename: nom du fichier (défaut : auto-construit depuis instrument_key)
    
    Retourne le Fp médian arrondi à l'entier, ou None si pas de données.
    """
    if specenv_dir is None:
        specenv_dir = SOL_DIR
    if specenv_filename is None:
        specenv_filename = f"FullSOL2020_specenv.db_{instrument_key}.txt"

    filepath = os.path.join(specenv_dir, specenv_filename)
    if not os.path.exists(filepath):
        # Fallback: try Yan_Adds directory with its naming pattern
        # Map SOL names to Yan_Adds names (underscore → hyphen, + → hyphen)
        ya_key = instrument_key.replace('_', '-')
        candidates = [
            f"Yan_Adds-Divers_specenv.db_{ya_key}.txt",
            f"Yan_Adds-Divers_specenv.db_{ya_key.replace('+', '-')}.txt",
            f"Yan_Adds-Divers_specenv.db_{instrument_key}.txt",
        ]
        found = False
        for cand in candidates:
            ya_filepath = os.path.join(YAN_DIR, cand)
            if os.path.exists(ya_filepath):
                filepath = ya_filepath
                found = True
                break
        if not found:
            return None

    # Bande Fp pour cet instrument
    band = FP_BANDS_SPECENV.get(instrument_key, (800, 1800))
    lo_hz, hi_hz = band

    fps = []
    try:
        with open(filepath, encoding='utf-8') as f:
            f.readline()  # en-tête
            for line in f:
                parts = line.strip().split(';')
                if len(parts) < 10:
                    continue
                path = parts[0]
                if not path.startswith('/'):
                    continue
                # Filtrer par technique (champ [3] du chemin /Famille/Instrument/technique/...)
                path_parts = path.split('/')
                if len(path_parts) < 5:
                    continue
                tech = path_parts[3]
                if tech not in techs:
                    continue
                try:
                    vals = [float(v) for v in parts[1:]]
                except ValueError:
                    continue
                if len(vals) < 100:
                    continue
                # Les valeurs specenv sont en dB (négatives) — on vérifie juste
                # que le tableau n'est pas vide et non constant
                max_amp = max(vals)
                min_amp = min(vals)
                if max_amp == min_amp:
                    continue   # spectre plat = données invalides
                fp = _band_centroid(vals, lo_hz, hi_hz)
                if fp:
                    fps.append(fp)
    except Exception:
        return None

    if not fps:
        return None

    fps.sort()
    median = fps[len(fps) // 2]
    return round(median)

# ─── Tableaux techniques HTML ────────────────────────────────
SUSTAINED_TECHS = {
    'ordinario', 'non-vibrato', 'non_vibrato', 'vibrato',
    'ordinario_open', 'ordinario_closed',
    'flautando', 'arco',
}
EXCLUDED_TECHS = {
    'pizzicato_secco', 'pizzicato_bartok', 'pizzicato_l_vib', 'pizzicato-l-vib',
    'col_legno_battuto', 'col_legno_tratto', 'tremolo', 'artificial_harmonic',
    'artificial_harmonic_tremolo', 'behind_the_bridge', 'behind_the_fingerboard',
    'on_the_bridge', 'on_the_tailpiece', 'on_the_tuning_pegs', 'on_the_frog',
    'hit_on_body', 'col_legno', 'natural_harmonics_glissandi',
}

def tech_table_html(inst_csv, filter_sustained=False):
    techs = sorted([(t, r) for (i, t), r in DATA.items() if i == inst_csv],
                   key=lambda x: x[0])
    if not techs:
        return ""
    rows = []
    for tech, r in techs:
        if filter_sustained and tech in EXCLUDED_TECHS:
            continue
        n = r['n_samples']
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        is_ord = any(s in tech.lower() for s in ('ordinario', 'non-vibrato', 'non_vibrato'))
        bg = ' style="background-color:#dff0d8;"' if is_ord else ''
        fvals = ''.join(f'<td{bg}>{fmt_hz(f)}</td>' for f in fs)
        rows.append(f'<tr><td{bg}><b>{tech}</b></td><td{bg}>{n}</td>{fvals}</tr>')
        # ── v5 : ligne stats σ / IQR pour les techniques ordinario ──
        if is_ord and r.get('F1_std') and sf(r.get('F1_std', '')) > 0:
            stats_cells = []
            for i in range(1, 4):  # F1–F3 only
                std_v = sf(r.get(f'F{i}_std', ''))
                q25_v = sf(r.get(f'F{i}_q25', ''))
                q75_v = sf(r.get(f'F{i}_q75', ''))
                if std_v > 0:
                    iqr = round(q75_v - q25_v)
                    stats_cells.append(
                        f'<td{bg} style="font-size:0.78em;color:#555;">'
                        f'σ={round(std_v)}<br>[{round(q25_v)}–{round(q75_v)}]<br>'
                        f'IQR={iqr}</td>')
                else:
                    stats_cells.append(f'<td{bg}>—</td>')
            # F4–F6 just σ
            for i in range(4, 7):
                std_v = sf(r.get(f'F{i}_std', ''))
                if std_v > 0:
                    stats_cells.append(
                        f'<td{bg} style="font-size:0.78em;color:#555;">σ={round(std_v)}</td>')
                else:
                    stats_cells.append(f'<td{bg}>—</td>')
            rows.append(
                f'<tr><td{bg} style="font-size:0.78em;color:#555;font-style:italic;">'
                f'↳ σ / [Q25–Q75]</td><td{bg}></td>'
                + ''.join(stats_cells) + '</tr>')
    if not rows:
        return ""
    return (f'<table class="tech-table">'
            f'<tr class="header"><th>Technique</th><th>N</th>'
            f'<th>F1</th><th>F2</th><th>F3</th><th>F4</th><th>F5</th><th>F6</th></tr>'
            + ''.join(rows) + '</table>')

# ─── Tableaux références sources HTML ───────────────────────
def ref_table_html(rows_data):
    """
    rows_data : liste de dicts avec clés :
      source, f1, f2, f3, f4, voyelle, n, accord, note
    """
    html = (
        '<table class="ref-table">'
        '<tr class="header">'
        '<th>Source</th><th>F1 (Hz)</th><th>F2 (Hz)</th>'
        '<th>F3 (Hz)</th><th>F4 (Hz)</th>'
        '<th>Voyelle</th><th>N</th><th>Accord</th>'
        '</tr>'
    )
    for r in rows_data:
        src = r.get('source', '—')
        is_sol = 'SOL' in src or 'Yan' in src
        bg = ' style="background-color:#dff0d8;"' if is_sol else ''
        accord = r.get('accord', '')
        if accord == 'oui':
            accord = '✓'
        elif accord == 'non':
            accord = '✗'
        elif accord == 'approx':
            accord = '~'
        html += (
            f'<tr>'
            f'<td{bg}><b>{src}</b></td>'
            f'<td{bg}>{r.get("f1","—")}</td>'
            f'<td{bg}>{r.get("f2","—")}</td>'
            f'<td{bg}>{r.get("f3","—")}</td>'
            f'<td{bg}>{r.get("f4","—")}</td>'
            f'<td{bg}>{r.get("voyelle","—")}</td>'
            f'<td{bg}>{r.get("n","—")}</td>'
            f'<td{bg}>{accord}</td>'
            f'</tr>'
        )
    html += '</table>'
    return html

# ─── Section doublures HTML ──────────────────────────────────
def doublures_html(items):
    """
    items : liste de dicts avec clés :
      instr, f1_a, f1_b, delta, quality, rapport, note
      rapport : 'Unisson' | 'Octave' | '2 octaves' | 'Complémentaire' | etc.
    """
    if not items:
        return ""
    rows = ""
    for it in items:
        delta = it.get('delta', '')
        if isinstance(delta, (int, float)):
            delta_str = f"Δ={delta} Hz"
        else:
            delta_str = str(delta)
        quality = it.get('quality', '')
        rapport = it.get('rapport', '—')
        # Colorier le rapport
        rapport_color = {
            'Unisson':       '#1B5E20',
            'Octave':        '#1565C0',
            '2 octaves':     '#4A148C',
            'Complémentaire':'#E65100',
        }.get(rapport, '#555')
        rapport_html = f'<span style="color:{rapport_color};font-weight:bold;">{rapport}</span>'
        rows += (
            f'<tr>'
            f'<td><b>{it.get("instr","")}</b></td>'
            f'<td>{it.get("f1_a","—")}</td>'
            f'<td>{it.get("f1_b","—")}</td>'
            f'<td>{delta_str}</td>'
            f'<td>{quality}</td>'
            f'<td>{rapport_html}</td>'
            f'<td>{it.get("note","")}</td>'
            f'</tr>'
        )
    return (
        '<div class="doublures-box">'
        '<h4>Doublures recommandées</h4>'
        '<table class="doublures-table">'
        '<tr class="header">'
        '<th>Instrument associé</th><th>F1 (instrument)</th>'
        '<th>F1 (associé)</th><th>Écart</th><th>Qualité</th>'
        '<th>Rapport de tessiture</th><th>Commentaire</th>'
        '</tr>'
        + rows +
        '</table></div>'
    )

# ─── CSS global ──────────────────────────────────────────────
CSS_GLOBAL = """
body {
  font-family: 'Segoe UI', Helvetica, Arial, sans-serif;
  max-width: 1100px; margin: 0 auto; padding: 24px;
  background: #f5f5f5; color: #333; line-height: 1.6;
}
h1 { color: #1a237e; border-bottom: 3px solid #2e7d32; padding-bottom: 10px; font-size: 1.9em; }
h2 { color: #2e7d32; margin-top: 50px; border-left: 5px solid #2e7d32;
     padding-left: 14px; font-size: 1.4em; }
h3 { color: #1b5e20; margin-top: 36px; font-size: 1.2em; }
h4 { color: #555; margin-top: 16px; font-size: 1.05em; }
.instrument-card {
  background: white; border: 1px solid #ddd; border-radius: 8px;
  padding: 22px; margin: 24px 0; box-shadow: 0 2px 6px rgba(0,0,0,0.08);
}
.formant-graph { max-width: 65%; border: 1px solid #eee; border-radius: 4px; display: block; }
.description {
  font-style: italic; color: #555; background: #e8f5e9;
  padding: 12px 16px; border-left: 4px solid #4caf50; margin: 12px 0; border-radius: 0 4px 4px 0;
}
.fp-note { color: #1b5e20; font-weight: bold; margin: 8px 0; }
.fp-explanation {
  background: #e8f5e9; border: 1px solid #a5d6a7; border-radius: 6px;
  padding: 14px 18px; margin: 16px 0; font-size: 0.95em;
}
.note-v4 {
  background: #fff3e0; border-left: 4px solid #ff9800;
  padding: 10px 14px; margin: 12px 0; font-size: 0.9em; color: #555;
  border-radius: 0 4px 4px 0;
}
.cluster-badge {
  display: inline-block; background: #e53935; color: white;
  padding: 2px 9px; border-radius: 10px; font-size: 0.8em;
  margin-left: 8px; font-weight: bold;
}
.yan-badge {
  display: inline-block; background: #ff9800; color: white;
  padding: 2px 8px; border-radius: 10px; font-size: 0.8em; margin-left: 8px;
}
/* Tables techniques */
.tech-table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 0.88em; }
.tech-table th, .tech-table td { border: 1px solid #ccc; padding: 5px 9px; text-align: center; }
.tech-table .header th { background: #1a3a5c; color: white; font-size: 0.9em; }
.tech-table tr:nth-child(even) { background: #fafafa; }
/* Tables références */
.ref-table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 0.88em; }
.ref-table th, .ref-table td { border: 1px solid #bbb; padding: 6px 10px; text-align: center; }
.ref-table .header th { background: #37474f; color: white; }
.ref-table tr:hover { background: #f0f4f8; }
/* Doublures */
.doublures-box {
  background: #fffde7; border: 1px solid #f9a825; border-radius: 6px;
  padding: 14px 18px; margin: 18px 0;
}
.doublures-box h4 { color: #f57f17; margin: 0 0 10px 0; }
.doublures-table { width: 100%; border-collapse: collapse; font-size: 0.88em; }
.doublures-table th, .doublures-table td { border: 1px solid #f9a825; padding: 5px 9px; text-align: center; }
.doublures-table .header th { background: #f57f17; color: white; }
/* Section intro */
.section-intro {
  padding: 16px 20px; border-radius: 6px; margin: 16px 0;
}
.section-intro.bois   { background: #e8f5e9; border-left: 5px solid #2e7d32; }
.section-intro.cuivres{ background: #fff3e0; border-left: 5px solid #e64a19; }
.section-intro.cordes { background: #e3f2fd; border-left: 5px solid #1565c0; }
.section-intro.sax    { background: #fce4ec; border-left: 5px solid #ad1457; }
.section-intro.general{ background: #f3e5f5; border-left: 5px solid #6a1b9a; }
.section-intro.intro  { background: #e8eaf6; border-left: 5px solid #283593; }
/* Source note */
.source-note {
  font-size: 0.83em; color: #888; margin-top: 40px;
  border-top: 1px solid #ddd; padding-top: 12px;
}
/* Document intro */
.doc-header {
  background: linear-gradient(135deg, #1a237e 0%, #283593 60%, #1565c0 100%);
  color: white; padding: 32px 36px; border-radius: 10px; margin-bottom: 32px;
}
.doc-header h1 { color: white; border: none; margin: 0 0 12px 0; font-size: 2em; }
.doc-header .subtitle { color: #b3c5ff; font-size: 1.05em; margin: 4px 0; }
.doc-header .meta { color: #90a4c8; font-size: 0.88em; margin-top: 16px; }
/* Tableau voyelles */
.vowel-table { width: 100%; border-collapse: collapse; margin: 16px 0; font-size: 0.9em; }
.vowel-table th, .vowel-table td { border: 1px solid #ccc; padding: 7px 12px; text-align: center; }
.vowel-table .header th { background: #4a148c; color: white; }
/* Convergence / matrice */
.matrix-container { overflow-x: auto; margin: 16px 0; }
.matrix-table { border-collapse: collapse; font-size: 0.72em; white-space: nowrap; }
.matrix-table th, .matrix-table td { border: 1px solid #ccc; padding: 4px 7px; text-align: center; }
.matrix-table th { background: #263238; color: white; }
"""

# ─── Utilitaires DOCX ────────────────────────────────────────
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_text(cell, text, bold=False, size=9, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)

def tech_table_docx(doc, inst_csv):
    techs = sorted([(t, r) for (i, t), r in DATA.items() if i == inst_csv],
                   key=lambda x: x[0])
    if not techs:
        return
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Technique', 'N', 'F1', 'F2', 'F3', 'F4', 'F5', 'F6']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], '1A3A5C')
    for tech, r in techs:
        row = table.add_row().cells
        fs = [round(sf(r[f'F{i}_hz'])) for i in range(1, 7)]
        vals = [tech, r['n_samples']] + [fmt_hz(f) for f in fs]
        is_ord = any(s in tech.lower() for s in ('ordinario', 'non-vibrato', 'non_vibrato'))
        fill = 'DFF0D8' if is_ord else None
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
            if fill:
                set_cell_shading(row[idx], fill)
        # ── v5 : ligne stats σ / IQR pour ordinario ──
        if is_ord and r.get('F1_std') and sf(r.get('F1_std', '')) > 0:
            srow = table.add_row().cells
            set_cell_text(srow[0], '↳ σ / [Q25–Q75]', size=7, color=(120, 120, 120))
            set_cell_text(srow[1], '', size=7)
            if fill:
                set_cell_shading(srow[0], fill)
                set_cell_shading(srow[1], fill)
            for i in range(1, 7):
                ci = i + 1  # column index
                std_v = sf(r.get(f'F{i}_std', ''))
                q25_v = sf(r.get(f'F{i}_q25', ''))
                q75_v = sf(r.get(f'F{i}_q75', ''))
                if std_v > 0 and i <= 3:
                    iqr = round(q75_v - q25_v)
                    txt = f"σ={round(std_v)} [{round(q25_v)}–{round(q75_v)}]"
                elif std_v > 0:
                    txt = f"σ={round(std_v)}"
                else:
                    txt = '—'
                set_cell_text(srow[ci], txt, size=7, color=(120, 120, 120))
                if fill:
                    set_cell_shading(srow[ci], fill)
    widths_cm = [4.8, 1.1, 1.3, 1.3, 1.3, 1.3, 1.3, 1.3]
    for row_obj in table.rows:
        for cell, width in zip(row_obj.cells, widths_cm):
            cell.width = Cm(width)

def ref_table_docx(doc, rows_data):
    if not rows_data:
        return
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Source', 'F1 (Hz)', 'F2 (Hz)', 'F3 (Hz)', 'F4 (Hz)', 'Voyelle', 'N', 'Accord']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], '37474F')
    for r in rows_data:
        row = table.add_row().cells
        is_sol = 'SOL' in r.get('source', '') or 'Yan' in r.get('source', '')
        fill = 'DFF0D8' if is_sol else None
        vals = [r.get('source','—'), r.get('f1','—'), r.get('f2','—'),
                r.get('f3','—'), r.get('f4','—'), r.get('voyelle','—'),
                r.get('n','—'), r.get('accord','—')]
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
            if fill:
                set_cell_shading(row[idx], fill)
    widths_cm = [3.0, 1.8, 1.8, 1.8, 1.8, 1.8, 1.0, 1.2]
    for row_obj in table.rows:
        for cell, width in zip(row_obj.cells, widths_cm):
            cell.width = Cm(width)

def clean_text(text):
    """
    Nettoie un texte pour le DOCX :
    - Supprime l'indentation Python des triples guillemets
    - Supprime les balises HTML simples (<strong>, <em>, <br>, etc.)
    - Retourne un texte propre sur une seule ligne cohérente
    """
    import re
    # Supprimer les balises HTML courantes
    text = re.sub(r'<strong>(.*?)</strong>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<em>(.*?)</em>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<b>(.*?)</b>', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'<br\s*/?>', ' ', text)
    text = re.sub(r'<[^>]+>', '', text)
    # Normaliser : strip chaque ligne puis rejoindre
    lines = [l.strip() for l in text.split('\n')]
    result = ' '.join(l for l in lines if l)
    # Nettoyer les espaces multiples
    result = re.sub(r'  +', ' ', result)
    return result.strip()

def doublures_table_docx(doc, dbl_items, header_color='F57F17'):
    """
    Rend un tableau de doublures en DOCX avec colonne Rapport de tessiture.
    dbl_items : même format que doublures_html (clé 'rapport' optionnelle).
    """
    if not dbl_items:
        return
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    headers = ['Instrument associé', 'F1', 'F1 associé', 'Écart', 'Qualité', 'Rapport tessiture']
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell_text(hdr[idx], h, bold=True, size=9, color=(255, 255, 255))
        set_cell_shading(hdr[idx], header_color)
    for it in dbl_items:
        row = table.add_row().cells
        delta = it.get('delta', '')
        delta_str = f"Δ={delta} Hz" if isinstance(delta, (int, float)) else str(delta)
        rapport = it.get('rapport', '—')
        vals = [it.get('instr', ''), it.get('f1_a', '—'), it.get('f1_b', '—'),
                delta_str, it.get('quality', ''), rapport]
        for idx, v in enumerate(vals):
            set_cell_text(row[idx], v, bold=(idx == 0), size=9)
    for row_obj in table.rows:
        for cell, w in zip(row_obj.cells, [3.2, 1.6, 1.6, 1.8, 2.5, 2.5]):
            cell.width = Cm(w)

def add_heading(doc, text, level=1, color=None):
    style_name = f'Heading {level}' if level <= 4 else 'Normal'
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def _make_bookmark_id():
    """Génère un ID numérique unique pour chaque bookmark."""
    if not hasattr(_make_bookmark_id, '_counter'):
        _make_bookmark_id._counter = 1
    bid = _make_bookmark_id._counter
    _make_bookmark_id._counter += 1
    return bid


def add_bookmark(paragraph, bookmark_name):
    """
    Ajoute un signet (bookmark) Word sur un paragraphe existant.
    Le signet entoure le run du paragraphe.
    bookmark_name doit être unique dans le document et sans espaces.
    """
    bid = str(_make_bookmark_id())
    # bookmarkStart
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), bid)
    bm_start.set(qn('w:name'), bookmark_name)
    paragraph._p.insert(0, bm_start)
    # bookmarkEnd
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), bid)
    paragraph._p.append(bm_end)


def add_heading_bookmarked(doc, text, level=1, color=None, bookmark=None):
    """
    Ajoute un titre avec un signet Word pour permettre les hyperliens TDM.
    bookmark : nom du signet (str, sans espaces). Si None, généré depuis text.
    """
    import unicodedata, re as _re
    p = add_heading(doc, text, level=level, color=color)
    if bookmark is None:
        # Générer un nom de signet depuis le texte
        bm = unicodedata.normalize('NFD', text)
        bm = ''.join(c for c in bm if unicodedata.category(c) != 'Mn')
        bm = _re.sub(r'[^a-zA-Z0-9]+', '_', bm).strip('_')
        bookmark = bm[:40]
    add_bookmark(p, bookmark)
    return p, bookmark


def add_toc_hyperlink(paragraph, text, bookmark, size=10, bold=False,
                      color=None):
    """
    Ajoute un hyperlien interne (vers un bookmark) dans un paragraphe.
    """
    # Relation hyperlien interne
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark)

    run_elem = OxmlElement('w:r')

    # Style du lien (souligné, couleur)
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    if color:
        clr = OxmlElement('w:color')
        if isinstance(color, RGBColor):
            clr.set(qn('w:val'), f'{color.red:02X}{color.green:02X}{color.blue:02X}')
        else:
            clr.set(qn('w:val'), color)
        rPr.append(clr)

    run_elem.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)

def add_paragraph(doc, text, italic=False, bold=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def new_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    return doc

# ─── HTML helpers ────────────────────────────────────────────
def html_head(title, extra_css=""):
    return f"""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
{CSS_GLOBAL}
{extra_css}
</style>
</head>
<body>
"""

def html_foot():
    return """</body>
</html>
"""



# ─── Carte spectrale vocalique + Analyse par registre ────────
import re as _re

def _note_to_midi(s):
    m = _re.match(r'([A-G])(#?)(\d+)', s)
    if not m: return None
    base = {'C':0,'D':2,'E':4,'F':5,'G':7,'A':9,'B':11}[m.group(1)]
    return (int(m.group(3))+1)*12 + base + (1 if m.group(2)=='#' else 0)

def cepstral_envelope(spectrum_linear, order=30):
    """Smoothed spectral envelope via cepstral truncation. Returns log-scale array."""
    n = len(spectrum_linear)
    log_spec = np.log(np.maximum(spectrum_linear, 1e-10))
    sym = np.concatenate([log_spec, log_spec[-2:0:-1]])
    cepstrum = np.real(np.fft.ifft(sym))
    liftered = np.zeros_like(cepstrum)
    liftered[:order+1] = cepstrum[:order+1]
    liftered[-(order):] = cepstrum[-(order):]
    return np.real(np.fft.fft(liftered))[:n]


# ─── Registres par instrument (source : registres.md) ────────
# Notes frontières incluses dans les DEUX registres adjacents.
REGISTERS = {
    'Flute': [
        ('Grave', 59, 69),       # B3-A4
        ('Médium', 69, 81),      # A4-A5 (A4 shared with Grave)
        ('Aigu', 81, 93),        # A5-A6 (A5 shared with Médium)
        ('Suraigu', 94, 999),
    ],
    'Oboe': [
        ('Grave', 58, 67),       # A#3-G4
        ('Médium', 67, 79),      # G4-G5 (G4 shared)
        ('Aigu', 79, 86),        # G5-D6 (G5 shared)
        ('Suraigu', 87, 999),
    ],
    'Clarinet_Bb': [
        ('Chalumeau', 50, 62),   # D3-D4
        ('Gorge', 63, 68),       # D#4-G#4
        ('Clairon', 69, 82),     # A4-A#5
        ('Suraigu', 83, 999),    # B5+
    ],
    'Bass_Clarinet_Bb': [
        ('Grave', 34, 50),       # A#1-D3
        ('Médium', 51, 56),      # D#3-G#3
        ('Aigu', 57, 70),        # A3-A#4
        ('Suraigu', 71, 999),    # B4+
    ],
    'Bassoon': [
        ('Grave', 34, 45),       # A#1-A2
        ('Bas médium', 46, 57),  # A#2-A3
        ('Haut médium', 58, 69), # A#3-A4
        ('Aigu', 70, 999),       # A#4+
    ],
    'Horn': [
        ('Pédale', 29, 33),      # F1-A1
        ('Grave', 34, 47),       # A#1-B2
        ('Médium', 48, 64),      # C3-E4
        ('Aigu', 65, 77),        # F4-F5
        ('Suraigu', 78, 999),    # F#5+
    ],
    'Trumpet_C': [
        ('Grave', 54, 59),       # F#3-B3
        ('Médium', 60, 67),      # C4-G4
        ('Aigu', 68, 84),        # G#4-C6
        ('Suraigu', 85, 999),
    ],
    'Trombone': [
        ('Pédale', 28, 34),      # E1-A#1 (A#1 shared with Grave)
        ('Grave', 34, 52),       # A#1-E3 (E3 shared with Médium)
        ('Médium', 52, 64),      # E3-E4
        ('Aigu', 65, 999),       # F4+
    ],
    'Bass_Tuba': [
        ('Grave', 0, 41),        # up to F2 (F2 shared with Médium)
        ('Médium', 41, 53),      # F2-F3 (F3 shared with Aigu)
        ('Aigu', 53, 62),        # F3-D4
        ('Suraigu', 63, 999),    # D#4+
    ],
    'Violin': [
        ('Grave', 55, 60),       # G3-C4 (C4 shared with Médium)
        ('Médium', 60, 72),      # C4-C5 (C5 shared with Aigu)
        ('Aigu', 72, 84),        # C5-C6 (C6 shared with Suraigu)
        ('Suraigu', 84, 999),    # C6+
    ],
    'Viola': [
        ('Grave', 48, 55),       # C3-G3 (G3 shared with Médium)
        ('Médium', 55, 67),      # G3-G4 (G4 shared with Aigu)
        ('Aigu', 67, 79),        # G4-G5 (G5 shared with Suraigu)
        ('Suraigu', 79, 999),    # G5+
    ],
    'Violoncello': [
        ('Grave', 36, 43),       # C2-G2 (G2 shared with Médium)
        ('Médium', 43, 55),      # G2-G3 (G3 shared with Aigu)
        ('Aigu', 55, 67),        # G3-G4 (G4 shared with Suraigu)
        ('Suraigu', 67, 999),    # G4+
    ],
    'Contrabass': [
        ('Grave', 24, 36),       # C1-C2 (C2 shared with Médium)
        ('Médium', 36, 48),      # C2-C3 (C3 shared with Aigu)
        ('Aigu', 48, 60),        # C3-C4 (C4 shared with Suraigu)
        ('Suraigu', 60, 999),    # C4+
    ],
}
REGISTERS['Violin_Ensemble'] = REGISTERS['Violin']
REGISTERS['Viola_Ensemble'] = REGISTERS['Viola']
REGISTERS['Violoncello_Ensemble'] = REGISTERS['Violoncello']
REGISTERS['Contrabass_Ensemble'] = REGISTERS['Contrabass']


def _get_registers(instrument_key, midi):
    """Return list of register names for a given MIDI note (frontier notes appear in multiple)."""
    regs = REGISTERS.get(instrument_key)
    if not regs:
        return []
    return [name for name, lo, hi in regs if lo <= midi <= hi]


def _find_specenv_file(instrument_key):
    filepath = os.path.join(SOL_DIR, f"FullSOL2020_specenv.db_{instrument_key}.txt")
    if os.path.exists(filepath):
        return filepath
    ya_key = instrument_key.replace('_', '-')
    for cand in [
        f"Yan_Adds-Divers_specenv.db_{ya_key}.txt",
        f"Yan_Adds-Divers_specenv.db_{ya_key.replace('+', '-')}.txt",
        f"Yan_Adds-Divers_specenv.db_{instrument_key}.txt",
    ]:
        p = os.path.join(YAN_DIR, cand)
        if os.path.exists(p):
            return p
    return None


def _find_spectrum_file(instrument_key):
    filepath = os.path.join(BASE, 'Data', 'FullSOL2020.spectrum_par_instrument',
                            f'{instrument_key}_spectrum.txt')
    if os.path.exists(filepath):
        return filepath
    ya_key = instrument_key.replace('_', '-')
    for cand in [f'{ya_key}_spectrum.txt', f'{instrument_key}_spectrum.txt']:
        p = os.path.join(BASE, 'Data', 'Yan_Adds-Divers.spectrum_par_instrument', cand)
        if os.path.exists(p):
            return p
    return None


def _load_grouped(filepath, techs, instrument_key):
    """Load data grouped by register. Frontier notes go into ALL matching registers."""
    groups = {}
    with open(filepath, encoding='utf-8') as f:
        f.readline()
        for line in f:
            parts = line.strip().split(';')
            if len(parts) < 10: continue
            path = parts[0]
            if not path.startswith('/'): continue
            pp = path.split('/')
            if len(pp) < 5 or pp[3] not in techs: continue
            try: vals = [float(v) for v in parts[1:]]
            except: continue
            if len(vals) < 100: continue
            fname = path.split('/')[-1]
            note_parts = fname.replace('.wav','').split('-')
            note = note_parts[2] if len(note_parts) >= 3 else None
            if not note: continue
            midi = _note_to_midi(note)
            if midi is None: continue
            regs = _get_registers(instrument_key, midi)
            for reg in regs:
                groups.setdefault(reg, {'data': [], 'notes': set()})
                groups[reg]['data'].append(np.array(vals))
                groups[reg]['notes'].add(note)
    return groups


def load_specenv_by_register(instrument_key, techs=('ordinario',)):
    filepath = _find_specenv_file(instrument_key)
    if not filepath:
        return {}
    return _load_grouped(filepath, techs, instrument_key)


def load_spectrum_by_register(instrument_key, techs=('ordinario',)):
    filepath = _find_spectrum_file(instrument_key)
    if not filepath:
        return {}
    return _load_grouped(filepath, techs, instrument_key)


def make_carte_spectrale(display, filename, mean_env, n_samples, fp_band=(800,1800),
                         family_color='#333', cep_env_db=None, note_range=''):
    """
    Generate a 'Carte spectrale vocalique' image with pixel-based anti-collision
    that checks both label-label AND label-curve overlaps.
    """
    _freqs = np.arange(len(mean_env)) * FREQ_RES
    mf = 5500
    mask = (_freqs >= 80) & (_freqs <= mf)

    # Find peaks
    lo = max(int(80/FREQ_RES),1); hi = min(int(6000/FREQ_RES), len(mean_env)-1)
    reg = list(mean_env[lo:hi]); mx = max(reg); th = mx - 25
    peaks = []
    for i in range(1, len(reg)-1):
        v = reg[i]
        if v < th: continue
        if v >= reg[i-1] and v >= reg[i+1]: peaks.append(((lo+i)*FREQ_RES, v))
    peaks.sort(key=lambda x: x[1], reverse=True)
    sel = []
    for f, a in peaks:
        if not any(abs(f-sf)<70 for sf,_ in sel): sel.append((f,a))
        if len(sel) >= 6: break
    sel.sort(key=lambda x: x[0])

    # Fp
    lo_b = int(fp_band[0]/FREQ_RES); hi_b = int(fp_band[1]/FREQ_RES)
    region = mean_env[lo_b:hi_b+1]
    linear = np.array([10**(v/10) for v in region])
    total = linear.sum()
    fp = sum(linear[i]*(lo_b+i)*FREQ_RES for i in range(len(linear)))/total if total > 0 else None
    fp_bin = min(int(fp/FREQ_RES), len(mean_env)-1) if fp else None

    y_min = mean_env[mask].min() - 6
    y_max = mean_env[mask].max() + 18
    VOWEL_Y = y_max - 2

    # Build labels
    labels = []
    for i, (f, a) in enumerate(sel):
        bin_idx = min(int(f/FREQ_RES), len(mean_env)-1)
        db_val = mean_env[bin_idx]
        labels.append({'x':f, 'y':db_val, 'text':f"F{i+1} {f:.0f} Hz ({db_val:.0f} dB)",
                       'priority':6-i, 'type':'formant'})
    if fp:
        labels.append({'x':fp, 'y':mean_env[fp_bin],
                       'text':f"Fp {fp:.0f} Hz (centroide)",
                       'priority':7, 'type':'fp'})
    labels.sort(key=lambda l: -l['priority'])

    # ── Draw figure ──
    fig, ax = plt.subplots(figsize=(9, 5), dpi=150)

    for lo_hz, hi_hz, color, label in VOWEL_ZONES:
        if lo_hz < mf:
            ax.axvspan(lo_hz, min(hi_hz, mf), alpha=0.20, color=color, zorder=0)
            mid = (lo_hz + min(hi_hz, mf)) / 2
            if mid < mf*0.95:
                ax.text(mid, VOWEL_Y, label, ha='center', va='top', fontsize=7,
                        color='#666', fontweight='bold')

    ax.axvspan(420, 550, alpha=0.12, color='red', zorder=1)
    ax.text(485, y_min+1, 'cluster /o/', ha='center', fontsize=6,
            color='#C62828', fontstyle='italic')

    ax.plot(_freqs[mask], mean_env[mask], color=family_color, lw=2, zorder=3,
            label=f'Enveloppe specenv (n={n_samples})')

    if cep_env_db is not None:
        ax.plot(_freqs[mask], cep_env_db[mask], '--', color='#9C27B0', lw=0.9, alpha=0.5,
                zorder=2, label='Cepstrale (ord.30)')

    for i, (f, a) in enumerate(sel):
        bin_idx = min(int(f/FREQ_RES), len(mean_env)-1)
        ax.plot(f, mean_env[bin_idx], 'v', color='#D32F2F', markersize=5, zorder=5)
    if fp:
        ax.plot(fp, mean_env[fp_bin], 'D', color='#1B5E20', markersize=5,
                markeredgecolor='black', markeredgewidth=0.8, zorder=6)

    # Set up axes BEFORE label placement (needed for transforms)
    ax.set_xscale('log')
    ticks = [t for t in [100,150,200,300,400,500,600,800,1000,1500,2000,3000,4000,5000] if t<=mf]
    ax.set_xticks(ticks)
    ax.set_xticklabels([str(t) for t in ticks], fontsize=7)
    ax.get_xaxis().set_major_formatter(matplotlib.ticker.ScalarFormatter())
    ax.set_xlim(80, mf); ax.set_ylim(y_min, y_max)
    fig.canvas.draw()

    trans = ax.transData
    inv_trans = ax.transData.inverted()

    # Pre-compute curve in pixel space for collision checking
    curve_freqs = _freqs[mask]
    curve_vals = mean_env[mask]
    curve_px = np.array([trans.transform((f, v)) for f, v in zip(curve_freqs, curve_vals)])

    placed_px = []  # (px_left, px_right, py_bottom, py_top)
    PX_PAD = 8

    def _collides(px_l, px_r, py_b, py_t):
        # Check label-label collision
        for ol, orr, ob, ot in placed_px:
            if px_l < orr + PX_PAD and px_r > ol - PX_PAD:
                if py_b < ot + PX_PAD and py_t > ob - PX_PAD:
                    return True
        # Check label-curve collision (does the curve cross through the label box?)
        curve_in_x = curve_px[(curve_px[:,0] >= px_l - 4) & (curve_px[:,0] <= px_r + 4)]
        if len(curve_in_x) > 0:
            # curve_px y is in display coords (y=0 at top)
            curve_y_min = curve_in_x[:,1].min()
            curve_y_max = curve_in_x[:,1].max()
            if py_b < curve_y_max + 6 and py_t > curve_y_min - 6:
                return True
        return False

    # Place labels
    for lab in labels:
        is_fp = lab['type'] == 'fp'
        clr = '#1B5E20' if is_fp else '#333'
        txt_w = len(lab['text']) * 5.5
        txt_h = 10

        anchor_px = trans.transform((lab['x'], lab['y']))
        ax_x, ax_y = anchor_px

        best = None
        best_cost = 999
        # In display coords, y DECREASES going up
        offsets_y_above = [18, 32, 46, 60, 76, 92, 110, 130]
        offsets_y_below = [20, 34, 48, 62, 78, 94]
        offsets_x = [0, -55, 55, -110, 110, -165, 165]

        for ox in offsets_x:
            tx = ax_x + ox
            for oy in offsets_y_above:
                ty = ax_y - oy  # up in display = lower y
                px_l, px_r = tx - txt_w/2, tx + txt_w/2
                py_b, py_t = ty - txt_h/2, ty + txt_h/2
                if not _collides(px_l, px_r, py_b, py_t):
                    cost = oy + abs(ox) * 0.3
                    if cost < best_cost:
                        best = (tx, ty, px_l, px_r, py_b, py_t)
                        best_cost = cost
                    break
            for oy in offsets_y_below:
                ty = ax_y + oy  # down in display = higher y
                px_l, px_r = tx - txt_w/2, tx + txt_w/2
                py_b, py_t = ty - txt_h/2, ty + txt_h/2
                if not _collides(px_l, px_r, py_b, py_t):
                    cost = oy + 3 + abs(ox) * 0.3
                    if cost < best_cost:
                        best = (tx, ty, px_l, px_r, py_b, py_t)
                        best_cost = cost
                    break

        if best is None:
            tx, ty = ax_x, ax_y - 18
            best = (tx, ty, tx-txt_w/2, tx+txt_w/2, ty-txt_h/2, ty+txt_h/2)

        tx, ty, px_l, px_r, py_b, py_t = best
        placed_px.append((px_l, px_r, py_b, py_t))
        data_pos = inv_trans.transform((tx, ty))

        ax.annotate(lab['text'], xy=(lab['x'], lab['y']),
                    xytext=data_pos, ha='center', va='center',
                    fontsize=5.5, fontweight='bold', color=clr,
                    arrowprops=dict(arrowstyle='->', color='#1B5E20' if is_fp else '#999',
                                  lw=0.7, ls='--'),
                    zorder=8 if is_fp else 6)

    ax.set_xlabel('Frequence (Hz) — echelle logarithmique', fontsize=9, fontweight='bold')
    ax.set_ylabel('Amplitude (dB)', fontsize=9, fontweight='bold')
    subtitle = f' — {note_range}' if note_range else ''
    ax.set_title(f'{display}{subtitle}\nCarte spectrale vocalique (ordinario, N={n_samples})',
                 fontsize=10, fontweight='bold', color=family_color)
    ax.grid(True, alpha=0.15, zorder=0)
    ax.legend(fontsize=7, loc='lower left')
    ax.tick_params(labelsize=7)
    plt.tight_layout()

    out = os.path.join(OUT_IMG, f"{filename}.png")
    fig.savefig(out, dpi=150, bbox_inches='tight')
    plt.close(fig)
    return out


def compute_register_profiles(instrument_key, techs=('ordinario',), fp_band=(800,1800)):
    """
    Compute per-register formant profiles + carte spectrale data.
    Returns ordered list of (register_name, profile_dict) plus 'GLOBAL'.
    """
    specenv_groups = load_specenv_by_register(instrument_key, techs)
    spectrum_groups = load_spectrum_by_register(instrument_key, techs)
    if not specenv_groups:
        return []

    _freqs = np.arange(1024) * FREQ_RES

    def _analyze(envs, spectra_lin=None):
        mean_env = np.mean(envs, axis=0)
        lo = max(int(80/FREQ_RES),1); hi = min(int(6000/FREQ_RES), len(mean_env)-1)
        reg = list(mean_env[lo:hi]); mx = max(reg); th = mx - 25
        peaks = []
        for i in range(1, len(reg)-1):
            v = reg[i]
            if v < th: continue
            if v >= reg[i-1] and v >= reg[i+1]: peaks.append(((lo+i)*FREQ_RES, v))
        peaks.sort(key=lambda x: x[1], reverse=True)
        sel = []
        for f, a in peaks:
            if not any(abs(f-sf)<70 for sf,_ in sel): sel.append((f,a))
            if len(sel) >= 7: break
        sel.sort(key=lambda x: x[0])
        lo_b = int(fp_band[0]/FREQ_RES); hi_b = int(fp_band[1]/FREQ_RES)
        region = mean_env[lo_b:hi_b+1]
        linear = np.array([10**(v/10) for v in region])
        total = linear.sum()
        fp = sum(linear[i]*(lo_b+i)*FREQ_RES for i in range(len(linear)))/total if total > 0 else None
        cep_db = None
        if spectra_lin is not None and len(spectra_lin) > 0:
            mean_spec = np.mean(spectra_lin, axis=0)
            cep_log = cepstral_envelope(mean_spec, order=30)
            cep_db_raw = (20/np.log(10)) * cep_log
            cmask = (_freqs >= 200) & (_freqs <= 4000)
            if cmask.any() and len(cep_db_raw) == len(mean_env):
                offset = np.median(mean_env[cmask]) - np.median(cep_db_raw[cmask])
                cep_db = cep_db_raw + offset
        return mean_env, sel, fp, cep_db

    reg_order = REGISTERS.get(instrument_key, [])
    results = []

    for reg_name, _, _ in reg_order:
        if reg_name not in specenv_groups:
            continue
        envs = specenv_groups[reg_name]['data']
        notes = sorted(specenv_groups[reg_name]['notes'],
                       key=lambda n: _note_to_midi(n) or 0)
        spectra = spectrum_groups.get(reg_name, {}).get('data', [])
        mean_env, sel, fp, cep_db = _analyze(envs, spectra)
        results.append((reg_name, {
            'n': len(envs), 'notes': notes,
            'note_range': f'{notes[0]}-{notes[-1]}' if notes else '',
            'peaks': [(f, a) for f, a in sel],
            'fp': fp, 'mean_env': mean_env, 'cep_db': cep_db,
        }))

    # Global
    all_envs = [e for g in specenv_groups.values() for e in g['data']]
    all_spectra = [s for g in spectrum_groups.values() for s in g.get('data', [])]
    all_notes = sorted(set(n for g in specenv_groups.values() for n in g['notes']),
                       key=lambda n: _note_to_midi(n) or 0)
    mean_env, sel, fp, cep_db = _analyze(all_envs, all_spectra if all_spectra else None)
    results.append(('GLOBAL', {
        'n': len(all_envs), 'notes': all_notes,
        'note_range': f'{all_notes[0]}-{all_notes[-1]}' if all_notes else '',
        'peaks': [(f, a) for f, a in sel],
        'fp': fp, 'mean_env': mean_env, 'cep_db': cep_db,
    }))

    return results


def make_register_table_html(register_profiles):
    """Generate HTML table with per-register F1-F7 (freq+dB interleaved) + Fp."""
    html = '<table class="tech-table" style="font-size:0.8em;">\n'
    html += '<tr class="header"><th>Registre</th><th>N</th><th>Notes</th>'
    for i in range(1, 8):
        html += f'<th>F{i} Hz</th><th>dB</th>'
    html += '<th>Fp Hz</th></tr>\n'

    for label, od in register_profiles:
        if label == 'GLOBAL':
            continue
        html += f'<tr><td><strong>{label}</strong></td><td>{od["n"]}</td><td>{od["note_range"]}</td>'
        for i in range(7):
            if i < len(od['peaks']):
                f, a = od['peaks'][i]
                html += f'<td>{f:.0f}</td><td>{a:.1f}</td>'
            else:
                html += '<td>&mdash;</td><td>&mdash;</td>'
        fp_str = f'{od["fp"]:.0f}' if od['fp'] else '&mdash;'
        html += f'<td>{fp_str}</td></tr>\n'

    for label, od in register_profiles:
        if label != 'GLOBAL':
            continue
        html += f'<tr style="font-weight:bold;background:#e8eaf6;"><td>GLOBAL</td><td>{od["n"]}</td><td>{od["note_range"]}</td>'
        for i in range(7):
            if i < len(od['peaks']):
                f, a = od['peaks'][i]
                html += f'<td>{f:.0f}</td><td>{a:.1f}</td>'
            else:
                html += '<td>&mdash;</td><td>&mdash;</td>'
        fp_str = f'{od["fp"]:.0f}' if od['fp'] else '&mdash;'
        html += f'<td>{fp_str}</td></tr>\n'

    html += '</table>\n'
    return html


def generate_per_register_html(instrument_key, display, techs=("ordinario",),
                                fp_band=(800,1800), family_color="#333"):
    """Generate complete per-register HTML block: table + carte spectrale per register."""
    profiles = compute_register_profiles(instrument_key, techs=techs, fp_band=fp_band)
    if not profiles or len(profiles) <= 1:
        return "", []

    import unicodedata as _ud
    def _slug(s):
        s = _ud.normalize("NFD", s)
        s = "".join(c for c in s if _ud.category(c) != "Mn")
        s = s.lower()
        s = _re.sub(r"[^a-z0-9]+", "_", s)
        return s.strip("_")

    slug = _slug(display)
    html = ""
    images = []

    html += "<h4>Analyse par registre</h4>\n"
    html += make_register_table_html(profiles)

    html += "<h4>Cartes spectrales vocaliques par registre</h4>\n"
    for label, od in profiles:
        if label == "GLOBAL":
            fname = f"carte_{slug}_global"
        else:
            reg_slug = _slug(label)
            fname = f"carte_{slug}_{reg_slug}"

        img_path = make_carte_spectrale(
            display, fname, od["mean_env"], od["n"],
            fp_band=fp_band, family_color=family_color,
            cep_env_db=od.get("cep_db"),
            note_range=f'{label} ({od["note_range"]})'
        )
        if img_path:
            rel = os.path.relpath(img_path, OUT_DIR).replace(os.sep, "/")
            oct_label = label if label != "GLOBAL" else "Global (toute la tessiture)"
            html += f"<p><strong>{oct_label}</strong></p>\n"
            html += f'<img src="{rel}" alt="Carte spectrale {display} {label}"/>\n'
            images.append(img_path)

    return html, images


# Backward-compatible alias
generate_per_octave_html = generate_per_register_html
