#!/usr/bin/env python3
"""
build_volume_html_docx.py — Section VIII : Étude du Volume (Koechlin)
Volume, densité, transparence, homogénéité, plans orchestraux.

Lit les données de Etude-Volume/Resultats-volume/*.csv
et les intègre dans le document de référence.
"""
import os, sys, csv
import numpy as np

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from common import *

load_all_csvs()

# ── Chemins des données ──
VOL_DIR = os.path.join(BASE, 'Etude-Volume', 'Resultats-volume')

# ── Mapping noms anglais → français ──
NAME_FR = {
    'Piccolo':'Piccolo','Flute':'Flûte','Bass_Flute':'Flûte basse',
    'Contrabass_Flute':'Flûte contrebasse',
    'Oboe':'Hautbois','Oboe+sordina':'Hautbois sourd.',
    'English_Horn':'Cor anglais',
    'Clarinet_Eb':'Clarinette Mib','Clarinet_Bb':'Clarinette Sib',
    'Bass_Clarinet_Bb':'Clarinette basse','Contrabass_Clarinet_Bb':'Clarinette cb',
    'Bassoon':'Basson','Bassoon+sordina':'Basson sourd.','Contrebasson':'Contrebasson',
    'Sax_Alto':'Saxophone alto',
    'Horn':'Cor','Horn+sordina':'Cor sourd.',
    'Trumpet_C':'Trompette','Trumpet_C+sordina_straight':'Trp straight',
    'Trumpet_C+sordina_cup':'Trp cup','Trumpet_C+sordina_harmon':'Trp harmon',
    'Trumpet_C+sordina_wah':'Trp wah',
    'Trombone':'Trombone','Trombone+sordina_straight':'Tbn straight',
    'Trombone+sordina_cup':'Tbn cup','Trombone+sordina_harmon':'Tbn harmon',
    'Trombone+sordina_wah':'Tbn wah',
    'Bass_Trombone':'Trombone basse',
    'Bass_Tuba':'Tuba basse','Bass_Tuba+sordina':'Tuba basse sourd.',
    'Contrabass_Tuba':'Tuba contrebasse',
    'Guitar':'Guitare','Harp':'Harpe','Marimba':'Marimba','Vibraphone':'Vibraphone',
    'Violin':'Violon','Violin+sordina':'Violon sourd.',
    'Violin+sordina_piombo':'Violon piombo',
    'Violin_Ensemble':'Ens. violons','Violin_Ensemble+sordina':'Ens. vn sourd.',
    'Viola':'Alto','Viola+sordina':'Alto sourd.',
    'Viola+sordina_piombo':'Alto piombo',
    'Viola_Ensemble':'Ens. altos','Viola_Ensemble+sordina':'Ens. alt. sourd.',
    'Violoncello':'Violoncelle','Violoncello+sordina':'Vc sourd.',
    'Violoncello+sordina_piombo':'Vc piombo',
    'Violoncello_Ensemble':'Ens. violoncelles',
    'Violoncello_Ensemble+sordina':'Ens. vc sourd.',
    'Contrabass':'Contrebasse','Contrabass+sordina':'Cb sourd.',
    'Contrabass_Ensemble':'Ens. contrebasses',
}

def fr(name):
    return NAME_FR.get(name, name)

def load_csv(filename):
    path = os.path.join(VOL_DIR, filename)
    with open(path, newline='', encoding='utf-8-sig') as f:
        return list(csv.DictReader(f))

# ═══════════════════════════════════════════════════════════
# SORTABLE TABLE CSS + JS (réutilisable)
# ═══════════════════════════════════════════════════════════
SORTABLE_CSS = """
.sortable-table { border-collapse: collapse; width: 100%; font-size: 0.82em; }
.sortable-table th, .sortable-table td { border: 1px solid #ccc; padding: 4px 7px; text-align: center; }
.sortable-table th { background: #37474f; color: white; cursor: pointer; user-select: none; position: sticky; top: 0; z-index: 2; }
.sortable-table th:hover { background: #546e7a; }
.sortable-table tr:nth-child(even) { background: #fafafa; }
.sortable-table tr:hover { background: #e8f5e9; }
.sortable-table .sort-arrow { font-size: 0.7em; margin-left: 3px; }
.filter-input { width: 100%; box-sizing: border-box; padding: 3px 6px; margin: 6px 0;
                border: 1px solid #ccc; border-radius: 4px; font-size: 0.9em; }
.table-scroll {
  overflow: auto; border: 1px solid #ddd; border-radius: 6px; margin: 12px 0 20px 0;
  position: relative; display: block; clear: both;
}
.table-scroll + * { clear: both; }
.note-v4, .description, .doublures-box, .section-intro, .fp-explanation {
  position: relative; clear: both; display: block;
}
h2, h3 { clear: both; }
.cat-fondu { background: #c8e6c9 !important; }
.cat-semi  { background: #fff9c4 !important; }
.cat-conv  { background: #e1f5fe !important; }
.cat-het   { background: #fff !important; }
"""

SORTABLE_JS = """
<script>
function makeSortable(tableId) {
  const table = document.getElementById(tableId);
  if (!table) return;
  const headers = table.querySelectorAll('th');
  let sortCol = -1, sortAsc = true;
  headers.forEach((th, idx) => {
    th.addEventListener('click', () => {
      if (sortCol === idx) sortAsc = !sortAsc;
      else { sortCol = idx; sortAsc = true; }
      const tbody = table.querySelector('tbody');
      const rows = Array.from(tbody.querySelectorAll('tr'));
      rows.sort((a, b) => {
        let va = a.cells[idx].getAttribute('data-val') || a.cells[idx].textContent.trim();
        let vb = b.cells[idx].getAttribute('data-val') || b.cells[idx].textContent.trim();
        let na = parseFloat(va), nb = parseFloat(vb);
        if (!isNaN(na) && !isNaN(nb)) return sortAsc ? na - nb : nb - na;
        return sortAsc ? va.localeCompare(vb, 'fr') : vb.localeCompare(va, 'fr');
      });
      rows.forEach(r => tbody.appendChild(r));
      headers.forEach(h => { let s = h.querySelector('.sort-arrow'); if(s) s.textContent = ''; });
      let arrow = th.querySelector('.sort-arrow');
      if (!arrow) { arrow = document.createElement('span'); arrow.className='sort-arrow'; th.appendChild(arrow); }
      arrow.textContent = sortAsc ? ' ▲' : ' ▼';
    });
  });
}
function makeFilterable(inputId, tableId, colIdx) {
  const input = document.getElementById(inputId);
  const table = document.getElementById(tableId);
  if (!input || !table) return;
  const container = table.closest('.table-scroll');
  const origMaxH = container ? container.style.maxHeight : '';
  input.addEventListener('input', () => {
    const val = input.value.toLowerCase();
    const rows = table.querySelectorAll('tbody tr');
    let visCount = 0;
    rows.forEach(r => {
      let match = false;
      if (colIdx === -1) {
        match = r.textContent.toLowerCase().includes(val);
      } else {
        match = r.cells[colIdx].textContent.toLowerCase().includes(val);
      }
      r.style.display = match ? '' : 'none';
      if (match) visCount++;
    });
    // Shrink container to fit filtered content, restore on clear
    if (container) {
      if (val === '') {
        container.style.maxHeight = origMaxH;
      } else {
        // Let it auto-size: small result = small box, big result = capped
        const rowH = 28; // approximate row height
        const needed = Math.min(visCount * rowH + 50, window.innerHeight * 0.7);
        container.style.maxHeight = needed + 'px';
      }
    }
  });
}
</script>
"""


# ═══════════════════════════════════════════════════════════
# BUILD HTML
# ═══════════════════════════════════════════════════════════
def build_html(output_path):
    print("  → Section VIII : Étude du Volume (Koechlin)")

    # Load data
    vol_data = load_csv('volume_koechlin_v3.csv')
    homo_data = load_csv('homogeneite_matrix_v3.csv')
    plans_data = load_csv('plans_orchestraux_koechlin.csv')
    conv_data = load_csv('convergences_par_registre.csv')

    html = html_head("VIII. Étude du Volume — Koechlin", extra_css=SORTABLE_CSS)

    # ── 8.0 Titre ──
    html += '<h1 id="viii-volume">VIII. Étude du Volume — Attributs du timbre selon Koechlin</h1>\n'

    # ── 8.1 Introduction ──
    html += '<h2 id="vol-intro">Cadre théorique</h2>\n'
    html += """<div class="section-intro general">
<p>Charles Koechlin, dans son <em>Traité de l'orchestration</em> (1954), a proposé un système d'attributs généraux
du timbre — <strong>volume</strong>, <strong>intensité</strong>, <strong>transparence</strong>, <strong>densité</strong> —
applicable à tous les instruments de l'orchestre, permettant de prédire quels mélanges sonneront
<em>fondus</em> (plan orchestral homogène) ou <em>hétérogènes</em> (plans distincts superposés).</p>
<p>Le <strong>volume</strong> selon Koechlin n'est pas l'intensité sonore, mais la <em>place qu'un son semble occuper
dans l'espace</em> : une note grave de flûte est volumineuse mais faible, tandis qu'une note médium de hautbois
est mince mais intense. Koechlin illustre cette distinction par une métaphore astronomique : la flûte comme
une étoile géante de densité minimale, le hautbois comme une étoile naine dense
(Chiasson &amp; Duchesneau, 2008&nbsp;; Chiasson, 2010).</p>
<p>Cette section opérationnalise ces concepts par l'analyse spectrale de <strong>53 instruments</strong>
(sourdines comprises), totalisant <strong>5&nbsp;835 échantillons</strong> sur <strong>193 segments de registre</strong>.</p>
</div>\n"""

    # ── 8.2 Méthodologie ──
    html += '<h2 id="vol-methodo">Méthodologie : Indice de Volume composite</h2>\n'
    html += """<div class="fp-explanation">
<p>L'Indice de Volume combine quatre descripteurs spectraux, chacun capturant une facette
de la « taille » perçue d'un son :</p>
<table class="tech-table" style="font-size:0.88em; margin:10px 0;">
<tr class="header"><th>Composante</th><th>Formule</th><th>Signification</th></tr>
<tr><td><b>V₁</b></td><td>Spread / Centroïde</td><td>Largeur spectrale relative — normalisée par la hauteur</td></tr>
<tr><td><b>V₂</b></td><td>E<sub>&lt;1kHz</sub> / E<sub>total</sub></td><td>Ratio d'énergie basse fréquence — poids grave</td></tr>
<tr><td><b>V₃</b></td><td>1 / ΔF̄ (formants)</td><td>Dispersion formantique inverse — taille apparente du résonateur</td></tr>
<tr><td><b>V₄</b></td><td>MFCC₁</td><td>Pente spectrale mel — énergie concentrée en basses = timbre « chaud »</td></tr>
</table>
<p>Chaque composante est normalisée en z-score sur l'ensemble du corpus, puis moyennée :</p>
<p style="text-align:center; font-size:1.05em;"><b>𝒱 = ¼ (V̂₁ + V̂₂ + V̂₃ + V̂₄)</b></p>
<p><b>Densité</b> = −𝒱 (Stevens, 1965 : Loudness = Volume × Density, à intensité constante).</p>
<p>La corrélation de Spearman avec l'échelle empirique de Koechlin (vol. I, p. 288) atteint
<b>ρ = −0,76</b> (p &lt; 0,002, n = 17 instruments de base).</p>
</div>\n"""

    # ── 8.3 Tableau Volume Index (sortable) ──
    html += '<h2 id="vol-table">Indice de Volume par instrument et registre</h2>\n'
    html += '<p>53 instruments × 193 registres. Cliquez sur un en-tête pour trier. Utilisez le filtre pour chercher.</p>\n'
    html += '<input type="text" class="filter-input" id="vol-filter" placeholder="Filtrer par instrument...">\n'
    html += '<div class="table-scroll" style="max-height:70vh;">\n'
    html += '<table class="sortable-table" id="vol-table-data">\n'
    html += '<thead><tr><th>Instrument</th><th>Registre</th><th>N</th><th>Centroïde (Hz)</th>'
    html += '<th>V₁</th><th>V₂</th><th>V₃</th><th>V₄</th>'
    html += '<th>Volume 𝒱</th><th>Densité 𝒟</th></tr></thead>\n<tbody>\n'
    for r in vol_data:
        vi = r.get('Volume_index', '')
        di = r.get('Density_index', '')
        html += f'<tr><td>{fr(r["instrument"])}</td><td>{r["register"]}</td>'
        html += f'<td data-val="{r["n_samples"]}">{r["n_samples"]}</td>'
        html += f'<td data-val="{r["mean_centroid_hz"]}">{float(r["mean_centroid_hz"]):.0f}</td>'
        for v in ['V1_spread_centroid','V2_energy_1khz','V3_inv_formant_disp','V4_mfcc1']:
            val = r.get(v, '')
            html += f'<td data-val="{val}">{val}</td>'
        html += f'<td data-val="{vi}"><b>{vi}</b></td>'
        html += f'<td data-val="{di}">{di}</td></tr>\n'
    html += '</tbody></table></div>\n'

    # ── 8.4 Effets des sourdines ──
    html += '<h2 id="vol-mutes">Effet des sourdines sur le volume spectral</h2>\n'
    html += """<div class="description">
<p>Les sourdines ne font pas qu'atténuer — elles <em>restructurent</em> le profil spectral,
créant des déplacements considérables de l'Indice de Volume.</p>
</div>
<table class="tech-table" style="max-width:600px;">
<tr class="header"><th>Combinaison</th><th>Δ𝒱</th><th>Effet</th></tr>
<tr><td>Trompette + cup</td><td>−0,72</td><td>Amincit fortement</td></tr>
<tr><td>Trompette + harmon</td><td>−0,69</td><td>Amincit fortement</td></tr>
<tr><td>Trompette + straight</td><td>−0,35</td><td>Amincissement modéré</td></tr>
<tr><td>Trombone + harmon</td><td><b>+0,63</b></td><td><em>Augmente</em> le volume</td></tr>
<tr><td>Cor + sourdine</td><td>−0,08</td><td>Négligeable</td></tr>
<tr><td>Cordes + sourdine</td><td>≈ 0</td><td>Léger assombrissement</td></tr>
<tr><td>Cordes + piombo</td><td>−0,3 à −0,8</td><td>Amincissement prononcé</td></tr>
</table>
<div class="note-v4">
<b>Découverte :</b> La sourdine harmon du trombone <em>augmente</em> le volume spectral de +0,63 en supprimant
les harmoniques supérieurs tout en renforçant l'énergie basse fréquence. Cela contraste fortement
avec la harmon de trompette, qui produit l'effet inverse.
</div>\n"""

    # ── 8.5 Matrice d'homogénéité ──
    html += '<h2 id="vol-homo">Matrice d\'homogénéité (53 × 53)</h2>\n'
    html += """<div class="description">
<p>L'homogénéité H(A,B) entre deux instruments combine la proximité de volume et la similarité
du profil MFCC (coefficients 1–12, distance cosinus) :</p>
<p style="text-align:center;"><b>H(A,B) = 1 − 0,5·|Δ𝒱| / 𝒱<sub>étendue</sub> − 0,5·d<sub>cos</sub>(MFCC<sub>A</sub>, MFCC<sub>B</sub>)</b></p>
<p>H ≥ 0,85 → plan fondu probable · H &lt; 0,70 → timbres hétérogènes</p>
</div>\n"""

    # Build homogeneity matrix as a compact HTML heatmap
    instrs = [r[''] for r in homo_data]
    html += '<p>Chercher une paire :</p>\n'
    html += '<input type="text" class="filter-input" id="homo-filter" placeholder="Filtrer (ex: Cor, Trombone)...">\n'
    html += '<div class="table-scroll" style="max-height:80vh; overflow-x:auto; overflow-y:auto;">\n'
    html += '<table class="sortable-table" id="homo-table" style="font-size:0.7em; width:auto; min-width:900px;">\n'
    html += '<thead><tr><th></th>'
    for i in instrs:
        html += f'<th style="writing-mode:vertical-rl;transform:rotate(180deg);padding:2px;max-width:18px;">{fr(i)}</th>'
    html += '</tr></thead>\n<tbody>\n'
    for row in homo_data:
        name = row['']
        html += f'<tr><td style="text-align:left;white-space:nowrap;font-weight:bold;">{fr(name)}</td>'
        for i in instrs:
            v = float(row[i])
            if name == i:
                html += '<td style="background:#ddd;">—</td>'
            else:
                r_val = int(min(255, max(0, (1-v)*400)))
                g_val = int(min(255, max(0, v*300)))
                bg = f'rgb({255-g_val//2},{200+g_val//5},{255-g_val//2})' if v >= 0.85 else (
                     f'rgb(255,{200+int(v*55)},{180+int(v*75)})' if v >= 0.70 else '#fff')
                html += f'<td data-val="{v:.3f}" style="background:{bg};font-size:0.85em;" title="{fr(name)} × {fr(i)} = {v:.3f}">{v:.2f}</td>'
        html += '</tr>\n'
    html += '</tbody></table></div>\n'

    # ── Key homogeneity findings ──
    html += """<div class="note-v4">
<b>Résultats principaux :</b>
<ul style="margin:4px 0;">
<li><b>Cor + Trombone</b> (H = 0,99) : paire inter-familles la plus homogène — mélange archétypal Brahms/Wagner</li>
<li><b>Clarinette basse + Cor</b> (H = 0,99) : la cl. basse comme « pont » spectral bois–cuivres</li>
<li><b>Cor + Marimba</b> (H = 0,96) : F1 quasi identiques au registre médium — paire de musique contemporaine</li>
<li><b>Trompette + Violon</b> (H = 0,98) : confirme la doublure classique de la ligne mélodique</li>
<li><b>Hautbois : le marginal</b> (H &lt; 0,70 avec cuivres/cordes graves) — profil spectral distinct = voix soliste</li>
</ul>
</div>\n"""

    # ── 8.6 Plans orchestraux (filterable) ──
    html += '<h2 id="vol-plans">Plans orchestraux — Prédiction de fusion timbrale</h2>\n'
    html += """<div class="description">
<p>Le croisement <b>Homogénéité (Volume + MFCC) × Convergence formantique (ΔF₁ + ΔFp)</b>
classe chaque paire en :</p>
<ul style="margin:4px 0;">
<li><span style="background:#c8e6c9;padding:1px 6px;border-radius:3px;">★ FONDU</span> : H ≥ 0,80 <em>et</em> (ΔF₁ ≤ 30 Hz ou ΔFp ≤ 50 Hz)</li>
<li><span style="background:#fff9c4;padding:1px 6px;border-radius:3px;">● SEMI-FONDU</span> : H ≥ 0,70 + ΔF₁ ≤ 80, ou H ≥ 0,80 seul</li>
<li><span style="background:#e1f5fe;padding:1px 6px;border-radius:3px;">◆ CONVERGENT</span> : ΔF₁ ≤ 30 mais H &lt; 0,70 (même zone vocalique, volume différent)</li>
<li>○ HÉTÉROGÈNE : autres paires</li>
</ul>
</div>\n"""

    # Stats
    cats = {}
    for r in plans_data:
        c = r['category']
        cats[c] = cats.get(c, 0) + 1
    html += '<p><b>1 378 paires analysées :</b> '
    for c in sorted(cats):
        html += f'{cats[c]} {c} · '
    html += '</p>\n'

    html += '<input type="text" class="filter-input" id="plans-filter" placeholder="Filtrer (instrument, catégorie...)">\n'
    html += '<div class="table-scroll" style="max-height:70vh;">\n'
    html += '<table class="sortable-table" id="plans-table">\n'
    html += '<thead><tr><th>Instrument A</th><th>Instrument B</th>'
    html += '<th>F1 A</th><th>F1 B</th><th>ΔF1</th><th>Fp A</th><th>Fp B</th><th>ΔFp</th>'
    html += '<th>H</th><th>Fusion</th><th>Catégorie</th></tr></thead>\n<tbody>\n'
    for r in plans_data:
        cat = r['category']
        cls = 'cat-fondu' if 'FONDU' in cat else ('cat-semi' if 'SEMI' in cat else
              ('cat-conv' if cat == 'CONVERGENT' else 'cat-het'))
        html += f'<tr class="{cls}">'
        html += f'<td>{fr(r["instrument_A"])}</td><td>{fr(r["instrument_B"])}</td>'
        for k in ['F1_A_hz','F1_B_hz','delta_F1_hz','Fp_A_hz','Fp_B_hz','delta_Fp_hz']:
            v = r.get(k, '')
            html += f'<td data-val="{v}">{v}</td>'
        html += f'<td data-val="{r["Homogeneity_H"]}">{r["Homogeneity_H"]}</td>'
        html += f'<td data-val="{r["fusion_score"]}">{r["fusion_score"]}</td>'
        html += f'<td>{cat}</td></tr>\n'
    html += '</tbody></table></div>\n'

    # ── Top fusions box ──
    html += '<div class="doublures-box"><h4>Top 15 fusions inter-familles prédites</h4>\n'
    html += '<table class="doublures-table"><tr class="header"><th>A</th><th>B</th><th>H</th><th>ΔF1</th><th>ΔFp</th><th>Score</th></tr>\n'
    cross_fam = [r for r in plans_data if 'FONDU' in r['category']
                 and r['instrument_A'].split('+')[0] != r['instrument_B'].split('+')[0]
                 and r['instrument_A'].replace('_Ensemble','') != r['instrument_B'].replace('_Ensemble','')]
    cross_fam.sort(key=lambda x: -float(x['fusion_score']))
    for r in cross_fam[:15]:
        html += f'<tr><td>{fr(r["instrument_A"])}</td><td>{fr(r["instrument_B"])}</td>'
        html += f'<td>{r["Homogeneity_H"]}</td><td>{r["delta_F1_hz"]}</td>'
        html += f'<td>{r["delta_Fp_hz"]}</td><td><b>{r["fusion_score"]}</b></td></tr>\n'
    html += '</table></div>\n'

    # ── 8.7 Convergences par zone vocalique ──
    html += '<h2 id="vol-zones">Convergences formantiques par zone vocalique</h2>\n'

    # Zone stats
    zones = {}
    for r in conv_data:
        z = r['zone']
        zones[z] = zones.get(z, 0) + 1

    html += """<div class="description">
<p>1 220 convergences formantiques (ΔF₁ ≤ 30 Hz) entre registres croisés,
classées en cinq zones vocaliques IPA :</p>
</div>
<table class="tech-table" style="max-width:700px;">
<tr class="header"><th>Zone</th><th>Bande F1</th><th>N paires</th><th>%</th><th>Caractère orchestral</th></tr>\n"""
    total = sum(zones.values())
    zone_info = [
        ('/u/', '≤ 250 Hz', 'Fondation — plénitude matérielle'),
        ('/o/', '251–400 Hz', 'Cluster de convergence central'),
        ('/å/', '401–520 Hz', 'Chaleur de transition'),
        ('/a/', '521–800 Hz', 'Présence, projection'),
        ('/e/–/i/', '> 800 Hz', 'Brillance — cuivres avec sourdine'),
    ]
    for z, band, desc in zone_info:
        n = zones.get(z, 0)
        pct = f'{100*n/total:.1f}' if total > 0 else '0'
        html += f'<tr><td><b>{z}</b></td><td>{band}</td><td>{n}</td><td>{pct} %</td><td>{desc}</td></tr>\n'
    html += '</table>\n'

    html += f"""<div class="note-v4">
<b>79 % des convergences</b> se situent dans les zones /u/ et /o/ (F₁ ≤ 400 Hz),
confirmant que le fondement du fondu orchestral réside dans la région formantique basse fréquence —
le <em>sous-médium</em> de Koechlin, lieu de la <em>plénitude matérielle</em>.
</div>\n"""

    # Convergences table (filterable)
    html += '<h3 id="vol-conv-table">Tableau complet des convergences</h3>\n'
    html += '<input type="text" class="filter-input" id="conv-filter" placeholder="Filtrer (instrument, zone...)">\n'
    html += '<div class="table-scroll" style="max-height:70vh;">\n'
    html += '<table class="sortable-table" id="conv-table">\n'
    html += '<thead><tr><th>Instrument A</th><th>Registre A</th><th>Instrument B</th><th>Registre B</th>'
    html += '<th>F1 A</th><th>F1 B</th><th>ΔF1</th><th>ΔFp</th><th>H</th><th>Zone</th></tr></thead>\n<tbody>\n'
    for r in conv_data:
        z = r['zone']
        cls = 'cat-fondu' if z in ('/u/','/o/') else ('cat-semi' if z == '/å/' else '')
        html += f'<tr class="{cls}">'
        html += f'<td>{fr(r["instrument_A"])}</td><td>{r["register_A"]}</td>'
        html += f'<td>{fr(r["instrument_B"])}</td><td>{r["register_B"]}</td>'
        for k in ['F1_A_hz','F1_B_hz','delta_F1_hz','delta_Fp_hz','Homogeneity_H']:
            v = r.get(k, '')
            html += f'<td data-val="{v}">{v}</td>'
        html += f'<td>{z}</td></tr>\n'
    html += '</tbody></table></div>\n'

    # ── 8.8 Découvertes au-delà du classicisme ──
    html += '<h2 id="vol-discoveries">Découvertes au-delà du classicisme</h2>\n'
    html += """<div class="section-intro general">
<p>Le cadre de Koechlin, développé pour les instruments ouverts de son époque, acquiert une puissance
inattendue lorsqu'on l'étend aux sourdines et aux instruments non-classiques :</p>
<ul>
<li><b>Trombone harmon + Violoncelle piombo</b> (ΔFp = 12 Hz) : la sourdine harmon déplace le spectre
du trombone dans une zone qui converge précisément avec le violoncelle sourdine de plomb.</li>
<li><b>Tuba basse + Guitare</b> (ΔFp = 1 Hz) : la guitare fonctionne comme un « tuba pincé » en termes spectraux.</li>
<li><b>Cor + Marimba</b> (ΔF₁ = 0 Hz, H = 0,96) : F₁ identiques au registre médium.</li>
<li><b>Clarinette Mib + Flûte</b> (ΔF₁ = 0, ΔFp = 5 Hz) : convergence quasi-parfaite.</li>
<li><b>Flûte contrebasse + Alto sourdine</b> (ΔFp = 23 Hz) : doublure d'orchestre contemporain.</li>
<li><b>Zone /e/–/i/ (&gt; 800 Hz)</b> : cluster entièrement composé de sourdines cup, harmon et piccolo —
sans équivalent dans la théorie classique.</li>
</ul>
</div>\n"""

    # ── JavaScript for sorting/filtering ──
    html += SORTABLE_JS
    html += """<script>
makeSortable('vol-table-data');
makeSortable('homo-table');
makeSortable('plans-table');
makeSortable('conv-table');
makeFilterable('vol-filter', 'vol-table-data', -1);
makeFilterable('homo-filter', 'homo-table', 0);
makeFilterable('plans-filter', 'plans-table', -1);
makeFilterable('conv-filter', 'conv-table', -1);
</script>\n"""

    html += html_foot()

    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html)
    print(f"  ✓ HTML: {output_path}")


# ═══════════════════════════════════════════════════════════
# BUILD DOCX
# ═══════════════════════════════════════════════════════════
def build_docx(output_path):
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    print("  → Section VIII DOCX : Étude du Volume")
    doc = new_docx()

    # ── Title ──
    add_heading(doc, "VIII. Étude du Volume — Attributs du timbre selon Koechlin", level=1,
                color=RGBColor(26, 35, 126))

    # ── Intro ──
    add_heading(doc, "Cadre théorique", level=2, color=RGBColor(46, 125, 50))
    add_paragraph(doc,
        "Charles Koechlin, dans son Traité de l'orchestration (1954), a proposé un système d'attributs "
        "généraux du timbre — volume, intensité, transparence, densité — permettant de prédire quels "
        "mélanges sonneront fondus ou hétérogènes. Le volume selon Koechlin n'est pas l'intensité sonore, "
        "mais la place qu'un son semble occuper dans l'espace. Cette section opérationnalise ces concepts "
        "par l'analyse spectrale de 53 instruments (sourdines comprises), totalisant 5 835 échantillons "
        "sur 193 segments de registre.")

    # ── Methodology ──
    add_heading(doc, "Méthodologie : Indice de Volume composite", level=2, color=RGBColor(46, 125, 50))
    add_paragraph(doc,
        "L'Indice de Volume combine quatre descripteurs spectraux : V₁ = Spread/Centroïde (largeur "
        "spectrale relative), V₂ = énergie sous 1 kHz (poids grave), V₃ = dispersion formantique inverse "
        "(taille du résonateur), V₄ = MFCC₁ (pente spectrale mel). Chaque composante est normalisée "
        "en z-score puis moyennée. La corrélation de Spearman avec l'échelle de Koechlin atteint "
        "ρ = −0,76 (p < 0,002, n = 17 instruments).")

    # ── Volume table (top 25) ──
    add_heading(doc, "Indice de Volume — Extraits", level=2, color=RGBColor(46, 125, 50))
    vol_data = load_csv('volume_koechlin_v3.csv')
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for idx, label in enumerate(['Instrument','Registre','N','Centroïde','Volume 𝒱','Densité 𝒟']):
        set_cell_text(hdr[idx], label, bold=True, size=8)
        set_cell_shading(hdr[idx], '37474F')
        for run in hdr[idx].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255,255,255)

    # Show top 25 entries
    for r in vol_data[:25]:
        row = table.add_row().cells
        set_cell_text(row[0], fr(r['instrument']), size=7)
        set_cell_text(row[1], r['register'], size=7)
        set_cell_text(row[2], r['n_samples'], size=7)
        set_cell_text(row[3], f"{float(r['mean_centroid_hz']):.0f}", size=7)
        set_cell_text(row[4], r.get('Volume_index',''), bold=True, size=7)
        set_cell_text(row[5], r.get('Density_index',''), size=7)
    add_paragraph(doc, "(Tableau complet : 193 entrées dans le fichier CSV et la version HTML en ligne.)",
                  italic=True, size=8, color=RGBColor(128,128,128))

    # ── Mutes table ──
    add_heading(doc, "Effet des sourdines", level=2, color=RGBColor(46, 125, 50))
    mute_data = [
        ('Trompette + cup', '−0,72', 'Amincit fortement'),
        ('Trompette + harmon', '−0,69', 'Amincit fortement'),
        ('Trombone + harmon', '+0,63', 'Augmente le volume'),
        ('Cor + sourdine', '−0,08', 'Négligeable'),
        ('Cordes + piombo', '−0,3 à −0,8', 'Amincissement prononcé'),
    ]
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    for idx, label in enumerate(['Combinaison','Δ𝒱','Effet']):
        set_cell_text(t2.rows[0].cells[idx], label, bold=True, size=8)
    for combo, dv, eff in mute_data:
        row = t2.add_row().cells
        set_cell_text(row[0], combo, size=8)
        set_cell_text(row[1], dv, bold=True, size=8)
        set_cell_text(row[2], eff, size=8)

    # ── Plans summary ──
    add_heading(doc, "Plans orchestraux", level=2, color=RGBColor(46, 125, 50))
    plans_data = load_csv('plans_orchestraux_koechlin.csv')
    cats = {}
    for r in plans_data:
        cats[r['category']] = cats.get(r['category'], 0) + 1
    add_paragraph(doc, f"1 378 paires analysées (53 instruments × registre médium) :")
    for c in sorted(cats):
        add_paragraph(doc, f"  • {cats[c]} paires {c}", size=9)

    # ── Convergences zones ──
    add_heading(doc, "Convergences par zone vocalique", level=2, color=RGBColor(46, 125, 50))
    add_paragraph(doc,
        "1 220 convergences (ΔF₁ ≤ 30 Hz) réparties en 5 zones : "
        "/u/ (544, fondation), /o/ (423, cluster central), /å/ (125, transition), "
        "/a/ (54, projection), /e/–/i/ (74, brillance/sourdines). "
        "79 % des convergences se situent dans /u/ et /o/ (F₁ ≤ 400 Hz).")

    # ── Discoveries ──
    add_heading(doc, "Découvertes au-delà du classicisme", level=2, color=RGBColor(46, 125, 50))
    discoveries = [
        "Trombone harmon + Violoncelle piombo (ΔFp = 12 Hz)",
        "Tuba basse + Guitare (ΔFp = 1 Hz)",
        "Cor + Marimba (ΔF₁ = 0, H = 0,96)",
        "Clarinette Mib + Flûte (ΔF₁ = 0, ΔFp = 5 Hz)",
        "Zone /e/–/i/ : cluster de sourdines cup, harmon et piccolo",
    ]
    for d in discoveries:
        add_paragraph(doc, f"• {d}", size=9)

    doc.save(output_path)
    print(f"  ✓ DOCX: {output_path}")


# ═══════════════════════════════════════════════════════════
if __name__ == '__main__':
    html_path = os.path.join(OUT_DIR, 'section_volume_v6.html')
    docx_path = os.path.join(OUT_DIR, 'section_volume_v6.docx')
    build_html(html_path)
    build_docx(docx_path)
    print(f"\n{'='*60}")
    print(f"HTML : {html_path}")
    print(f"DOCX : {docx_path}")
